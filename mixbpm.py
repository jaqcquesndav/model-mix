import streamlit as st
import pandas as pd
from openai import OpenAI
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from markdown_pdf import MarkdownPdf, Section
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import re
import base64
import streamlit.components.v1 as components
from bs4 import BeautifulSoup
import os

# Initialisation de la persistance des donn√©es
def init_session_state():
    """Initialise les variables de session pour maintenir la persistance des donn√©es"""
    if 'initialized' not in st.session_state:
        st.session_state.initialized = True
        # Variables principales
        if 'nom_entreprise' not in st.session_state:
            st.session_state.nom_entreprise = ""
        if 'type_entreprise' not in st.session_state:
            st.session_state.type_entreprise = ""
        # Donn√©es collect√©es
        if 'persona_data' not in st.session_state:
            st.session_state.persona_data = {}
        if 'analyse_marche_data' not in st.session_state:
            st.session_state.analyse_marche_data = {}
        if 'facteurs_limitants_data' not in st.session_state:
            st.session_state.facteurs_limitants_data = {}
        if 'concurrence_data' not in st.session_state:
            st.session_state.concurrence_data = {}
        if 'problem_tree_data' not in st.session_state:
            st.session_state.problem_tree_data = {}
        if 'financial_data' not in st.session_state:
            st.session_state.financial_data = {}
        if 'business_model_precedent' not in st.session_state:
            st.session_state.business_model_precedent = ""

# Appel de l'initialisation
init_session_state()



hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
    """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Configuration de l'API OpenAI

api_key = st.secrets["API_KEY"]

if not api_key:
    st.error("üîë Cl√© API OpenAI requise. Veuillez la configurer dans les secrets Streamlit.")
    st.stop()
    
client = OpenAI(api_key=api_key)
# ----------------------------------------------------------------------------
# Business Model 
# ----------------------------------------------------------------------------

def collect_persona_pme():
    """
    Fonction pour s√©lectionner le type de persona.
    """
    type_persona = st.selectbox(
        "S√©lectionnez le type de Persona",
        ["B2C", "B2B", "M√©nage"],
        key="type_persona_selectbox"
    )
    st.write(f"**Type de Persona s√©lectionn√© :** {type_persona}")
    return type_persona

def collect_persona_details(type_persona):
    """
    Fonction pour collecter les d√©tails du persona en fonction du type s√©lectionn√©.
    """
    persona = {"type": type_persona}

    if type_persona == "B2C":
        st.subheader("Persona - B2C")
        
        # Donn√©es D√©mographiques B2C
        st.subheader("Donn√©es D√©mographiques")
        age = st.number_input("√Çge", min_value=18, max_value=100, value=st.session_state.get('b2c_age', 28), key="b2c_age")
        sexe = st.text_input("Sexe", placeholder="Homme, Femme, Autre", value=st.session_state.get('b2c_sexe', ''), key="b2c_sexe")
        localisation_detail = st.text_input(
            "Localisation G√©ographique (ex: Kinshasa, Lubumbashi, Goma)",
            value=st.session_state.get('b2c_localisation', 'Kinshasa, RDC'),
            key="b2c_localisation"
        )
        education = st.text_input(
            "Niveau d'√âducation",
            placeholder="Primaire, Secondaire, Sup√©rieur, Alphab√©tis√©, Non-alphab√©tis√©",
            value=st.session_state.get('b2c_education', ''),
            key="b2c_education"
        )
        profession = st.text_input("Profession", value=st.session_state.get('b2c_profession', 'Ex: Commer√ßant, Fonctionnaire'), key="b2c_profession")
        revenu_moyen = st.number_input("Revenu Moyen (FC)", min_value=0, step=50000, value=st.session_state.get('b2c_revenu', 300000), key="b2c_revenu")
        
        # Param√®tres Comportementaux B2C
        st.subheader("Param√®tres Comportementaux")
        sensibilite_prix = st.text_input(
            "Sensibilit√© au Prix",
            placeholder="Tr√®s sensible, Mod√©r√©ment sensible, Peu sensible",
            value=st.session_state.get('b2c_sensibilite_prix', ''),
            key="b2c_sensibilite_prix"
        )
        
        frequence_achat = st.text_input(
            "Fr√©quence d'Achat",
            placeholder="Quotidienne, Hebdomadaire, Mensuelle, Occasionnelle",
            value=st.session_state.get('b2c_frequence_achat', ''),
            key="b2c_frequence_achat"
        )
        
        volume_achat = st.text_input(
            "Volume d'Achat",
            placeholder="Petites quantit√©s, Quantit√©s moyennes, Gros volumes",
            value=st.session_state.get('b2c_volume_achat', ''),
            key="b2c_volume_achat"
        )
        if volume_achat and volume_achat not in ["Faible", "Moyen", "√âlev√©"]:
            st.warning("Veuillez entrer une valeur valide pour le Volume d'Achat : Faible, Moyen ou √âlev√©.")
        
        perception_qualite = st.text_area(
            "Perception de la Qualit√©",
            "D√©crivez la perception de la qualit√©...",
            key="b2c_perception_qualite"
        )
        utilisation_tech = st.text_area(
            "Utilisation Technologique",
            "D√©crivez l'utilisation technologique...",
            key="b2c_utilisation_tech"
        )
        acces_transport = st.text_area(
            "Accessibilit√© (Transport)",
            "D√©crivez l'accessibilit√© via le transport...",
            key="b2c_acces_transport"
        )
        temps_disponible = st.text_area(
            "Temps Disponible",
            "D√©crivez le temps disponible...",
            key="b2c_temps_disponible"
        )
        besoins_specifiques = st.text_area(
            "Besoins Sp√©cifiques",
            "D√©crivez les besoins sp√©cifiques...",
            key="b2c_besoins_specifiques"
        )
        motivations = st.text_area(
            "Motivations",
            "D√©crivez les motivations des clients...",
            key="b2c_motivations"
        )
        
        # Capacit√© d‚ÄôAdoption de l‚ÄôInnovation B2C
        st.subheader("Capacit√© d‚ÄôAdoption de l‚ÄôInnovation")
        familiarite_tech = st.text_area(
            "Familiarit√© avec certaines Technologies",
            "D√©crivez la familiarit√© technologique...",
            key="b2c_familiarite_tech"
        )
        ouverture_changement = st.text_input(
            "Ouverture au Changement",
            placeholder="Faible, Moyenne, √âlev√©e",
            key="b2c_ouverture_changement"
        )
        if ouverture_changement and ouverture_changement not in ["Faible", "Moyenne", "√âlev√©e"]:
            st.warning("Veuillez entrer une valeur valide pour l'Ouverture au Changement : Faible, Moyenne ou √âlev√©e.")
        
        barri√®res = st.text_area(
            "Barri√®res Psychologiques/Culturelles",
            "D√©crivez les barri√®res psychologiques ou culturelles...",
            key="b2c_barrieres"
        )
        
        # Compilation du persona B2C
        persona.update({
            "√¢ge": age,
            "sexe": sexe,
            "localisation": localisation_detail,
            "√©ducation": education,
            "profession": profession,
            "revenu_moyen": revenu_moyen,
            "sensibilite_prix": sensibilite_prix,
            "frequence_achat": frequence_achat,
            "volume_achat": volume_achat,
            "perception_qualite": perception_qualite,
            "utilisation_tech": utilisation_tech,
            "acces_transport": acces_transport,
            "temps_disponible": temps_disponible,
            "besoins_specifiques": besoins_specifiques,
            "motivations": motivations,
            "familiarite_tech": familiarite_tech,
            "ouverture_changement": ouverture_changement,
            "barrieres": barri√®res
        })
        
        # Sauvegarder dans le session state pour persistance
        st.session_state.persona_data = persona

    elif type_persona == "B2B":
        st.subheader("Persona - B2B")
        
        # Donn√©es D√©mographiques B2B
        st.subheader("Donn√©es D√©mographiques")
        taille_entreprise = st.text_input(
            "Taille de l'Entreprise",
            placeholder="PME, Grande Entreprise, Multinationale",
            key="b2b_taille_entreprise"
        )
        secteur_activite = st.text_input(
            "Secteur d'Activit√©",
            "Ex: Technologie, Sant√©",
            key="b2b_secteur_activite"
        )
        localisation_entreprise = st.text_input(
            "Localisation de l'Entreprise",
            value=st.session_state.get('b2b_localisation_entreprise', 'Kinshasa, RDC'),
            key="b2b_localisation_entreprise"
        )
        chiffre_affaires = st.number_input(
            "Chiffre d'Affaires (FC)",
            min_value=0,
            step=5000000,
            value=st.session_state.get('b2b_chiffre_affaires', 500000000),
            key="b2b_chiffre_affaires"
        )
        nombre_employes = st.number_input(
            "Nombre d'Employ√©s",
            min_value=1,
            step=1,
            value=st.session_state.get('b2b_nombre_employes', 25),
            key="b2b_nombre_employes"
        )
        
        # D√©cideurs et Influenceurs B2B
        st.subheader("D√©cideurs et Influenceurs")
        role_decideur = st.text_input(
            "R√¥le du D√©cideur",
            "Ex: Directeur des Achats",
            key="b2b_role_decideur"
        )
        influenceur = st.text_input(
            "Influenceurs Internes",
            "Ex: √âquipe IT, Marketing",
            key="b2b_influenceur"
        )
        
        # Param√®tres Comportementaux B2B
        st.subheader("Param√®tres Comportementaux")
        sensibilite_prix_b2b = st.text_input(
            "Sensibilit√© au Prix",
            placeholder="Faible, Moyenne, √âlev√©e",
            key="b2b_sensibilite_prix"
        )
        if sensibilite_prix_b2b and sensibilite_prix_b2b not in ["Faible", "Moyenne", "√âlev√©e"]:
            st.warning("Veuillez entrer une valeur valide pour la Sensibilit√© au Prix : Faible, Moyenne ou √âlev√©e.")
        
        cycle_achat = st.text_input(
            "Cycle d'Achat",
            placeholder="Long, Moyen, Court",
            key="b2b_cycle_achat"
        )
        if cycle_achat and cycle_achat not in ["Long", "Moyen", "Court"]:
            st.warning("Veuillez entrer une valeur valide pour le Cycle d'Achat : Long, Moyen ou Court.")
        
        volume_achat_b2b = st.text_input(
            "Volume d'Achat",
            placeholder="Faible, Moyen, √âlev√©",
            key="b2b_volume_achat"
        )
        if volume_achat_b2b and volume_achat_b2b not in ["Faible", "Moyen", "√âlev√©"]:
            st.warning("Veuillez entrer une valeur valide pour le Volume d'Achat : Faible, Moyen ou √âlev√©.")
        
        perception_qualite_b2b = st.text_area(
            "Perception de la Qualit√©",
            "D√©crivez la perception de la qualit√©...",
            key="b2b_perception_qualite"
        )
        besoins_specifiques_b2b = st.text_area(
            "Besoins Sp√©cifiques",
            "D√©crivez les besoins sp√©cifiques de l'entreprise...",
            key="b2b_besoins_specifiques"
        )
        motivations_b2b = st.text_area(
            "Motivations",
            "D√©crivez les motivations de l'entreprise...",
            key="b2b_motivations"
        )
        
        # Capacit√© d‚ÄôAdoption de l‚ÄôInnovation B2B
        st.subheader("Capacit√© d‚ÄôAdoption de l‚ÄôInnovation")
        familiarite_tech_b2b = st.text_area(
            "Familiarit√© avec certaines Technologies",
            "D√©crivez la familiarit√© technologique de l'entreprise...",
            key="b2b_familiarite_tech"
        )
        ouverture_changement_b2b = st.text_input(
            "Ouverture au Changement",
            placeholder="Faible, Moyenne, √âlev√©e",
            key="b2b_ouverture_changement"
        )
        if ouverture_changement_b2b and ouverture_changement_b2b not in ["Faible", "Moyenne", "√âlev√©e"]:
            st.warning("Veuillez entrer une valeur valide pour l'Ouverture au Changement : Faible, Moyenne ou √âlev√©e.")
        
        barri√®res_b2b = st.text_area(
            "Barri√®res Psychologiques/Culturelles",
            "D√©crivez les barri√®res psychologiques ou culturelles...",
            key="b2b_barrieres"
        )
        
        # Compilation du persona B2B
        persona.update({
            "taille_entreprise": taille_entreprise,
            "secteur_activite": secteur_activite,
            "localisation_entreprise": localisation_entreprise,
            "chiffre_affaires": chiffre_affaires,
            "nombre_employes": nombre_employes,
            "role_decideur": role_decideur,
            "influenceur": influenceur,
            "sensibilite_prix": sensibilite_prix_b2b,
            "cycle_achat": cycle_achat,
            "volume_achat": volume_achat_b2b,
            "perception_qualite": perception_qualite_b2b,
            "besoins_specifiques": besoins_specifiques_b2b,
            "motivations": motivations_b2b,
            "familiarite_tech": familiarite_tech_b2b,
            "ouverture_changement": ouverture_changement_b2b,
            "barrieres": barri√®res_b2b
        })

    elif type_persona == "M√©nage":
        st.subheader("Persona - M√©nage")
        
        # Donn√©es D√©mographiques M√©nage
        st.subheader("Donn√©es D√©mographiques")
        taille_menage = st.number_input(
            "Nombre de Personnes dans le M√©nage",
            min_value=1,
            step=1,
            value=4,
            key="menage_taille_menage"
        )
        revenu_menage = st.number_input(
            "Revenu Mensuel du M√©nage (FC)",
            min_value=0,
            step=25000,
            value=st.session_state.get('menage_revenu_menage', 400000),
            key="menage_revenu_menage"
        )
        localisation_menage = st.text_input(
            "Localisation G√©ographique (ex: Kinshasa, Gombe)",
            value=st.session_state.get('menage_localisation', 'Kinshasa, Gombe'),
            key="menage_localisation"
        )
        type_logement = st.text_input(
            "Type de Logement",
            placeholder="Villa, Appartement, Maison de passage, Parcelle, Studio",
            value=st.session_state.get('menage_type_logement', ''),
            key="menage_type_logement"
        )
        
        # Param√®tres Comportementaux M√©nage
        st.subheader("Param√®tres Comportementaux")
        sensibilite_prix_menage = st.text_input(
            "Sensibilit√© au Prix",
            placeholder="Faible, Moyenne, √âlev√©e",
            key="menage_sensibilite_prix"
        )
        if sensibilite_prix_menage and sensibilite_prix_menage not in ["Faible", "Moyenne", "√âlev√©e"]:
            st.warning("Veuillez entrer une valeur valide pour la Sensibilit√© au Prix : Faible, Moyenne ou √âlev√©e.")
        
        frequence_achat_menage = st.text_input(
            "Fr√©quence d'Achat",
            placeholder="Rarement, Mensuellement, Hebdomadairement",
            key="menage_frequence_achat"
        )
        if frequence_achat_menage and frequence_achat_menage not in ["Rarement", "Mensuellement", "Hebdomadairement"]:
            st.warning("Veuillez entrer une valeur valide pour la Fr√©quence d'Achat : Rarement, Mensuellement ou Hebdomadairement.")
        
        volume_achat_menage = st.text_input(
            "Volume d'Achat",
            placeholder="Faible, Moyen, √âlev√©",
            key="menage_volume_achat"
        )
        if volume_achat_menage and volume_achat_menage not in ["Faible", "Moyen", "√âlev√©"]:
            st.warning("Veuillez entrer une valeur valide pour le Volume d'Achat : Faible, Moyen ou √âlev√©.")
        
        perception_qualite_menage = st.text_area(
            "Perception de la Qualit√©",
            "D√©crivez la perception de la qualit√©...",
            key="menage_perception_qualite"
        )
        utilisation_tech_menage = st.text_area(
            "Utilisation Technologique",
            "D√©crivez l'utilisation technologique dans le m√©nage...",
            key="menage_utilisation_tech"
        )
        acces_transport_menage = st.text_area(
            "Accessibilit√© (Transport)",
            "D√©crivez l'accessibilit√© via le transport pour le m√©nage...",
            key="menage_acces_transport"
        )
        temps_disponible_menage = st.text_area(
            "Temps Disponible",
            "D√©crivez le temps disponible pour le m√©nage...",
            key="menage_temps_disponible"
        )
        besoins_specifiques_menage = st.text_area(
            "Besoins Sp√©cifiques",
            "D√©crivez les besoins sp√©cifiques du m√©nage...",
            key="menage_besoins_specifiques"
        )
        motivations_menage = st.text_area(
            "Motivations",
            "D√©crivez les motivations du m√©nage...",
            key="menage_motivations"
        )
        
        # Capacit√© d‚ÄôAdoption de l‚ÄôInnovation M√©nage
        st.subheader("Capacit√© d‚ÄôAdoption de l‚ÄôInnovation")
        familiarite_tech_menage = st.text_area(
            "Familiarit√© avec certaines Technologies",
            "D√©crivez la familiarit√© technologique du m√©nage...",
            key="menage_familiarite_tech"
        )
        ouverture_changement_menage = st.text_input(
            "Ouverture au Changement",
            placeholder="Faible, Moyenne, √âlev√©e",
            key="menage_ouverture_changement"
        )
        if ouverture_changement_menage and ouverture_changement_menage not in ["Faible", "Moyenne", "√âlev√©e"]:
            st.warning("Veuillez entrer une valeur valide pour l'Ouverture au Changement : Faible, Moyenne ou √âlev√©e.")
        
        barri√®res_menage = st.text_area(
            "Barri√®res Psychologiques/Culturelles",
            "D√©crivez les barri√®res psychologiques ou culturelles...",
            key="menage_barrieres"
        )
        
        # Compilation du persona M√©nage
        persona.update({
            "taille_menage": taille_menage,
            "revenu_menage": revenu_menage,
            "localisation_menage": localisation_menage,
            "type_logement": type_logement,
            "sensibilite_prix": sensibilite_prix_menage,
            "frequence_achat": frequence_achat_menage,
            "volume_achat": volume_achat_menage,
            "perception_qualite": perception_qualite_menage,
            "utilisation_tech": utilisation_tech_menage,
            "acces_transport": acces_transport_menage,
            "temps_disponible": temps_disponible_menage,
            "besoins_specifiques": besoins_specifiques_menage,
            "motivations": motivations_menage,
            "familiarite_tech": familiarite_tech_menage,
            "ouverture_changement": ouverture_changement_menage,
            "barrieres": barri√®res_menage
        })

    return persona




def collect_analyse_marche_pme():
    st.header("Analyse du March√© - PME")
    
    # Taille du March√©
    st.subheader("Taille du March√©")
    taille_marche = st.text_area("Taille du March√©", "D√©crivez la taille du march√©, les segments et la valeur totale.")
    
    # Segments du March√©
    st.subheader("Segments du March√©")
    segments_marche = st.text_area("Segments du March√©", "D√©crivez les segments du march√©...")
    
    # Valeur Totale du March√© ($)
    st.subheader("Valeur Totale du March√© ($)")
    valeur_totale = st.text_area("Valeur Totale du March√© ($)", "D√©crivez la valeur totale du march√©...")
    
    # Offres Concurrentes
    st.subheader("Offres Concurrentes")
    offres_concurrentes = st.text_area("Offres Concurrentes", "D√©crivez les offres concurrentes...")
    
    # Niveau de Satisfaction
    st.subheader("Niveau de Satisfaction")
    niveau_satisfaction = st.text_area("Niveau de Satisfaction", "D√©crivez le niveau de satisfaction...")
    
    # Tendances du March√©
    st.subheader("Tendances du March√©")
    tendances = st.text_area("Tendances du March√©", "D√©crivez les tendances du march√©...")
    
    # Innovations √âmergentes
    st.subheader("Innovations √âmergentes")
    innovations = st.text_area("Innovations √âmergentes", "D√©crivez les innovations √©mergentes...")
    
    # Comportements √âmergents
    st.subheader("Comportements √âmergents")
    comportements_emergents = st.text_area("Comportements √âmergents", "D√©crivez les comportements √©mergents...")
    
    analyse_marche = {
        "taille_marche": taille_marche,
        "segments_marche": segments_marche,
        "valeur_totale": valeur_totale,
        "offres_concurrentes": offres_concurrentes,
        "niveau_satisfaction": niveau_satisfaction,
        "tendances": tendances,
        "innovations": innovations,
        "comportements_emergents": comportements_emergents
    }
    
    return analyse_marche

def collect_facteurs_limitants_pme():
    st.header("Facteurs Limitants - PME")
    
    # Contraintes Technologiques
    st.subheader("Contraintes Technologiques")
    contraintes_techno = st.text_area("Contraintes Technologiques", "D√©crivez les contraintes technologiques...")
    
    # Contraintes √âconomiques
    st.subheader("Contraintes √âconomiques")
    contraintes_economiques = st.text_area("Contraintes √âconomiques", "D√©crivez les contraintes √©conomiques...")
    
    # Contraintes Culturelles
    st.subheader("Contraintes Culturelles")
    contraintes_culturelles = st.text_area("Contraintes Culturelles", "D√©crivez les contraintes culturelles...")
    
    # Contraintes Psychologiques et Physiologiques
    st.subheader("Contraintes Psychologiques et Physiologiques")
    contraintes_psych_phys = st.text_area("Contraintes Psychologiques et Physiologiques", "D√©crivez ces contraintes...")
    
    # Contraintes R√©glementaires
    st.subheader("Contraintes R√©glementaires")
    contraintes_reglementaires = st.text_area("Contraintes R√©glementaires", "D√©crivez les contraintes r√©glementaires...")
    
    facteurs_limitants = {
        "contraintes_technologiques": contraintes_techno,
        "contraintes_economiques": contraintes_economiques,
        "contraintes_culturelles": contraintes_culturelles,
        "contraintes_psych_phys": contraintes_psych_phys,
        "contraintes_reglementaires": contraintes_reglementaires
    }
    
    return facteurs_limitants

def collect_concurrence_pme():
    st.header("√âvaluation de la Concurrence - PME")
    
    # Concurrents Directs
    concurrents_directs = st.text_area("Concurrents Directs", "Listez les concurrents directs...")
    
    # Concurrents Indirects
    concurrents_indirects = st.text_area("Concurrents Indirects", "Listez les concurrents indirects...")
    
    # Forces des Concurrents
    forces_concurrents = st.text_area("Forces des Concurrents", "D√©crivez les forces des concurrents...")
    
    # Faiblesses des Concurrents
    faiblesses_concurrents = st.text_area("Faiblesses des Concurrents", "D√©crivez les faiblesses des concurrents...")
    
    # Niveau de Satisfaction des Clients envers les Concurrents
    satisfaction_concurrence = st.slider("Satisfaction des Clients envers les Concurrents", 0, 10, 5)
    
    # Niveau de Confiance des Clients envers les Concurrents
    confiance_concurrence = st.slider("Confiance des Clients envers les Concurrents", 0, 10, 5)
    
    concurrence = {
        "concurrents_directs": concurrents_directs,
        "concurrents_indirects": concurrents_indirects,
        "forces_concurrents": forces_concurrents,
        "faiblesses_concurrents": faiblesses_concurrents,
        "satisfaction_concurrence": satisfaction_concurrence,
        "confiance_concurrence": confiance_concurrence
    }
    
    return concurrence


def collect_problem_tree_pme():
    st.header("Arbre √† Probl√®me")
    
    # 1. Contexte
    st.subheader("Contexte")
    contexte = st.text_area(
        "D√©crire le Contexte",
        placeholder="Exemple : Dans le domaine de la sant√© publique dans les zones rurales..."
    )
    
    st.markdown("---")  # S√©parateur visuel
    
    # 2. Probl√®me Principal
    st.subheader("Probl√®me Principal")
    probleme_principal = st.text_input(
        "D√©crire le Probl√®me Principal",
        placeholder="Exemple : Un acc√®s limit√© aux soins de sant√© de base."
    )
    
    st.markdown("---")  # S√©parateur visuel
    
    # 3. Causes Principales
    st.subheader("Causes Principales")
    cause1 = st.text_input(
        "Cause 1",
        placeholder="Exemple : Manque d'infrastructures m√©dicales..."
    )
    cause2 = st.text_input(
        "Cause 2",
        placeholder="Exemple : Faible financement gouvernemental..."
    )
    cause3 = st.text_input(
        "Cause 3 (Facultatif)",
        placeholder="Exemple : Isolement g√©ographique..."
    )
    
    st.markdown("---")  # S√©parateur visuel
    
    # 4. Impact
    st.subheader("Impact")
    impact = st.text_area(
        "D√©crire l'Impact",
        placeholder="Exemple : Augmentation de la mortalit√© infantile, retard de d√©veloppement √©conomique."
    )
    
    st.markdown("---")  # S√©parateur visuel
    
    # 5. Parties Prenantes
    st.subheader("Parties Prenantes")
    partie1 = st.text_input(
        "Partie Prenante 1",
        placeholder="Exemple : Gouvernement local..."
    )
    partie2 = st.text_input(
        "Partie Prenante 2",
        placeholder="Exemple : ONG..."
    )
    partie3 = st.text_input(
        "Partie Prenante 3 (Facultatif)",
        placeholder="Exemple : R√©sidents des zones rurales..."
    )
    
    st.markdown("---")  # S√©parateur visuel
    
    # 6. Opportunit√©s
    st.subheader("Opportunit√©s")
    opportunite1 = st.text_input(
        "Opportunit√© 1",
        placeholder="Exemple : Introduction de cliniques mobiles..."
    )
    opportunite2 = st.text_input(
        "Opportunit√© 2",
        placeholder="Exemple : Formation d'agents de sant√© communautaire..."
    )
    opportunite3 = st.text_input(
        "Opportunit√© 3 (Facultatif)",
        placeholder="Exemple : Partenariats avec des organisations internationales..."
    )
    
    # Collecte des donn√©es
    problem_tree = {
        "contexte": contexte,
        "probleme_principal": probleme_principal,
        "causes_principales": [cause for cause in [cause1, cause2, cause3] if cause.strip()],
        "impact": impact,
        "parties_prenantes": [partie for partie in [partie1, partie2, partie3] if partie.strip()],
        "opportunites": [opp for opp in [opportunite1, opportunite2, opportunite3] if opp.strip()]
    }
    
    return problem_tree



# ----------------------------------------------------------------------------
# 2) Fonctions de collecte des donn√©es pour Startups
# ----------------------------------------------------------------------------

def collect_persona_startup():
    st.header("Persona - Startup")
    
    # Donn√©es D√©mographiques
    st.subheader("Donn√©es D√©mographiques")
    age = st.number_input("√Çge", min_value=18, max_value=100, value=30)
    sexe = st.text_input("Sexe", value=st.session_state.get('startup_sexe', 'Homme/Femme/Autre'))
    localisation_detail = st.text_input("Localisation G√©ographique (ex: Kinshasa, Lubumbashi)", value=st.session_state.get('startup_localisation', 'Kinshasa, RDC'))
    education = st.text_input("Niveau d'√âducation", value=st.session_state.get('startup_education', 'Ex: Licence, Master'))
    profession = st.text_input("Profession", value=st.session_state.get('startup_profession', 'Ex: Entrepreneur, Commer√ßant'))
    revenu_moyen = st.number_input("Revenu Moyen (FC)", min_value=0, step=50000, value=st.session_state.get('startup_revenu', 250000))
    
    # Param√®tres Comportementaux
    st.subheader("Param√®tres Comportementaux")
    
    sensibilite_prix = st.text_input("Sensibilit√© au Prix", "D√©crivez la sensibilit√© au prix...")
    frequence_achat = st.text_input("Fr√©quence d'Achat", "D√©crivez la fr√©quence d'achat...")
    volume_achat = st.text_input("Volume d'Achat", "D√©crivez le volume d'achat...")
    perception_qualite = st.text_area("Perception de la Qualit√©", "D√©crivez la perception de la qualit√©...")
    utilisation_tech = st.text_area("Utilisation Technologique", "D√©crivez l'utilisation technologique...")
    acces_transport = st.text_area("Accessibilit√© (Transport)", "D√©crivez l'accessibilit√© via le transport...")
    temps_disponible = st.text_area("Temps Disponible", "D√©crivez le temps disponible...")
    besoins_specifiques = st.text_area("Besoins Sp√©cifiques", "D√©crivez les besoins sp√©cifiques...")
    motivations = st.text_area("Motivations", "D√©crivez les motivations des clients...")
    
    # Capacit√© d‚ÄôAdoption de l‚ÄôInnovation
    st.subheader("Capacit√© d‚ÄôAdoption de l‚ÄôInnovation")
    
    familiarite_tech = st.text_area("Familiarit√© avec certaines Technologies", "D√©crivez la familiarit√© technologique...")
    ouverture_changement = st.text_input("Ouverture au Changement", "Faible/Moyenne/√âlev√©e")
    barri√®res = st.text_area("Barri√®res Psychologiques/Culturelles", "D√©crivez les barri√®res psychologiques ou culturelles...")
    
    persona = {
        "√¢ge": age,
        "sexe": sexe,
        "localisation": localisation_detail,
        "√©ducation": education,
        "profession": profession,
        "revenu_moyen": revenu_moyen,
        "sensibilite_prix": sensibilite_prix,
        "frequence_achat": frequence_achat,
        "volume_achat": volume_achat,
        "perception_qualite": perception_qualite,
        "utilisation_tech": utilisation_tech,
        "acces_transport": acces_transport,
        "temps_disponible": temps_disponible,
        "besoins_specifiques": besoins_specifiques,
        "motivations": motivations,
        "familiarite_tech": familiarite_tech,
        "ouverture_changement": ouverture_changement,
        "barrieres": barri√®res
    }
    
    return persona

def collect_analyse_marche_startup():
    st.header("Analyse du March√© - Startup")
    
    # Taille du March√©
    st.subheader("Taille du March√©")
    taille_marche = st.text_area("Taille du March√©", "D√©crivez la taille du march√©, les segments et la valeur totale.")
    
    # Segments du March√©
    st.subheader("Segments du March√©")
    segments_marche = st.text_area("Segments du March√©", "D√©crivez les segments du march√©...")
    
    # Valeur Totale du March√© ($)
    st.subheader("Valeur Totale du March√© ($)")
    valeur_totale = st.text_area("Valeur Totale du March√© ($)", "D√©crivez la valeur totale du march√©...")
    
    # Offres Concurrentes
    st.subheader("Offres Concurrentes")
    offres_concurrentes = st.text_area("Offres Concurrentes", "D√©crivez les offres concurrentes...")
    
    # Niveau de Satisfaction
    st.subheader("Niveau de Satisfaction")
    niveau_satisfaction = st.text_area("Niveau de Satisfaction", "D√©crivez le niveau de satisfaction...")
    
    # Tendances du March√©
    st.subheader("Tendances du March√©")
    tendances = st.text_area("Tendances du March√©", "D√©crivez les tendances du march√©...")
    
    # Innovations √âmergentes
    st.subheader("Innovations √âmergentes")
    innovations = st.text_area("Innovations √âmergentes", "D√©crivez les innovations √©mergentes...")
    
    # Comportements √âmergents
    st.subheader("Comportements √âmergents")
    comportements_emergents = st.text_area("Comportements √âmergents", "D√©crivez les comportements √©mergents...")
    
    analyse_marche = {
        "taille_marche": taille_marche,
        "segments_marche": segments_marche,
        "valeur_totale": valeur_totale,
        "offres_concurrentes": offres_concurrentes,
        "niveau_satisfaction": niveau_satisfaction,
        "tendances": tendances,
        "innovations": innovations,
        "comportements_emergents": comportements_emergents
    }
    
    return analyse_marche

def collect_facteurs_limitants_startup():
    st.header("Facteurs Limitants - Startup")
    
    # Contraintes Technologiques
    st.subheader("Contraintes Technologiques")
    contraintes_techno = st.text_area("Contraintes Technologiques", "D√©crivez les contraintes technologiques...")
    
    # Contraintes √âconomiques
    st.subheader("Contraintes √âconomiques")
    contraintes_economiques = st.text_area("Contraintes √âconomiques", "D√©crivez les contraintes √©conomiques...")
    
    # Contraintes Culturelles
    st.subheader("Contraintes Culturelles")
    contraintes_culturelles = st.text_area("Contraintes Culturelles", "D√©crivez les contraintes culturelles...")
    
    # Contraintes Psychologiques et Physiologiques
    st.subheader("Contraintes Psychologiques et Physiologiques")
    contraintes_psych_phys = st.text_area("Contraintes Psychologiques et Physiologiques", "D√©crivez ces contraintes...")
    
    # Contraintes R√©glementaires
    st.subheader("Contraintes R√©glementaires")
    contraintes_reglementaires = st.text_area("Contraintes R√©glementaires", "D√©crivez les contraintes r√©glementaires...")
    
    facteurs_limitants = {
        "contraintes_technologiques": contraintes_techno,
        "contraintes_economiques": contraintes_economiques,
        "contraintes_culturelles": contraintes_culturelles,
        "contraintes_psych_phys": contraintes_psych_phys,
        "contraintes_reglementaires": contraintes_reglementaires
    }
    
    return facteurs_limitants

def collect_concurrence_startup():
    st.header("√âvaluation de la Concurrence - Startup")
    
    # Concurrents Directs
    concurrents_directs = st.text_area("Concurrents Directs", "Listez les concurrents directs...")
    
    # Concurrents Indirects
    concurrents_indirects = st.text_area("Concurrents Indirects", "Listez les concurrents indirects...")
    
    # Forces des Concurrents
    forces_concurrents = st.text_area("Forces des Concurrents", "D√©crivez les forces des concurrents...")
    
    # Faiblesses des Concurrents
    faiblesses_concurrents = st.text_area("Faiblesses des Concurrents", "D√©crivez les faiblesses des concurrents...")
    
    # Niveau de Satisfaction des Clients envers les Concurrents
    satisfaction_concurrence = st.slider("Satisfaction des Clients envers les Concurrents", 0, 10, 5)
    
    # Niveau de Confiance des Clients envers les Concurrents
    confiance_concurrence = st.slider("Confiance des Clients envers les Concurrents", 0, 10, 5)
    
    concurrence = {
        "concurrents_directs": concurrents_directs,
        "concurrents_indirects": concurrents_indirects,
        "forces_concurrents": forces_concurrents,
        "faiblesses_concurrents": faiblesses_concurrents,
        "satisfaction_concurrence": satisfaction_concurrence,
        "confiance_concurrence": confiance_concurrence
    }
    
    return concurrence

# ----------------------------------------------------------------------------
# 3) Fonctions pour appeler ChatGPT et g√©n√©rer le Business Model Canvas
# ----------------------------------------------------------------------------

def get_metaprompt(type_entreprise):
    """
    Retourne un metaprompt sp√©cifique bas√© sur le type d'entreprise.
    """
    metaprompts = {
        "PME": """**M√©ta-Prompt pour l‚Äô√âlaboration d‚Äôun Business Model pour PME Traditionnelle (Int√©grant des Innovations Low-Tech et Adapt√©es aux Contextes Africains ou √âmergents)**

        **Votre R√¥le :**  
        Vous √™tes un expert en strat√©gie d‚Äôentreprise, marketing, UX, innovation frugale (low-tech et √©ventuellement high-tech), et √©laboration de Business Models. Vous devez g√©n√©rer un Business Model complet, clair, chiffr√©, coh√©rent et innovant, adapt√© √† une PME qui op√®re dans un environnement local (par exemple en Afrique ou dans d‚Äôautres pays √©mergents) o√π les r√©alit√©s technologiques, √©conomiques, culturelles et r√©glementaires diff√®rent des contextes occidentaux fortement num√©ris√©s.  
        L‚Äôinnovation ne sera pas seulement technologique de pointe (high-tech), mais aussi low-tech (solutions simples, robustes, faciles d‚Äôentretien, peu consommatrices de ressources), et tenant compte des infrastructures limit√©es, des pr√©f√©rences culturelles, de la disponibilit√© intermittente de l‚Äô√©lectricit√©, du co√ªt de la connectivit√©, de l‚Äôimportance du lien social, etc.

        Votre t√¢che s‚Äôorganise en trois phases :  
        1. Configuration Initiale (Collecte et Structuration des Donn√©es)  
        2. √âtapes Interm√©diaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)  
        3. Production Finale (Business Model Canvas)

        Pour chaque phase, suivez les instructions et veillez √† :  
        - Prendre en compte la persona (donn√©es d√©mographiques, comportementales, capacit√©s d‚Äôadoption de l‚Äôinnovation)
        - Pendre en compte l'arbre √† problemes(Probl√®me Principal,Causes Principales,Impact,Parties Prenantes , Opportunit√©s)  
        - Analyser le march√© (taille, segments, offres existantes formelles et informelles, niveau de satisfaction, tendances locales, disponibilit√© de ressources, logistique)  
        - Int√©grer les facteurs limitants (technologiques, √©conomiques, culturels, psychologiques, physiologiques, r√©glementaires, infrastructures limit√©es)  
        - √âvaluer la concurrence (locale, informelle, substituts traditionnels), comprendre les niveaux de satisfaction et de confiance  
        - Comprendre le parcours client (avant, pendant, apr√®s), int√©grer la carte d‚Äôempathie, identifier les gains et souffrances sp√©cifiques au contexte (par exemple : importance du bouche-√†-oreille, confiance interpersonnelle, exigence de robustesse, maintenance locale)  
        - V√©rifier syst√©matiquement la coh√©rence, proposer des optimisations et ajustements  
        - Avant d‚Äôintroduire une innovation (low-tech ou high-tech), s‚Äôassurer que la persona est pr√™te √† l‚Äôadopter, en tenant compte de l‚Äôaccessibilit√©, du co√ªt, de la simplicit√© et de la r√©putation  
        - Produire un Business Model Canvas complet (9 blocs), avec des m√©ta-prompts sp√©cifiques pour chacun des blocs, adapt√©s au contexte local

        ---

        ### Phase 1 : Configuration Initiale (Entr√©e de Donn√©es)

        1. **Recueille et structure les informations suivantes :**  
        - **Persona** :  
            - Donn√©es d√©mographiques : √¢ge, sexe, localisation (zones urbaines, p√©ri-urbaines, rurales), niveau d‚Äô√©ducation (alphab√©tisation, langues parl√©es), profession (artisans, commer√ßants, agriculteurs, employ√©s, ind√©pendants), revenu moyen.  
            - Param√®tres comportementaux : sensibilit√© au prix (budgets limit√©s, n√©cessit√© de micro-paiements), fr√©quence et volume d‚Äôachat (achats ponctuels, saisonniers, hebdomadaires), perception de la qualit√© (fiabilit√©, durabilit√©), utilisation technologique (t√©l√©phones basiques, smartphones d‚Äôentr√©e de gamme, acc√®s limit√© √† Internet), accessibilit√© (distance aux points de vente, transport limit√©), temps disponible (horaires de travail, saison des r√©coltes), besoins sp√©cifiques (ex : acc√®s √† l‚Äôeau, √©nergie, outils agricoles, services financiers de base, √©ducation des enfants, soins de sant√©).  
            - Capacit√© d‚Äôadoption de l‚Äôinnovation : Familiarit√© avec certaines technologies (mobile money, radios communautaires, solutions solaires), ouverture au changement d√©pendant de la preuve sociale, de la confiance dans la communaut√©, de la simplicit√© et robustesse du produit/service. Barri√®res psychologiques/culturelles (m√©fiance envers les nouvelles solutions √©trang√®res, pr√©f√©rence pour le contact humain, importance de la recommandation de la famille ou du chef de village).  
        
        - **Arbre √† Probl√®me** :
            - Contexte:Description g√©n√©rale du domaine ou de la situation actuelle (secteur d'activit√©, environnement g√©ographique, tendances actuelles du march√©), Facteurs externes influen√ßant la situation (r√©glementations, conditions √©conomiques, technologies √©mergentes),Facteurs internes pertinents (ressources disponibles, comp√©tences cl√©s, structure organisationnelle).
            - Probl√®me Principal : Identification du d√©fi ou de l'obstacle central (nature du probl√®me, circonstances sp√©cifiques),Impact imm√©diat sur l'organisation ou le projet (effets sur les op√©rations, la performance financi√®re, la r√©putation).
            - Causes Principales :Causes internes contribuant au probl√®me (processus inefficaces, manque de comp√©tences, ressources limit√©es),Causes externes contribuant au probl√®me (concurrence accrue, changements de march√©, √©volutions technologiques),Interaction entre les causes internes et externes (comment elles se renforcent mutuellement).
            - Impact:Cons√©quences financi√®res du probl√®me (pertes de revenus, augmentation des co√ªts, rentabilit√© r√©duite),Effets op√©rationnels (d√©lai dans les projets, baisse de productivit√©, qualit√© des services ou produits affect√©e),Impact sur les parties prenantes (satisfaction des clients, moral des employ√©s, relations avec les partenaires).
            - Parties Prenantes :Identification des acteurs concern√©s ou impact√©s (clients, employ√©s, fournisseurs, investisseurs, communaut√© locale),Int√©r√™ts et attentes de chaque partie prenante vis-√†-vis du probl√®me (besoins sp√©cifiques, priorit√©s, pr√©occupations).
            - Opportunit√©s :Pistes d‚Äôam√©lioration ou de r√©solution du probl√®me (solutions innovantes, meilleures pratiques),Strat√©gies pour att√©nuer les causes principales (formation, r√©organisation, investissement technologique),Actions pour maximiser les impacts positifs (exploitation des forces, diversification, partenariats strat√©giques),

        - **Analyse du March√©** :  
            - Taille du march√© local : estimer la population concern√©e, le pouvoir d‚Äôachat moyen, les infrastructures disponibles.  
            - Segments : populations urbaines vs rurales, artisans, commer√ßants, coop√©ratives, PME locales, secteur informel.  
            - Offres concurrentes existantes : solutions traditionnelles (artisanales, informelles), importations bas de gamme, programmes d‚ÄôONG, concurrents locaux ou √©trangers, mod√®les low-cost.  
            - Niveau de satisfaction actuel : Les clients sont-ils satisfaits des solutions actuelles ? Y a-t-il un manque de fiabilit√©, de formation, de SAV ?  
            - Tendances : adoption progressive du mobile money, sensibilisation croissante √† l‚Äô√©nergie solaire, √©mergence de petites coop√©ratives, engouement pour des solutions durables et r√©parables.  
            - Innovations et comportements √©mergents : r√©emploi, √©conomie circulaire, mise en commun de ressources, augmentation des transferts d‚Äôargent via mobile.  
        
        - **Facteurs Limitants** :  
            - Contraintes technologiques : faible acc√®s √† l‚Äô√©lectricit√© stable, couverture internet in√©gale, outils technologiques rudimentaires, importance de solutions low-tech (pompes manuelles, panneaux solaires simples, syst√®mes de filtration d‚Äôeau basiques).  
            - Contraintes √©conomiques : revenus limit√©s, volatilit√© des prix, acc√®s restreint au cr√©dit, n√©cessit√© d‚Äô√©taler les paiements (micro-paiements, cr√©dit rotatif, tontines).  
            - Contraintes culturelles : langues locales, importance de la confiance interpersonnelle, r√©ticence √† adopter des produits inconnus sans d√©monstration ou validation par la communaut√©.  
            - Contraintes psychologiques et physiologiques : besoin de solutions simples d‚Äôutilisation, ergonomiques, adapt√©es aux conditions climatiques (chaleur, poussi√®re), faible taux d‚Äôalphab√©tisation n√©cessitant des modes d‚Äôemploi visuels.  
            - Contraintes r√©glementaires : normes locales, barri√®res douani√®res, absence de normes formelles dans certains secteurs, difficult√© √† obtenir des certifications officielles.  

        **Apr√®s avoir recueilli ces donn√©es, effectue une premi√®re analyse critique** :  
        - V√©rifie la coh√©rence des informations.  
        - Identifie les lacunes (par exemple, manque d‚Äôinformations sur le pouvoir d‚Äôachat r√©el, sur le r√©seau de distribution informel, sur le r√¥le des leaders d‚Äôopinion locaux).  
        - Propose des compl√©ments ou ajustements pour optimiser la qualit√© des donn√©es (ajouter des donn√©es sur la saisonnalit√© du march√©, l‚Äôinfluence des ONG, l‚Äôimpact des conditions climatiques, la pr√©sence ou non de microfinance).

        ---

        ### Phase 2 : √âtapes Interm√©diaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)

        2. **Analyse du Parcours Client & Carte d‚ÄôEmpathie** :  
        - D√©cris le parcours client (avant, pendant, apr√®s) en tenant compte des conditions locales :  
            - Avant : Le client prend conscience de son besoin par le bouche-√†-oreille, via un ami, un voisin, un leader communautaire, ou en √©coutant la radio locale. Il compare avec les solutions d√©j√† connues (artisan local, r√©parations informelles, solutions import√©es). Il √©value la confiance, le prix, la disponibilit√©.  
            - Pendant : Achat sur un march√© local, essai d‚Äôune d√©monstration concr√®te (d√©monstration en conditions r√©elles, sur une place de village), informations donn√©es par un vendeur itin√©rant ou un agent de confiance. Utilisation d‚Äôun mode de paiement adapt√© (cash, mobile money).  
            - Apr√®s : Suivi du produit, entretien, besoin de pi√®ces d√©tach√©es, possibilit√© de contact direct avec l‚Äôentreprise (ligne t√©l√©phonique, point de service local), √©change d‚Äôexp√©riences avec d‚Äôautres utilisateurs, √©ventuel SAV simplifi√© (r√©parations locales, pi√®ces d√©tach√©es bon march√©).  
        
        - Identifie les points de contact (march√©s, boutiques, interm√©diaires locaux, radios communautaires, SMS informatifs), obstacles (faible connectivit√©, manque d‚Äôinformations d√©taill√©es, barri√®res linguistiques), moments de v√©rit√© (premier essai du produit, premi√®re panne et r√©activit√© du SAV), frustrations (produit pas adapt√©, manuel incompr√©hensible, manque de fiabilit√©).  
        
        - Int√®gre les contraintes physiologiques, psychologiques, √©conomiques, culturelles, technologiques, r√©glementaires : par exemple, l‚Äôimportance de la simplicit√© et de la robustesse pour r√©duire la crainte d‚Äôune technologie trop complexe, la n√©cessit√© de support en langue locale, la possibilit√© de s‚Äôadapter aux normes informelles.  
        
        - Cr√©e une carte d‚Äôempathie :  
            - Pens√©es : ¬´ Est-ce que cette solution est fiable, reconnue par ma communaut√© ? Est-ce que je vais perdre mon argent si √ßa ne marche pas ? ¬ª  
            - Sentiments : M√©fiance, curiosit√©, besoin de r√©assurance, fiert√© s‚Äôil s‚Äôagit d‚Äôune innovation locale valoris√©e.  
            - Actions : Demande de conseils √† d‚Äôautres, observation d‚Äôexemples concrets, volont√© d‚Äôessayer avant d‚Äôacheter.

        3. **Gains et Souffrances** :  
        - Liste les gains : par exemple, acc√®s facilit√© √† un service vital (eau, √©nergie, outil de gestion commerciale simple), r√©duction du temps et de l‚Äôeffort, robustesse (moins de pannes), acc√®s √† un SAV local, meilleure rentabilit√© ou productivit√©.  
        - Liste les souffrances : manque de solutions adapt√©es, probl√®mes de maintenance, co√ªts initiaux trop √©lev√©s sans option de paiement flexible, manque de formation pour utiliser correctement le produit.

        4. **√âlaboration de la Carte de Valeur** :  
        - D√©finis la mission de consommation principale : r√©pondre √† un besoin fondamental (ex : un outil agricole robuste, une solution d‚Äô√©clairage solaire fiable, un service financier simple via mobile, un appareil domestique low-tech adapt√© aux pannes d‚Äô√©lectricit√©).  
        - Identifie les gains d√©j√† fournis par les offres actuelles (ex : disponibilit√© locale, prix bas) et les souffrances non adress√©es (faible qualit√©, pas de SAV, pas d‚Äôadaptation aux conditions r√©elles).  
        - Esquisse une proposition de valeur pr√©liminaire adapt√©e √† la capacit√© d‚Äôadoption de l‚Äôinnovation par la persona :  
            - Une solution simple, robuste, facilement compr√©hensible, qui peut √™tre test√©e avant achat.  
            - Un mod√®le de distribution local (agents sur le terrain), un SAV accessible, un support en langue locale, des options de paiement flexible (mobile money, tontines, microcr√©dit).  
            - Int√©gration progressive d‚Äôinnovations low-tech (p. ex. appareils m√©caniques robustes, panneaux solaires portables) ou high-tech simple (SMS, USSD, application mobile l√©g√®re) si l‚Äôutilisateur est pr√™t.

        5. **D√©termination du Segment de Clients** :  
        - Choisis le type de relation (B2C direct, B2B via des coop√©ratives, B2B2C via des distributeurs locaux).  
        - Priorise les segments qui correspondent le mieux :  
            - Par exemple, petits commer√ßants urbains ayant un pouvoir d‚Äôachat limit√© mais stables, agriculteurs n√©cessitant un outil fiable en milieu rural, coop√©ratives d‚Äôartisans pr√™ts √† adopter une solution pour am√©liorer leur productivit√©.  
        - Tient compte de leur sensibilit√© au prix, de leur ouverture √† l‚Äôinnovation, de leur capacit√© √† comprendre et utiliser la solution, de la n√©cessit√© de formation.

        6. **Analyse des Probl√®mes et Solutions (Canvas de Probl√®me)** :  
        - Identifie les probl√®mes majeurs : par exemple, la difficult√© √† acc√©der √† un produit fiable, le manque d‚Äôinformations, la complexit√© du produit, le co√ªt trop √©lev√© d‚Äôune solution import√©e haut de gamme.  
        - Associe chaque probl√®me √† une solution :  
            - Probl√®me : manque de SAV ‚Üí Solution : r√©seau de r√©parateurs locaux form√©s.  
            - Probl√®me : prix √©lev√© d‚Äôentr√©e ‚Üí Solution : offres en micro-paiements, location-vente, partenariats avec microfinance.  
            - Probl√®me : manque de confiance ‚Üí Solution : d√©monstrations, t√©moignages de pairs, communication via radios locales et leaders d‚Äôopinion.  
        - Justifie en quoi les solutions sont meilleures que l‚Äôexistant : plus adapt√©es, plus abordables, plus simples, prenant en compte la r√©alit√© du terrain (faible infrastructure, besoin de r√©silience, faible taux d‚Äôalphab√©tisation).

        **Apr√®s ces √©tapes, fais une analyse interm√©diaire** :  
        - V√©rifie la coh√©rence du contexte, du parcours client, des solutions propos√©es.  
        - Assure-toi que les innovations (low-tech, partenariats locaux, solutions de paiement flexible) sont compr√©hensibles et adoptables par la persona.  
        - Propose des ajustements strat√©giques : simplification du produit, ajustement du prix, ajout d‚Äôun canal de distribution plus local, formation des utilisateurs, partenariats avec des ONG ou des radios locales.

        ---

        ### Phase 3 : Production Finale du Business Model (Business Model Canvas)

        Sur la base des analyses pr√©c√©dentes, g√©n√®re un Business Model Canvas complet. Utilise les m√©ta-prompts suivants pour chaque bloc, en tenant compte du contexte local, des solutions low-tech et des infrastructures limit√©es :

        1. **Segments de Clients**  
        M√©ta-Prompt :  
        ¬´ D√©finis pr√©cis√©ment les segments de clients cibl√©s, en tenant compte :  
        - De leurs caract√©ristiques sociod√©mographiques (√¢ge, sexe, localisation, niveau d‚Äô√©ducation, profession, revenu, langue).  
        - De leurs comportements d‚Äôachat (fr√©quence, volume, sensibilit√© au prix, recours au cr√©dit informel, canaux de confiance : march√©s locaux, revendeurs informels, chefs de village, radios).  
        - De leur maturit√© technologique (t√©l√©phones basiques, usage de SMS/USSD, familiarit√© avec le mobile money, radio, bouche-√†-oreille, rencontres physiques).  
        - De leur capacit√© d‚Äôadoption de l‚Äôinnovation (ouverture au changement si d√©monstration concr√®te, barri√®res culturelles, besoin de preuves, pr√©f√©rences pour du low-tech robuste plut√¥t que du high-tech fragile).  
        - De leurs contraintes (faible pouvoir d‚Äôachat, saisons de r√©colte, temps de disponibilit√©, acc√®s difficile √† l‚Äô√©lectricit√© ou √† internet).  
        Int√®gre √©galement des sc√©narios √©volutifs (si l‚Äô√©conomie se d√©grade, r√©duction de l‚Äôachat ou passage √† des solutions plus frugales ; si la technologie progresse, adoption graduelle de services num√©riques simples).  
        Justifie pourquoi ces segments sont retenus : potentiel de rentabilit√©, facilit√© d‚Äôacc√®s via des canaux locaux, r√©ceptivit√© √† la proposition de valeur (am√©liorer leur vie de fa√ßon concr√®te, fiable, abordable). ¬ª

        2. **Proposition de Valeur**  
        M√©ta-Prompt :  
        ¬´ D√©taille la proposition de valeur en explicitant :  
        - Les besoins fondamentaux (eau, √©nergie, information, outils productifs, services financiers simples).  
        - Les souffrances clientes (manque de fiabilit√©, difficult√© d‚Äôentretien, complexit√© des produits, m√©fiance) et comment elles sont r√©solues (simplicit√©, robustesse, support local, preuves sociales).  
        - Les gains fournis (am√©lioration de la productivit√©, √©conomies de temps, durabilit√©, r√©duction de la d√©pendance √† des syst√®mes complexes, meilleure gestion financi√®re) et inclure les b√©n√©fices √©motionnels (confiance, fiert√©, reconnaissance sociale).  
        - La diff√©renciation par rapport aux offres concurrentes : int√©gration dans le tissu local, formation d‚Äôagents locaux, facilit√© d‚Äôentretien, pricing adapt√©, low-tech combin√© avec technologie simple (mobile money), SAV local.  
        - L‚Äôintroduction progressive de l‚Äôinnovation : d√©monstrations pratiques, formation sur le terrain, tutoriels en langue locale, partenariat avec leaders communautaires.  
        - Variantes selon les segments : option premium (un meilleur SAV, une maintenance plus pouss√©e) pour les clients plus solvables, version ultra-simplifi√©e pour les segments plus conservateurs ou √† tr√®s faible pouvoir d‚Äôachat. ¬ª

        3. **Canaux de Distribution**  
        M√©ta-Prompt :  
        ¬´ D√©finis les canaux par lesquels les clients seront inform√©s, convaincus, ach√®teront et utiliseront le produit/service, en tenant compte des r√©alit√©s locales :  
        - Canaux hors ligne : march√©s locaux, boutiques physiques, vente itin√©rante, radios communautaires, affichages, d√©monstrations sur place, coop√©ratives agricoles, leaders religieux ou communautaires.  
        - Canaux digitaux l√©gers : SMS, USSD, appels t√©l√©phoniques, WhatsApp, Facebook local, mobile money.  
        - N√©cessit√© d‚Äôomnicanalit√© adapt√©e au contexte : coh√©rence entre communication radio, d√©monstration physique, et suivi par t√©l√©phone.  
        - Simplicit√© d‚Äôacc√®s et besoin d‚Äôaccompagnement p√©dagogique (formation dans les march√©s, brochures visuelles, tutoriels audio).  
        - Adaptabilit√© des canaux si le march√© √©volue (ex: introduction progressive d‚Äôune application mobile si la connectivit√© s‚Äôam√©liore).  
        Justifie chaque canal (co√ªt, accessibilit√©, confiance) et comment il r√©duit les obstacles √† l‚Äôadoption, am√©liore la satisfaction, et s‚Äôint√®gre dans le parcours client local. ¬ª

        4. **Relations Clients**  
        M√©ta-Prompt :  
        ¬´ D√©cris la nature et la qualit√© des relations √©tablies avec les clients :  
        - Personnalisation via un r√©seau d‚Äôagents locaux qui connaissent la langue, la culture, et les besoins.  
        - Communaut√© : cr√©ation de groupes d‚Äôutilisateurs, d‚Äôassociations locales, de rencontres de d√©monstration, √©v√©nements communautaires o√π les clients √©changent leurs exp√©riences.  
        - Automatisation : mise en place d‚Äôun service SMS de rappel, d‚Äôune hotline t√©l√©phonique simple, d‚Äôun chatbot vocal si la technologie le permet (ou service d‚Äôappels humains en langue locale).  
        - Fid√©lisation : r√©ductions pour clients fid√®les, options de maintenance pr√©ventive, acc√®s √† des mises √† jour techniques simples, partenariats avec des ONG pour aider √† la formation continue.  
        - Gestion des plaintes et retours : politique claire de SAV, r√©paration locale, garantie adapt√©e, d√©lais de r√©ponse rapides.  
        Int√®gre la dimension culturelle (contact humain valoris√©), psychologique (confiance, besoin de r√©assurance), r√©glementaire (respect des r√®gles locales, si existantes). Explique comment ces relations √©voluent au fil du temps et renforcent la CLV dans un contexte de march√© volatile. ¬ª

        5. **Sources de Revenus**  
        M√©ta-Prompt :
        ¬´ D√©taille les m√©canismes de g√©n√©ration de revenus :  
        - Mod√®les de tarification : vente directe √† prix abordable, options de micro-paiements √©chelonn√©s, cr√©dit via partenaire de microfinance, location-vente, abonnement l√©ger (maintenance), freemium (d√©monstration gratuite, paiement pour les pi√®ces d√©tach√©es).  
        - Justification des prix : aligner le prix sur le pouvoir d‚Äôachat, offrir un excellent rapport qualit√©/durabilit√©/prix, tenir compte des r√©f√©rentiels locaux (si les concurrents informels sont tr√®s bon march√©, justifier la valeur par la fiabilit√©).  
        - R√©ductions des freins √©conomiques : essai avant achat, garantie satisfait ou rembours√©, partenariats avec ONG ou institutions locales.  
        - Diversification des revenus : ventes crois√©es (pi√®ces d√©tach√©es, formation), partenariats B2B (ventes en gros √† des coop√©ratives), publicit√© locale, sponsorisation par des institutions de d√©veloppement.  
        - Adaptation aux changements : si le march√© se contracte, proposer des mod√®les encore plus frugaux, si la r√©glementation change, s‚Äôadapter par des produits conformes.  
        Explique comment cette structure de revenus soutient la viabilit√© √† long terme et reste coh√©rente avec la proposition de valeur et la sensibilit√© au prix de la persona. ¬ª

        6. **Ressources Cl√©s**  
        M√©ta-Prompt :  
        ¬´ Identifie toutes les ressources indispensables :  
        - Ressources Humaines : agents locaux (form√©s aux langues et contextes locaux), r√©parateurs, formateurs, personnels de SAV.  
        - Ressources Technologiques : outils de communication simples (t√©l√©phones basiques, logiciels l√©gers), syst√®mes de paiement mobile, √©ventuellement une plateforme centralis√©e mais l√©g√®re.  
        - Ressources Intellectuelles : savoir-faire sur l‚Äôadaptation du produit au contexte local, guides visuels, partenariats de R&D avec des instituts techniques locaux.  
        - Ressources Mat√©rielles : pi√®ces d√©tach√©es robustes, mat√©riaux durables, √©quipements simples qui ne n√©cessitent pas une infrastructure complexe.  
        - Ressources Financi√®res : capital initial, fonds de roulement, acc√®s √† la microfinance ou √† des investisseurs sociaux, tr√©sorerie pour faire face aux saisons difficiles.  
        - Ressources Relationnelles : liens solides avec les communaut√©s, chefs traditionnels, radios locales, ONG, institutions de d√©veloppement.  
        Pour chaque ressource, justifie pourquoi elle est critique (ex. sans agents locaux, pas de confiance ; sans mat√©riaux robustes, produit inutilisable), et comment ces ressources assurent un avantage concurrentiel durable. ¬ª

        7. **Activit√©s Cl√©s**  
        M√©ta-Prompt :  
        ¬´ D√©cris les activit√©s indispensables :  
        - D√©veloppement & Innovation : adapter le produit aux conditions locales (climat, langue), am√©liorer la durabilit√©, simplifier l‚Äôusage.  
        - Production & Livraison : fabrication locale ou semi-locale, contr√¥le de la qualit√©, approvisionnement en pi√®ces robustes, logistique simple (transport par camions, motos, √¢nes si n√©cessaire).  
        - Marketing & Ventes : communication via radios communautaires, d√©monstrations physiques, formation d‚Äôagents, distribution de brochures visuelles.  
        - Relation Client & Support : formation du personnel de SAV, mise en place d‚Äôune hotline t√©l√©phonique, ateliers pratiques, visites r√©guli√®res sur le terrain.  
        - Partenariats & N√©gociations : conclure des partenariats avec ONG, coop√©ratives, associations villageoises, n√©gocier des conditions avantageuses avec fournisseurs locaux.  
        Int√®gre une perspective adaptative : si la demande fluctue, ajuster les stocks, si une nouvelle r√©glementation appara√Æt, adapter le produit. Justifie comment chaque activit√© soutient la proposition de valeur. ¬ª

        8. **Partenaires Cl√©s**  
        M√©ta-Prompt :  
        ¬´ Liste et justifie les partenaires strat√©giques :  
        - Fournisseurs locaux : garantissant disponibilit√© et qualit√© des mati√®res premi√®res.  
        - Distributeurs locaux et interm√©diaires informels : acc√®s direct √† la client√®le, r√©duction des co√ªts d‚Äôacquisition.  
        - Partenaires technologiques locaux ou ONG : formation, maintenance, R&D adapt√©e.  
        - Organismes de certification locaux, influenceurs communautaires, m√©dias (radios, journaux locaux) : augmentent la cr√©dibilit√© et la confiance.  
        - Institutions financi√®res (microfinance) : faciliter l‚Äôacc√®s au cr√©dit, au paiement √©chelonn√©.  
        Anticipe les risques (un partenaire cl√© fait d√©faut, troubles politiques, p√©nuries) et pr√©vois des alternatives (autres fournisseurs, diversification g√©ographique). Explique comment ces partenariats renforcent la proposition de valeur et l‚Äôefficacit√© op√©rationnelle. ¬ª

        9. **Structure de Co√ªts**  
        M√©ta-Prompt :  
        ¬´ D√©taille les co√ªts :  
        - Co√ªts fixes : salaires des agents locaux, loyers de petits entrep√¥ts, licences minimales, amortissement de mat√©riel de base.  
        - Co√ªts variables : achat des mati√®res premi√®res, commission aux revendeurs, campagnes radio, formation continue, SAV.  
        - Co√ªts li√©s √† l‚Äôinnovation : R&D pour adapter le produit, formation des √©quipes, tests terrain.  
        Analyse la rentabilit√© :  
        - Le mod√®le de revenus couvre-t-il ces co√ªts ?  
        - Possibilit√©s de r√©duire les co√ªts (sourcing local moins cher, √©conomies d‚Äô√©chelle, recyclage, revente de pi√®ces usag√©es).  
        - Strat√©gies pour faire face aux fluctuations (augmenter la part de services, moduler les prix, limiter le stock).  
        Explique comment la structure de co√ªts reste en ligne avec la proposition de valeur, le niveau de vie local, et comment elle assure la p√©rennit√© financi√®re √† long terme. ¬ª

        ---

        **Instructions Finales** :  
        Apr√®s avoir utilis√© ces m√©ta-prompts pour chaque bloc du Business Model Canvas, effectue une derni√®re v√©rification :  
        - Assure-toi que tous les blocs sont coh√©rents et align√©s avec la proposition de valeur, le parcours client et les r√©alit√©s locales.  
        - V√©rifie que l‚Äôinnovation (low-tech ou high-tech adapt√©e) est r√©ellement adoptable par la persona, apporte un avantage concurrentiel durable, et que les contraintes (culturelles, √©conomiques, r√©glementaires, infrastructurelles) sont prises en compte.  
        - Contr√¥le la rentabilit√©, la viabilit√© √† long terme, et la flexibilit√© face aux changements (variations saisonni√®res, crises √©conomiques, √©volution des r√©glementations ou de la p√©n√©tration technologique).  
        - Ajuste les √©l√©ments (segments, prix, canaux, partenariats) si n√©cessaire pour am√©liorer la robustesse du mod√®le.  
        - Fournis un r√©capitulatif global du Business Model, mettant en avant la logique, la coh√©rence, la proposition de valeur diff√©renciante et quelques chiffres (taille du march√© estim√©e, co√ªts, revenus, marge, etc.) pour valider la viabilit√© √©conomique.

        Le r√©sultat final doit √™tre un Business Model clair, complet, adapt√© au contexte local, pr√™t √† √™tre test√© ou impl√©ment√©, avec une feuille de route pour l‚Äôadoption progressive de l‚Äôinnovation et une vision claire des points de diff√©renciation face aux solutions traditionnelles ou informelles existantes.
        """,
        
        
        "Startup": """ Tu es un assistant expert en strat√©gie d‚Äôentreprise, marketing, UX, innovation et √©laboration de Business Models. Ton r√¥le est de g√©n√©rer un Business Model complet, clair, chiffr√©, coh√©rent et innovant, en suivant trois phases : Configuration Initiale, √âtapes Interm√©diaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation) et Production Finale (Business Model Canvas).

        Tout au long du processus, tu dois :
        - Prendre en compte la persona (donn√©es d√©mographiques, comportementales, capacit√©s d‚Äôadoption de l‚Äôinnovation).
        - Pendre en compte l'arbre √† problemes(Probl√®me Principal,Causes Principales,Impact,Parties Prenantes , Opportunit√©s)  
        - Analyser le march√© (taille, segments, offres existantes, niveau de satisfaction, tendances).
        - Int√©grer les facteurs limitants (technologiques, √©conomiques, culturels, psychologiques, physiologiques, r√©glementaires).
        - √âvaluer la concurrence et comprendre le niveau de satisfaction actuel.
        - Comprendre le parcours client (avant, pendant, apr√®s), la carte d‚Äôempathie, les gains et souffrances.
        - V√©rifier syst√©matiquement la coh√©rence, proposer des optimisations et ajustements.
        - Avant d‚Äôintroduire une innovation, t‚Äôassurer que la persona est pr√™te √† l‚Äôadopter.
        - Produire un Business Model Canvas complet (9 blocs), avec des meta-prompts sp√©cifiques pour chacun des blocs.

        Voici les √©tapes :

        ### Phase 1 : Configuration Initiale (Entr√©e de Donn√©es)

        1. Recueille et structure les informations suivantes :
        - **Persona :**
            - Donn√©es d√©mographiques : √Çge, sexe, localisation, niveau d‚Äô√©ducation, profession, revenu.
            - Param√®tres comportementaux : Sensibilit√© au prix, budget, fr√©quence et volume d‚Äôachat, perception de la qualit√©, utilisation technologique, accessibilit√©, temps disponible, besoins, motivations.
            - Capacit√© d‚Äôadoption de l‚Äôinnovation : Familiarit√© technologique, ouverture au changement, barri√®res psychologiques ou culturelles.
        
        - **Arbre √† Probl√®me** :
            - Contexte:Description g√©n√©rale du domaine ou de la situation actuelle (secteur d'activit√©, environnement g√©ographique, tendances actuelles du march√©), Facteurs externes influen√ßant la situation (r√©glementations, conditions √©conomiques, technologies √©mergentes),Facteurs internes pertinents (ressources disponibles, comp√©tences cl√©s, structure organisationnelle).
            - Probl√®me Principal : Identification du d√©fi ou de l'obstacle central (nature du probl√®me, circonstances sp√©cifiques),Impact imm√©diat sur l'organisation ou le projet (effets sur les op√©rations, la performance financi√®re, la r√©putation).
            - Causes Principales :Causes internes contribuant au probl√®me (processus inefficaces, manque de comp√©tences, ressources limit√©es),Causes externes contribuant au probl√®me (concurrence accrue, changements de march√©, √©volutions technologiques),Interaction entre les causes internes et externes (comment elles se renforcent mutuellement).
            - Impact:Cons√©quences financi√®res du probl√®me (pertes de revenus, augmentation des co√ªts, rentabilit√© r√©duite),Effets op√©rationnels (d√©lai dans les projets, baisse de productivit√©, qualit√© des services ou produits affect√©e),Impact sur les parties prenantes (satisfaction des clients, moral des employ√©s, relations avec les partenaires).
            - Parties Prenantes :Identification des acteurs concern√©s ou impact√©s (clients, employ√©s, fournisseurs, investisseurs, communaut√© locale),Int√©r√™ts et attentes de chaque partie prenante vis-√†-vis du probl√®me (besoins sp√©cifiques, priorit√©s, pr√©occupations).
            - Opportunit√©s :Pistes d‚Äôam√©lioration ou de r√©solution du probl√®me (solutions innovantes, meilleures pratiques),Strat√©gies pour att√©nuer les causes principales (formation, r√©organisation, investissement technologique),Actions pour maximiser les impacts positifs (exploitation des forces, diversification, partenariats strat√©giques),
        
        - **Analyse du March√© :**
            - Taille du march√©, segments, valeur totale.
            - Offres concurrentes, niveau de satisfaction, tendances, innovations, comportements √©mergents.
        - **Facteurs Limitants :**
            - Contraintes technologiques, √©conomiques, culturelles, r√©glementaires, physiologiques, psychologiques.
        
        Apr√®s avoir recueilli ces donn√©es, effectue une premi√®re analyse critique :
        - V√©rifie la coh√©rence des informations.
        - Identifie les lacunes.
        - Propose des compl√©ments ou ajustements pour optimiser la qualit√© des donn√©es.

        ### Phase 2 : √âtapes Interm√©diaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)

        2. **Analyse du Parcours Client & Carte d‚ÄôEmpathie :**
        - D√©cris le parcours client (avant, pendant, apr√®s consommation).
        - Identifie les points de contact, obstacles, moments de v√©rit√©, frustrations.
        - Int√®gre les contraintes physiologiques, psychologiques, √©conomiques, culturelles, technologiques, r√©glementaires.
        - Cr√©e une carte d‚Äôempathie (pens√©es, sentiments, actions) pour comprendre l‚Äôexp√©rience du client √† chaque √©tape.

        3. **Gains et Souffrances :**
        - √Ä partir du parcours client et de la carte d‚Äôempathie, liste les gains (b√©n√©fices, r√©assurance, sentiment de comp√©tence) et les souffrances (probl√®mes non r√©solus, frustrations, co√ªts d‚Äôopportunit√©).

        4. **√âlaboration de la Carte de Valeur :**
        - D√©finis la mission de consommation principale (besoin fondamental).
        - Identifie les gains d√©j√† fournis par les offres actuelles.
        - Mets en √©vidence les souffrances non adress√©es.
        - Esquisse une proposition de valeur pr√©liminaire, adapt√©e √† la capacit√© d‚Äôadoption de l‚Äôinnovation par la persona.

        5. **D√©termination du Segment de Clients :**
        - Choisis le type de relation (B2C, B2B, B2B2C‚Ä¶).
        - Priorise les segments (taille, pouvoir d‚Äôachat, sensibilit√© au prix, ouverture √† l‚Äôinnovation, contraintes) qui correspondent le mieux √† la proposition de valeur.

        6. **Analyse des Probl√®mes et Solutions (Canvas de Probl√®me) :**
        - Identifie clairement les probl√®mes majeurs √† r√©soudre.
        - Associe chaque probl√®me √† une solution sp√©cifique, justifie en quoi elle est meilleure que les offres existantes.

        Apr√®s ces √©tapes, effectue une analyse interm√©diaire :
        - V√©rifie la coh√©rence du contexte, du parcours client, des solutions propos√©es.
        - Assure-toi que les innovations sont compr√©hensibles, utiles et adoptables par la persona.
        - Propose des ajustements strat√©giques (simplification de l‚Äôoffre, ajustement du prix, s√©lection de segments plus pertinents, etc.) si n√©cessaire.

        ### Phase 3 : Production Finale du Business Model (Business Model Canvas)

        Sur la base des analyses pr√©c√©dentes, g√©n√®re un Business Model Canvas complet. Utilise les meta-prompts suivants pour chaque bloc :

        1. **Segments de Clients**  
        M√©ta-Prompt :  
        ¬´ D√©finis pr√©cis√©ment les segments de clients cibl√©s, en tenant compte :  
        - De leurs caract√©ristiques sociod√©mographiques (√¢ge, sexe, localisation, niveau d‚Äô√©ducation, profession, revenu).  
        - De leurs comportements d‚Äôachat (fr√©quence, volume, sensibilit√© au prix, crit√®res de qualit√©) et de leur maturit√© technologique (utilisation d‚Äôoutils num√©riques, appareils connect√©s, plateformes en ligne).  
        - De leur capacit√© d‚Äôadoption de l‚Äôinnovation (ouverture au changement, barri√®res psychologiques, √©ventuelle r√©ticence culturelle).  
        - De leurs contraintes physiologiques (accessibilit√©, ergonomie), psychologiques (stress, anxi√©t√©, besoin de r√©assurance), √©conomiques (pouvoir d‚Äôachat, rapport qualit√©/prix), culturelles (normes, tabous) et r√©glementaires (normes l√©gales, certifications).  
        Int√®gre √©galement des sc√©narios √©volutifs :  
        - Si la technologie √©volue, comment ce segment r√©agit-il ?  
        - S‚Äôil y a une crise √©conomique, ces clients r√©duisent-ils leur consommation ?  
        - Une partie du segment est-elle pr√™te √† payer plus pour des options premium ?  
        Justifie pourquoi ces segments sont retenus, comment ils se distinguent de segments non cibl√©s, et comment leur potentiel de rentabilit√©, leur facilit√© d‚Äôacc√®s, et leur r√©ceptivit√© √† la proposition de valeur justifient leur inclusion. ¬ª

        2. **Proposition de Valeur**  
        M√©ta-Prompt :  
        ¬´ D√©taille la proposition de valeur en explicitant :  
        - Les besoins fondamentaux adress√©s (mission de consommation principale).  
        - Les souffrances clientes (manque de temps, complexit√©, mauvaise qualit√©, manque de confiance, crainte face √† l‚Äôinnovation) et comment elles sont r√©solues.  
        - Les gains fournis (gain de temps, √©conomie d‚Äôargent, facilit√© d‚Äôutilisation, statut social, tranquillit√© d‚Äôesprit), y compris les b√©n√©fices √©motionnels et symboliques.  
        - La diff√©renciation par rapport aux offres concurrentes (qualit√© sup√©rieure, innovation plus accessible, prix comp√©titifs, service client exemplaire, partenariats de prestige).  
        - L‚Äôint√©gration de l‚Äôinnovation : montre comment elle est introduite progressivement, comment l‚Äô√©ducation ou la formation du client est assur√©e, et comment les barri√®res √† l‚Äôadoption sont lev√©es (essais gratuits, d√©monstrations, tutoriels, certifications reconnues).  
        - Pr√©vois des variantes de proposition de valeur en fonction des segments, si n√©cessaire (une version premium pour les early adopters innovants, une version simplifi√©e pour les plus conservateurs). ¬ª

        3. **Canaux de Distribution**  
        M√©ta-Prompt :  
        ¬´ D√©finis les canaux par lesquels les clients seront inform√©s, convaincus, ach√®teront et utiliseront le produit/service. Consid√®re :  
        - Les canaux en ligne (site web, application mobile, plateformes e-learning, r√©seaux sociaux, partenariats avec marketplaces, influenceurs, SEO, SEA).  
        - Les canaux hors ligne (magasins physiques, salons professionnels, conf√©rences, revendeurs, agents sur le terrain).  
        - La n√©cessit√© de coh√©rence entre les points de contact (omnicanal), la simplicit√© d‚Äôacc√®s, le besoin d‚Äôaccompagnement p√©dagogique (webinaires, tutoriels vid√©o), et les contraintes technologiques de la persona (faible bande passante, pr√©f√©rence pour un canal mobile vs desktop).  
        - L‚Äôadaptabilit√© des canaux si les conditions du march√© changent (p√©nurie d‚Äôun canal, √©volution l√©gale, concurrence d‚Äôun nouveau canal).  
        Justifie pourquoi chaque canal est choisi, comment il s‚Äôint√®gre dans le parcours client, comment il favorise l‚Äôadoption de l‚Äôinnovation, et comment il est optimis√© pour r√©duire les co√ªts d‚Äôacquisition et am√©liorer la satisfaction. ¬ª

        4. **Relations Clients**  
        M√©ta-Prompt :  
        ¬´ D√©cris la nature et la qualit√© des relations que l‚Äôentreprise √©tablira avec ses clients :  
        - Personnalisation : existe-t-il un accompagnement individuel, des conseils sur mesure, une assistance humaine ou une IA conversationnelle ?  
        - Communaut√© : les clients peuvent-ils interagir entre eux (forums, r√©seaux sociaux, clubs, rencontres physiques) pour renforcer leur sentiment d‚Äôappartenance et √©changer des exp√©riences ?  
        - Automatisation : y a-t-il des √©l√©ments de self-service, de chatbots, de bases de connaissances en ligne ? Est-ce adapt√© aux cibles moins technophiles ?  
        - Fid√©lisation : cartes de fid√©lit√©, programmes de r√©compenses, contenus exclusifs, mises √† jour gratuites, offres sp√©ciales pour clients fid√®les.  
        - Gestion des plaintes et retours : proc√©dures de remboursement, garantie de satisfaction, SLA pour r√©pondre aux demandes critiques.  
        Int√®gre la dimension psychologique (rassurer les clients sur l‚Äôinnovation), culturelle (certains clients pr√©f√®rent un contact humain), r√©glementaire (besoin de conformit√© avec les lois sur la protection des donn√©es).  
        Explique comment ces relations √©voluent au fil du temps (du premier contact √† la fid√©lisation), comment elles am√©liorent la CLV, et comment elles s‚Äôadaptent aux changements de march√© (nouveaux concurrents, crises √©conomiques). ¬ª

        5. **Sources de Revenus**  
        M√©ta-Prompt :  
        ¬´ D√©taille les m√©canismes de g√©n√©ration de revenus :  
        - Mod√®le de tarification : abonnement mensuel, paiement √† l‚Äôusage, achat unique, freemium avec options premium, licences, commissions.  
        - Justification des prix : comment le prix refl√®te-t-il la valeur per√ßue par le client ? Est-il align√© avec le pouvoir d‚Äôachat du segment, la concurrence, la qualit√© et l‚Äôinnovation propos√©e ?  
        - Options de r√©duction des freins √©conomiques : essais gratuits, garantie satisfait ou rembours√©, paiement √©chelonn√©, remises pour les early adopters.  
        - Diversification des revenus : ventes crois√©es, upselling, partenariats, publicit√©, formation compl√©mentaire, mon√©tisation de donn√©es (en respectant la r√©glementation).  
        - Adaptation √† des changements de contexte : si le march√© se contracte, proposer un mod√®le plus flexible ? Si une r√©glementation limite certains types de revenus, anticiper une alternative ?  
        Explique comment cette structure de revenus soutient la croissance, la rentabilit√©, et s‚Äôint√®gre avec les co√ªts pr√©vus. V√©rifie la coh√©rence avec la proposition de valeur et la sensibilit√© au prix de la persona. ¬ª

        6. **Ressources Cl√©s**  
        M√©ta-Prompt :  
        ¬´ Identifie toutes les ressources indispensables :  
        - Ressources Humaines : √©quipes multidisciplinaires (ing√©nieurs, designers UX, experts marketing, formateurs, support client multilingue) n√©cessaires √† la cr√©ation, maintenance, am√©lioration de l‚Äôoffre.  
        - Ressources Technologiques : plateformes e-learning, serveurs, logiciels de personnalisation, outils d‚ÄôIA, applications mobiles, infrastructure IT s√©curis√©e.  
        - Ressources Intellectuelles : brevets, marques, contenus propri√©taires, m√©thodologies exclusives, licences de tiers, donn√©es clients prot√©g√©es.  
        - Ressources Financi√®res : capitaux n√©cessaires au lancement, tr√©sorerie pour r√©sister √† une p√©riode de faible demande, fonds pour R&D.  
        - Ressources Relationnelles : partenariats strat√©giques, acc√®s √† un r√©seau d‚Äôinfluenceurs, certification par des organismes reconnus.  
        Explique pour chaque ressource pourquoi elle est critique, comment elle se combine avec les autres pour d√©livrer la proposition de valeur, soutenir l‚Äôadoption de l‚Äôinnovation, et maintenir un avantage concurrentiel. Prends en compte la robustesse de la cha√Æne d‚Äôapprovisionnement, la r√©silience face aux crises, et la propri√©t√© intellectuelle. ¬ª

        7. **Activit√©s Cl√©s**  
        M√©ta-Prompt :  
        ¬´ D√©cris les activit√©s indispensables pour que le Business Model fonctionne :  
        - D√©veloppement & Innovation : R&D, am√©lioration continue, int√©gration de nouvelles fonctionnalit√©s, veille concurrentielle, tests utilisateurs.  
        - Production & Livraison : cr√©ation de contenu, mise √† jour r√©guli√®re, gestion du stock (si produit physique), maintenance technique, logistique.  
        - Marketing & Ventes : campagnes publicitaires, r√©f√©rencement, webinaires de d√©monstration, √©ducation du march√©, gestion des promotions.  
        - Relation Client & Support : formation du personnel du support, chatbots, assistance multicanal, traitement des plaintes, suivi de la satisfaction.  
        - Partenariats & N√©gociations : recherche, signature et entretien des partenariats cl√©s, mise en place de conditions avantageuses.  
        Int√®gre une perspective adaptative :  
        - Quelles activit√©s mener si la demande fluctue fortement ?  
        - Comment r√©allouer les ressources si une nouvelle r√©glementation √©merge ?  
        Justifie comment chaque activit√© soutient la proposition de valeur, favorise l‚Äôadoption de l‚Äôinnovation, et contribue √† la rentabilit√© globale. ¬ª

        8. **Partenaires Cl√©s**  
        M√©ta-Prompt :  
        ¬´ Liste et justifie les partenaires strat√©giques critiques :  
        - Fournisseurs : apportant des ressources rares, de haute qualit√© ou √† un co√ªt avantageux.  
        - Distributeurs : offrant un acc√®s facilit√© √† certains segments, r√©duisant les co√ªts d‚Äôacquisition, am√©liorant la visibilit√©.  
        - Partenaires technologiques : fournissant une infrastructure fiable, des outils d‚ÄôIA performants, ou des solutions compl√©mentaires (API, int√©grations).  
        - Organismes de certification, influenceurs, m√©dias sp√©cialis√©s : augmentant la cr√©dibilit√©, validant la qualit√©, rassurant sur l‚Äôinnovation.  
        - Associations professionnelles, clusters, √©cosyst√®mes sectoriels : permettant de suivre les tendances, d‚Äôanticiper les changements r√©glementaires, d‚Äô√©changer les bonnes pratiques.  
        Explique comment ces partenariats renforcent la proposition de valeur, am√©liorent la confiance du client, augmentent l‚Äôefficacit√© op√©rationnelle, r√©duisent les co√ªts ou les risques, et soutiennent la strat√©gie √† long terme. Anticipe les risques : et si un partenaire cl√© fait d√©faut ? Quels sont les plans B ? ¬ª

        9. **Structure de Co√ªts**  
        M√©ta-Prompt :  
        ¬´ D√©taille tous les co√ªts engendr√©s par les ressources, activit√©s et partenariats cl√©s :  
        - Co√ªts fixes (salaires, loyers, licences, amortissement de l‚Äôinfrastructure).  
        - Co√ªts variables (marketing, support client, acquisition de nouveaux outils, commission aux partenaires).  
        - Co√ªts li√©s √† l‚Äôinnovation (R&D, tests, formations du personnel), et comment ils sont amortis dans le temps.  
        Analyse la rentabilit√© :  
        - Le mod√®le de revenus couvre-t-il ces co√ªts ?  
        - Quelles mesures de r√©duction de co√ªts sont possibles (automatisation, sourcing moins cher, √©conomies d‚Äô√©chelle) ?  
        - Comment r√©agir face √† des fluctuations du march√© (baisse de la demande, hausse des prix des ressources) ?  
        Explique comment la structure de co√ªts reste align√©e avec la proposition de valeur, les segments, et les moyens financiers de l‚Äôentreprise. Justifie la p√©rennit√© financi√®re en montrant que les marges sont satisfaisantes, que le CAC est raisonnable par rapport √† la CLV, et que le mod√®le reste rentable m√™me en cas de stress. ¬ª

        ### Instructions Finales

        Apr√®s avoir utilis√© ces m√©ta-prompts pour chaque bloc du Business Model Canvas, effectue une derni√®re v√©rification :

        - Assure-toi que tous les blocs sont coh√©rents entre eux et s‚Äôalignent parfaitement avec la proposition de valeur et le parcours client.
        - V√©rifie que l‚Äôinnovation propos√©e est bien adoptable par la persona, qu‚Äôelle apporte un avantage concurrentiel durable, et que les contraintes sont g√©r√©es.  
        - Contr√¥le la rentabilit√©, la viabilit√© √† long terme, et la flexibilit√© pour s‚Äôadapter aux changements de march√©.
        - Ajuste les √©l√©ments (segments, prix, canaux, partenariats) si n√©cessaire pour am√©liorer la robustesse du mod√®le.

        Le r√©sultat final doit √™tre un Business Model clair, complet, et pr√™t √† √™tre test√© ou impl√©ment√©, avec une feuille de route pour l‚Äôadoption de l‚Äôinnovation et une vision claire des points de diff√©renciation face √† la concurrence.


        Enfin, fournis un r√©capitulatif global du Business Model, mettant en avant la logique, la coh√©rence, et la proposition de valeur diff√©renciante. Indique, si possible, des chiffres (taille du march√©, CAC, CLV, taux de conversion, CA projet√©) pour valider la viabilit√© √©conomique.""",
       
        "Autre": "Fournissez une approche g√©n√©rale adapt√©e √† votre entreprise."
    }
    return metaprompts.get(type_entreprise, metaprompts["Autre"])


def obtenir_business_model(nom_entreprise, type_entreprise,previousdata, rubriques,generation=1):
    
    """
    Interroge ChatGPT (API OpenAI) pour g√©n√©rer le contenu textuel
    des diff√©rents blocs du Business Model Canvas.
    'type_entreprise' peut √™tre "PME", "Startup", "Grande Entreprise", etc.
    'previousdata' peut etre du contenue html generer precedement par chatgpt
    """

    
    
    # R√©cup√©rer le metaprompt bas√© sur le type d'entreprise
    metaprompt = get_metaprompt(type_entreprise)
    print(rubriques)
    
    if generation == 1:
        # Premi√®re g√©n√©ration avec les nouvelles rubriques
        prompt = f"""
        {metaprompt}
        
        Mener la reflexions du generation du business modele sur base des indications(M√©ta-Prompt) precedents du metaprompts; 
        Chercher les chiffres et autres donn√©es sur internet, assurer-vous d'etre trop precis et excat en fonction fonction des donn√©es collecter sur internet 
        G√©n√®re le contenu d'un Business Model Canvas en format HTML et CSS encapsul√© dans des blocs de code sans aucun autre texte pour une entreprise nomm√©e '{nom_entreprise}'.
        Le type d'entreprise est : {type_entreprise}.
        
        Utilisez les donn√©es ci apr√®s(dans la rubriques) comme donn√©es collect√© lors de la Phase 1 (Configuration Initiale (Entr√©e de Donn√©es)): {rubriques}
        Certains partie du rubriques peuvent etre vide, si c'est les cas generer les donn√©es manquantes. les chiffres entrer pour l'utilisateur doivent etre imperativement tenue en compte
        
        √Ä faire imp√©rativement :
        Je veux imp√©rativement 9 blocs distincts, r√©dig√©s en fran√ßais, avec les titres en gras et des listes √† puces si n√©cessaire :
          - Partenaires cl√©s
          - Activit√©s cl√©s
          - Offre (proposition de valeur)
          - Relation client
          - Segments de client√®le
          - Ressources cl√©s
          - Canaux de distribution
          - Structure des co√ªts
          - Sources de revenus
        Fournissez 5 √† 10 points ou √©l√©ments (phrases) par bloc pour un contenu riche et adapt√©, soyez concis.
        """
    else:
        # Deuxi√®me g√©n√©ration (am√©lioration) en utilisant le BMC pr√©c√©dent et les nouvelles rubriques
        # Prompt ajust√© sans num√©rotation dans les titres
        prompt = f"""
        {metaprompt}
        
        
        Voici les donn√©es generer precedement {previousdata}
        Ameliorer ces business modeles modeles sur bases de metaprompt, et des informations fournit pour chaque rubriques
        Mener la reflexions du generation du business modele sur base des indications(M√©ta-Prompt) precedents du metaprompts; 
        Chercher les chiffres et autres donn√©es sur internet, assurer-vous d'etre trop precis et excat en fonction fonction des donn√©es collecter sur internet 
        G√©n√®re le contenu d'un Business Model Canvas en format HTML pour une entreprise nomm√©e '{nom_entreprise}'.
        Le type d'entreprise est : {type_entreprise}.
        
        
        et dont les donn√©es complementaires (non obligatoire pour l'utilisateur) pour chaque bloc se trouve dans : {rubriques}.
        si l'utlisateur a donner les donn√©es complementaires, veuillez en tenir compte dans la generation, et ca doit etre imperativement prioritaire.
        Si dans un bloque un utilisateur n'as pas donner des informations (elements), veuillez generer,
        Si l'utilisateur √† donn√© des elements que vous juger peu, generer d'autres et les ajout√©es √† ce que l'utlisateur √† fournit.
        
        √† faire imperativement est:
        Je veux imp√©rativement 9 blocs distincts, r√©dig√©s en fran√ßais, avec les titres en gras et des listes √† puces si n√©cessaire :
        - Partenaires cl√©s
        - Activit√©s cl√©s
        - Offre (proposition de valeur)
        - Relation client
        - Segments de client√®le
        - Ressources cl√©s
        - Canaux de distribution
        - Structure des co√ªts
        - Sources de revenus
        Fournis 5 √† 10 points ou √©lements(phrases) , meme plus pour chacun afin d'avoir un contenu riche et adapt√©, soyez concis.
        """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Tu es un assistant expert en g√©n√©ration de business  et business plan."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=5000,
            temperature=0.7
        )
        html_genere = response.choices[0].message.content.strip()
        return html_genere
    except Exception as e:
        st.error(f"Erreur lors de la g√©n√©ration du contenu : {e}")
        return ""

# ----------------------------------------------------------------------------
# 2) Fonction pour cr√©er le fichier Word (format tableau) avec python-docx
# ----------------------------------------------------------------------------

def generer_docx_business_model(nom_entreprise, date_bmc, contenu_business_model, doc, value=1):
    """
    Construit un document Word reproduisant un tableau avec la disposition souhait√©e
    pour le Business Model Canvas. La mise en forme inclut des titres en gras et
    des listes √† puces.
    'contenu_business_model' : le contenu HTML renvoy√© par ChatGPT,
    qu'on d√©coupe ensuite pour remplir chaque bloc.
    """
    # Cr√©er un nouveau document Word
    if value == 1:
        doc = Document()

    # D√©finir les styles de base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titre principal
    titre = doc.add_heading(level=1)
    titre_run = titre.add_run(f"Business Model Canvas de {nom_entreprise}")
    titre_run.bold = True
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date : {date_bmc}")
    date_run.bold = True
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Ajouter un saut de ligne
    doc.add_paragraph("")

    # Cr√©er un tableau de 6 lignes √ó 5 colonnes
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    # Ajuster les largeurs des colonnes (en pouces)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.8)  # Ajustez selon vos besoins

    # 1) Ligne 0 : Titre (fusion des 5 colonnes)
    cell00 = table.cell(0, 0)
    cell00_merge = cell00.merge(table.cell(0, 4))
    cell00_merge.text = f"Business Model Canvas de {nom_entreprise}"
    for paragraph in cell00_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 2) Ligne 1 : Nom de l'entreprise et Date (fusion des colonnes)
    cell10 = table.cell(1, 0)
    cell10_merge = cell10.merge(table.cell(1, 2))
    cell10_merge.text = f"**Nom de l'entreprise**: {nom_entreprise}"
    for paragraph in cell10_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    cell13 = table.cell(1, 3)
    cell13_merge = cell13.merge(table.cell(1, 4))
    cell13_merge.text = f"**Date**: {date_bmc}"
    for paragraph in cell13_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 3) Ligne 2 : Headers des 5 blocs
    headers = ["Partenaires cl√©s", "Activit√©s cl√©s", "Offre (proposition de valeur)", 
               "Relation client", "Segments de client√®le"]
    for idx, header in enumerate(headers):
        cell = table.cell(2, idx)
        paragraphe = cell.paragraphs[0]
        run = paragraphe.add_run(header)
        run.bold = True
        paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 4) Ligne 3 : Contenus des 5 blocs
    # Initialiser les blocs
    blocs = {
        "Partenaires cl√©s": "",
        "Activit√©s cl√©s": "",
        "Offre (proposition de valeur)": "",
        "Relation client": "",
        "Segments de client√®le": ""
    }

    # Utiliser BeautifulSoup pour parser le HTML
    soup = BeautifulSoup(contenu_business_model, 'html.parser')

    # Fonction pour trouver le bon header tag (h3 par d√©faut, avec flexibilit√©)
    def trouver_header(soup, header):
        # Regex pour capturer optionnellement des num√©ros suivis de points et espaces
        pattern = rf"^(?:\d+\.\s*)?{re.escape(header)}$"
        # Chercher dans les balises h3
        header_tag = soup.find(['h2', 'h3', 'h4', 'h5', 'h6'], text=re.compile(pattern, re.IGNORECASE))
        return header_tag

    # Extraire chaque bloc
    for header in blocs.keys():
        h_tag = trouver_header(soup, header)
        if h_tag:
            content = []
            for sibling in h_tag.find_next_siblings():
                if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                    break  # Arr√™ter si un nouveau header est trouv√©
                if sibling.name == 'ul':
                    for li in sibling.find_all('li'):
                        content.append(f"- {li.get_text(strip=True)}")
                elif sibling.name == 'p':
                    content.append(sibling.get_text(strip=True))
                elif isinstance(sibling, str):
                    text = sibling.strip()
                    if text:
                        content.append(text)
            blocs[header] = '\n'.join(content)

    # Debug: Afficher les blocs extraits (√† d√©sactiver en production)
    # st.write("Blocs extraits :", blocs)

    # Fonction pour ajouter du contenu format√© dans une cellule
    def ajouter_contenu(cell, titre, contenu):
        """
        Ajoute du contenu format√© dans une cellule Word.
        Le titre est en gras, suivi de listes √† puces si n√©cessaire.
        """
        # Supprimer le texte initial (par d√©faut) dans la cellule
        cell.text = ""

        # Ajouter le titre en gras
        paragraphe = cell.add_paragraph()
        run = paragraphe.add_run(titre)
        run.bold = True

        # Ajouter le contenu
        # Diviser le contenu par les sauts de ligne
        lignes = contenu.split('\n')
        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne:
                continue
            # V√©rifier si la ligne commence par '-', '+', '‚Ä¢' pour une liste √† puces
            if re.match(r'^[-+‚Ä¢]\s+', ligne):
                # Ajouter une puce
                item = re.sub(r'^[-+‚Ä¢]\s+', '', ligne)
                p = cell.add_paragraph(item, style='List Bullet')
            else:
                # Ajouter un paragraphe normal
                p = cell.add_paragraph(ligne)

    # Remplir les cellules de la ligne 3
    ordre_blocs = [
        "Partenaires cl√©s", "Activit√©s cl√©s", "Offre (proposition de valeur)",
        "Relation client", "Segments de client√®le"
    ]

    for idx, bloc in enumerate(ordre_blocs):
        cell = table.cell(3, idx)
        ajouter_contenu(cell, bloc, blocs[bloc])

    # 5) Ligne 4 : Structure de co√ªts (fusion 3 cols) et Sources de revenus (fusion 2 cols)
    # Fusionner les cellules pour "Structure de co√ªts" (colonnes 0-2)
    cell40 = table.cell(4, 0)
    cell40_merge = cell40.merge(table.cell(4, 2))
    cell40_merge.text = f"**Structure de co√ªts**:\n\n"

    # Fusionner les cellules pour "Sources de revenus" (colonnes 3-4)
    cell43 = table.cell(4, 3)
    cell43_merge = cell43.merge(table.cell(4, 4))
    cell43_merge.text = f"**Sources de revenus**:\n\n"

    # Extraire les contenus pour ces blocs
    structure_couts = ""
    sources_revenus = ""

    # Structure des co√ªts
    strong_tag = trouver_header(soup, "Structure des co√ªts")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        structure_couts = '\n'.join(content)

    # Sources de revenus
    strong_tag = trouver_header(soup, "Sources de revenus")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        sources_revenus = '\n'.join(content)

    # Remplir les cellules fusionn√©es
    ajouter_contenu(cell40_merge, "Structure de co√ªts", structure_couts)
    ajouter_contenu(cell43_merge, "Sources de revenus", sources_revenus)

    # Ajuster les paragraphes existants
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

    # Ajouter un saut de ligne √† la fin
    doc.add_paragraph("")

    # Convertir le document en binaire pour t√©l√©chargement via Streamlit
    fichier_io = BytesIO()
    doc.save(fichier_io)
    fichier_io.seek(0)
    return fichier_io


def page_generer_business_model():
    st.header("√âtape 2 : G√©n√©rer le Business Model Canvas")
    
    if st.session_state.get('business_model_precedent'):
        st.write("Le Business Model Canvas initial a √©t√© g√©n√©r√©. Vous pouvez le t√©l√©charger ci-dessous ou proc√©der √† son am√©lioration.")
        
        # Afficher le contenu g√©n√©r√©
        st.subheader("Business Model Canvas Initial G√©n√©r√©")
        html_content = st.session_state.business_model_precedent

        # Encoder le contenu HTML en Base64
        encoded_html = base64.b64encode(html_content.encode('utf-8')).decode('utf-8')

        # Cr√©er l'URL de donn√©es
        data_url = f"data:text/html;base64,{encoded_html}"

        st.markdown(
            f"""
            <iframe src="{data_url}" width="100%" height="1500" frameborder="0" scrolling="yes"></iframe>
            """,
            unsafe_allow_html=True
        )
        
        # Vous pouvez √©galement proposer d'autres actions ici si n√©cessaire
    else:
        st.info("Veuillez d'abord collecter toutes les donn√©es et g√©n√©rer le Business Model Canvas initial dans l'onglet 'Collecte des Donn√©es'.")

def page_ameliore_business_model():
    st.header("√âtape 3 : Am√©liorer le Business Model Canvas")
    
    if st.session_state.get('business_model_precedent'):
        st.write("Utilisez cette section pour am√©liorer le Business Model Canvas g√©n√©r√© pr√©c√©demment.")
        
        with st.form("form_ameliore"):
            st.write("Veuillez ajouter des informations suppl√©mentaires pour am√©liorer le Business Model Canvas.")
            
            # Formulaire pour les 9 rubriques du BMC (d√©plac√© vers la fin)
            expand_all = st.checkbox("√âtendre / R√©duire tout le formulaire", value=False)

            with st.expander("Partenaires cl√©s", expanded=expand_all):
                st.markdown("""
                **Partenaires cl√©s :**  
                Identifiez les organisations ou individus essentiels √† votre activit√©. Par exemple :  
                - Fournisseurs  
                - Banques ou institutions financi√®res  
                - Partenaires strat√©giques  
                - Associations ou gouvernements locaux  
                """)
                partenaire_cles = st.text_area("Listez vos principaux partenaires.", key="partenaires_cles_ameliore")

            with st.expander("Activit√©s cl√©s", expanded=expand_all):
                st.markdown("""
                **Activit√©s cl√©s :**  
                D√©crivez les t√¢ches ou processus les plus importants pour ex√©cuter votre mod√®le √©conomique. Par exemple :  
                - Fabrication de produits  
                - Marketing et vente  
                - Recherche et d√©veloppement  
                - Gestion des relations avec les clients  
                """)
                activites_cles = st.text_area("D√©crivez vos activit√©s principales.", key="activites_cles_ameliore")

            with st.expander("Offre (proposition de valeur)", expanded=expand_all):
                st.markdown("""
                **Proposition de valeur :**  
                Expliquez ce que vous offrez √† vos clients et ce qui vous diff√©rencie de vos concurrents. Par exemple :  
                - R√©solution d'un probl√®me sp√©cifique  
                - Am√©lioration d'un besoin existant  
                - Caract√©ristiques uniques de vos produits ou services  
                """)
                offre_valeur = st.text_area("D√©crivez votre proposition de valeur.", key="offre_valeur_ameliore")

            with st.expander("Relation client", expanded=expand_all):
                st.markdown("""
                **Relation client :**  
                D√©crivez comment vous interagissez avec vos clients. Par exemple :  
                - Assistance personnalis√©e  
                - Automatisation des services (chatbots, self-service)  
                - Programmes de fid√©lisation  
                """)
                relation_client = st.text_area("D√©crivez comment vous g√©rez vos relations clients.", key="relation_client_ameliore")

            with st.expander("Segments de client√®le", expanded=expand_all):
                st.markdown("""
                **Segments de client√®le :**  
                Identifiez vos diff√©rents groupes de clients cibles. Par exemple :  
                - Particuliers (par revenus, √¢ge, localisation)  
                - Entreprises (par secteur ou taille)  
                - March√©s de niche  
                """)
                segments_clientele = st.text_area("D√©finissez vos segments de client√®le.", key="segments_clientele_ameliore")

            with st.expander("Ressources cl√©s", expanded=expand_all):
                st.markdown("""
                **Ressources cl√©s :**  
                Listez les ressources n√©cessaires pour ex√©cuter vos activit√©s. Par exemple :  
                - Ressources physiques (locaux, machines)  
                - Ressources humaines (comp√©tences cl√©s, √©quipes)  
                - Ressources financi√®res (fonds, pr√™ts)  
                """)
                ressources_cles = st.text_area("Listez vos ressources principales.", key="ressources_cles_ameliore")

            with st.expander("Canaux de distribution", expanded=expand_all):
                st.markdown("""
                **Canaux de distribution :**  
                D√©crivez comment vos produits ou services atteignent vos clients. Par exemple :  
                - Boutiques physiques  
                - Plateformes en ligne  
                - Distributeurs tiers  
                """)
                canaux_distribution = st.text_area("D√©crivez vos canaux de distribution.", key="canaux_distribution_ameliore")

            with st.expander("Structure de co√ªts", expanded=expand_all):
                st.markdown("""
                **Structure de co√ªts :**  
                √ânum√©rez les principaux co√ªts li√©s √† votre activit√©. Par exemple :  
                - Co√ªts de production  
                - Salaires et charges sociales  
                - D√©penses marketing et publicitaires  
                """)
                structure_couts = st.text_area("D√©crivez votre structure de co√ªts.", key="structure_couts_ameliore")

            with st.expander("Sources de revenus", expanded=expand_all):
                st.markdown("""
                **Sources de revenus :**  
                D√©crivez comment vous g√©n√©rez des revenus. Par exemple :  
                - Vente de produits ou services  
                - Abonnements  
                - Publicit√© ou partenariats  
                """)
                sources_revenus = st.text_area("D√©crivez vos sources de revenus.", key="sources_revenus_ameliore")

            submit_ameliore = st.form_submit_button("Valider les Informations d'Am√©lioration")
        
        if submit_ameliore:
            # R√©cup√©rer les rubriques pour la deuxi√®me g√©n√©ration
            rubriques_ameliore = {
                "Partenaires cl√©s": partenaire_cles,
                "Activit√©s cl√©s": activites_cles,
                "Offre (proposition de valeur)": offre_valeur,
                "Relation client": relation_client,
                "Segments de client√®le": segments_clientele,
                "Ressources cl√©s": ressources_cles,
                "Canaux de distribution": canaux_distribution,
                "Structure de co√ªts": structure_couts,
                "Sources de revenus": sources_revenus
            }
            
            # G√©n√©rer le BMC am√©lior√© en utilisant le BMC pr√©c√©dent et les nouvelles rubriques
            contenu_bmc_ameliore = obtenir_business_model(
                nom_entreprise=st.session_state.nom_entreprise,
                type_entreprise=st.session_state.type_entreprise,
                rubriques=rubriques_ameliore,
                previousdata=st.session_state.business_model_precedent,
                generation=2
            )
            
            if not contenu_bmc_ameliore:
                st.error("Erreur lors de la g√©n√©ration du contenu am√©lior√©. Veuillez r√©essayer.")
            else:
                # G√©n√©rer le document Word en m√©moire
                doc = Document()
                docx_bytes_ameliore = generer_docx_business_model(
                    nom_entreprise=st.session_state.nom_entreprise,
                    date_bmc=st.session_state.get('date_bmc_generate', datetime.date.today()).strftime("%d %B %Y"),
                    contenu_business_model=contenu_bmc_ameliore,
                    doc=doc,
                    value=1)
                
                
                st.success("Business Model Canvas am√©lior√© g√©n√©r√© avec succ√®s !")
                
                # Proposer le t√©l√©chargement du document Word am√©lior√©
                st.download_button(
                    label="T√©l√©charger le Business Model Canvas Am√©lior√© (Word)",
                    data=docx_bytes_ameliore,
                    file_name=f"BMC_Ameliore_{st.session_state.nom_entreprise.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Stocker le BMC am√©lior√© dans la session
                st.session_state.business_model_precedent = contenu_bmc_ameliore
                
                # Optionnel : Afficher le contenu g√©n√©r√© pour v√©rification
                st.subheader("Contenu Am√©lior√© G√©n√©r√© par ChatGPT")
                st.markdown(contenu_bmc_ameliore, unsafe_allow_html=True)


def page_collecte_donnees():
    st.header("√âtape 1 : Collecte des Donn√©es")
    st.write("Veuillez remplir les informations initiales pour g√©n√©rer le Business Model Canvas.")
    
    # Cr√©er des sous-onglets pour chaque section de collecte
    collecte_tabs = st.tabs([
        "Persona",
        "Arbre √† Probl√®me",
        "Analyse du March√©",
        "Facteurs Limitants",
        "Concurrence",
    ])
    
    # Collecte de Persona
    with collecte_tabs[0]:
        if 'type_entreprise' not in st.session_state:
            st.warning("Veuillez s√©lectionner le type d'entreprise dans la barre lat√©rale.")
        else:
            if st.session_state.type_entreprise == "PME":
                st.subheader("Collecte de Persona pour PME")
                
                # S√©lection du type de persona **hors** du formulaire
                type_persona = collect_persona_pme()
                
                # Formulaire pour collecter les d√©tails du persona
                with st.form("form_persona_pme"):
                    persona = collect_persona_details(type_persona)
                    submit_persona = st.form_submit_button("Valider Persona")
                
                if submit_persona:
                    st.session_state.persona = persona
                    st.success("Donn√©es Persona enregistr√©es avec succ√®s !")
            
            elif st.session_state.type_entreprise == "Startup":
                st.subheader("Collecte de Persona pour Startup")
                
                # S√©lection du type de persona **hors** du formulaire
                type_persona = collect_persona_pme()
                
                # Formulaire pour collecter les d√©tails du persona
                with st.form("form_persona_startup"):
                    persona = collect_persona_details(type_persona)  # Assurez-vous d'avoir une fonction similaire pour les startups si n√©cessaire
                    submit_persona = st.form_submit_button("Valider Persona")
                
                if submit_persona:
                    st.session_state.persona = persona
                    st.success("Donn√©es Persona enregistr√©es avec succ√®s !")

    # Collecte de l'Arbre √† Probl√®me
    with collecte_tabs[1]:
        if st.session_state.type_entreprise == "PME":
            with st.form("collect_problem_tree_pme"):
                problem_tree = collect_problem_tree_pme()
                submit_problem_tree = st.form_submit_button("Valider Arbre √† Probl√®me")
            
            if submit_problem_tree:
                st.session_state.problem_tree = problem_tree
                st.success("Arbre √† Probl√®me enregistr√© avec succ√®s !")
        elif st.session_state.type_entreprise == "Startup":
            with st.form("collect_problem_tree_startup"):
                problem_tree = collect_problem_tree_pme()  # Assurez-vous d'utiliser la bonne fonction
                submit_problem_tree = st.form_submit_button("Valider Arbre √† Probl√®me")
            
            if submit_problem_tree:
                st.session_state.problem_tree = problem_tree
                st.success("Arbre √† Probl√®me enregistr√© avec succ√®s !")

    # Collecte d'Analyse du March√©
    with collecte_tabs[2]:
        if st.session_state.type_entreprise == "PME":
            with st.form("form_analyse_marche_pme"):
                analyse_marche = collect_analyse_marche_pme()
                submit_analyse_marche = st.form_submit_button("Valider Analyse du March√©")
            
            if submit_analyse_marche:
                st.session_state.analyse_marche = analyse_marche
                st.success("Analyse du March√© enregistr√©e avec succ√®s !")
        elif st.session_state.type_entreprise == "Startup":
            with st.form("form_analyse_marche_startup"):
                analyse_marche = collect_analyse_marche_startup()
                submit_analyse_marche = st.form_submit_button("Valider Analyse du March√©")
            
            if submit_analyse_marche:
                st.session_state.analyse_marche = analyse_marche
                st.success("Analyse du March√© enregistr√©e avec succ√®s !")

    # Collecte de Facteurs Limitants
    with collecte_tabs[3]:
        if st.session_state.type_entreprise == "PME":
            with st.form("form_facteurs_limitants_pme"):
                facteurs_limitants = collect_facteurs_limitants_pme()
                submit_facteurs_limitants = st.form_submit_button("Valider Facteurs Limitants")
            
            if submit_facteurs_limitants:
                st.session_state.facteurs_limitants = facteurs_limitants
                st.success("Facteurs Limitants enregistr√©s avec succ√®s !")
        elif st.session_state.type_entreprise == "Startup":
            with st.form("form_facteurs_limitants_startup"):
                facteurs_limitants = collect_facteurs_limitants_startup()
                submit_facteurs_limitants = st.form_submit_button("Valider Facteurs Limitants")
            
            if submit_facteurs_limitants:
                st.session_state.facteurs_limitants = facteurs_limitants
                st.success("Facteurs Limitants enregistr√©s avec succ√®s !")

    # Collecte de Concurrence
    with collecte_tabs[4]:
        if st.session_state.type_entreprise == "PME":
            with st.form("form_concurrence_pme"):
                concurrence = collect_concurrence_pme()
                submit_concurrence = st.form_submit_button("Valider Concurrence")
            
            if submit_concurrence:
                st.session_state.concurrence = concurrence
                st.success("√âvaluation de la Concurrence enregistr√©e avec succ√®s !")
        elif st.session_state.type_entreprise == "Startup":
            with st.form("form_concurrence_startup"):
                concurrence = collect_concurrence_startup()
                submit_concurrence = st.form_submit_button("Valider Concurrence")
            
            if submit_concurrence:
                st.session_state.concurrence = concurrence
                st.success("√âvaluation de la Concurrence enregistr√©e avec succ√®s !")

    # Bouton pour G√©n√©rer le BMC Initial apr√®s avoir collect√© toutes les donn√©es
    with st.form("form_generate_initial"):
        st.write("Apr√®s avoir collect√© toutes les donn√©es, cliquez sur le bouton ci-dessous pour g√©n√©rer le Business Model Canvas initial.")
        submit_generate_initial = st.form_submit_button("G√©n√©rer BMC Initial")
    
    if submit_generate_initial:
        # V√©rifier que toutes les donn√©es sont collect√©es
        required_fields = [
            'persona', 'analyse_marche', 'facteurs_limitants',
            'concurrence'
        ]
        missing_fields = [field for field in required_fields if field not in st.session_state]
        
        if missing_fields:
            st.error(f"Veuillez compl√©ter toutes les sections de collecte des donn√©es avant de g√©n√©rer le BMC. Sections manquantes : {', '.join(missing_fields)}")
        elif not st.session_state.nom_entreprise:
            st.error("Veuillez entrer le nom de votre entreprise dans la barre lat√©rale.")
        else:
            # Combiner toutes les rubriques initiales en un seul dictionnaire
            rubriques_initiales = {
                "persona": st.session_state.persona,
                "problem_tree": st.session_state.problem_tree,
                "analyse_marche": st.session_state.analyse_marche,
                "facteurs_limitants": st.session_state.facteurs_limitants,
                "concurrence": st.session_state.concurrence,
            }
            
            # R√©cup√©rer la date du BMC
            date_bmc = st.date_input("Date du BMC", value=datetime.date.today(), key="date_bmc_generate")
            
            # G√©n√©rer le premier BMC
            contenu_bmc_initial = obtenir_business_model(                          
                nom_entreprise=st.session_state.nom_entreprise,
                type_entreprise=st.session_state.type_entreprise,
                rubriques=rubriques_initiales,
                previousdata="",
                generation=1
            )
            
            if not contenu_bmc_initial:
                st.error("Erreur lors de la g√©n√©ration du contenu initial. Veuillez r√©essayer.")
            else:
                # G√©n√©rer le document Word en m√©moire
                doc = Document()
                docx_bytes_initial = generer_docx_business_model(
                    nom_entreprise=st.session_state.nom_entreprise,
                    date_bmc=date_bmc.strftime("%d %B %Y"),
                    contenu_business_model=contenu_bmc_initial,
                    doc=doc,
                    value=1
                )
                
                st.success("Business Model Canvas initial g√©n√©r√© avec succ√®s !")
                
                # Proposer le t√©l√©chargement du document Word
                st.download_button(
                    label="T√©l√©charger le Business Model Canvas Initial (Word)",
                    data=docx_bytes_initial,
                    file_name=f"BMC_Initial_{st.session_state.nom_entreprise.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Stocker le BMC initial dans la session pour la deuxi√®me g√©n√©ration
                st.session_state.business_model_precedent = contenu_bmc_initial
                
                # Optionnel : Afficher le contenu g√©n√©r√© pour v√©rification
                st.subheader("Contenu Initial G√©n√©r√© par ChatGPT")
                st.markdown(contenu_bmc_initial, unsafe_allow_html=True)






# ----------------------------------------------------------------------------
# Business plan 
# ----------------------------------------------------------------------------




# Initialiser le dictionnaire principal dans session_state
if "data" not in st.session_state:
    st.session_state["data"] = {}
    
# Section 1 : Informations G√©n√©rales
def page_informations_generales():
    st.title("Informations G√©n√©rales")
    
    # Acc√®s au dictionnaire principal
    data = st.session_state["data"]
    
    # Collecte des entr√©es et stockage dans le dictionnaire principal
    data["informations_generales"] = data.get("informations_generales", {})
    info = data["informations_generales"]
    info["prenom_nom"] = st.text_input("Pr√©nom, nom :", value=info.get("prenom_nom", ""))
    info["intitule_projet"] = st.text_input("Intitul√© de votre projet :", value=info.get("intitule_projet", ""))
    info["statut_juridique"] = st.selectbox(
        "Votre statut juridique :",
        ["Micro-entreprise", "EURL", "SARL", "SAS", "SASU"],
        index=["Micro-entreprise", "EURL", "SARL", "SAS", "SASU"].index(info.get("statut_juridique", "Micro-entreprise"))
    )
    info["telephone"] = st.text_input("Votre num√©ro de t√©l√©phone :", value=info.get("telephone", ""))
    info["email"] = st.text_input("Votre adresse e-mail :", value=info.get("email", ""))
    info["ville"] = st.text_input("Votre ville ou commune d'activit√© :", value=info.get("ville", ""))
    info["type_vente"] = st.selectbox(
        "Vente de marchandises ou de services ?",
        ["Marchandises", "Services", "Mixte"],
        index=["Marchandises", "Services", "Mixte"].index(info.get("type_vente", "Marchandises"))
    )
    
    # Mise √† jour des donn√©es dans le dictionnaire principal
    st.session_state["data"]["informations_generales"] = info

def page_besoins_demarrage():
    st.title("Besoins de D√©marrage")
    
    # Acc√®s au dictionnaire principal
    data = st.session_state.get("data", {})
    
    # Liste des besoins r√©organis√©e
    besoins = [
        "Frais d‚Äô√©tablissement", 
        "Logiciels, formations",
        "Mat√©riel professionnel",
        "Mat√©riel autre",
        "Mat√©riel de bureau",
        "Stock de mati√®res et produits",
        "Enseigne et √©l√©ments de communication",
        "V√©hicule",
        "Frais de dossier",
        "Frais de notaire",
        "Tr√©sorerie de d√©part",
        "Frais d‚Äôouverture de compteurs",
        "D√©p√¥t de marque",
        "Droits d‚Äôentr√©e",
        "Achat fonds de commerce ou parts",
        "Droit au bail",
        "Caution ou d√©p√¥t de garantie"
    ]
    
    # Initialiser le dictionnaire pour stocker les besoins
    data["besoins_demarrage"] = data.get("besoins_demarrage", {})
    besoins_demarrage = data["besoins_demarrage"]
    
    total_besoins = 0.0
    
    for besoin in besoins:
        montant = st.number_input(
            f"{besoin} ($)",
            min_value=0.0,
            key=f"besoin_{besoin}",
            value=besoins_demarrage.get(besoin, 0.0)
        )
        besoins_demarrage[besoin] = montant
        total_besoins += montant
    
    data["total_besoins"] = total_besoins
    
    st.write("---")
    st.markdown(f"**Total des Besoins de D√©marrage :** {total_besoins:.2f} $")
    
    # Dur√©e d'amortissement
    data["duree_amortissement"] = st.number_input(
        "Dur√©e d'amortissement des investissements (en ann√©es) :",
        min_value=1,
        key="duree_amortissement",
        value=data.get("duree_amortissement", 3)
    )
    
    # Mise √† jour des donn√©es dans le dictionnaire principal
    st.session_state["data"] = data



def calculer_pret_interet_fixe(montant, taux_annuel, duree_mois):
    """
    Calcule les d√©tails d'un pr√™t avec int√©r√™ts fixes par mois.

    Args:
        montant (float): Montant du pr√™t en euros.
        taux_annuel (float): Taux d'int√©r√™t annuel en pourcentage.
        duree_mois (int): Dur√©e du pr√™t en mois.

    Returns:
        dict: D√©tails du pr√™t incluant mensualit√©, total √† rembourser, principal mensuel,
              int√©r√™ts totaux et int√©r√™ts par ann√©e.
    """
    if duree_mois <= 0:
        return {
            "mensualite": 0.0,
            "total_a_rembourser": 0.0,
            "principal_mensuel": 0.0,
            "interet_mensuel": 0.0,
            "interets_totaux": 0.0,
            "interets_annee1": 0.0,
            "interets_annee2": 0.0,
            "interets_annee3": 0.0
        }

    taux_mensuel = taux_annuel / 100 / 12

    # Calcul de la mensualit√© en utilisant la formule PMT
    try:
        mensualite = (taux_mensuel * montant) / (1 - (1 + taux_mensuel) ** (-duree_mois))
    except ZeroDivisionError:
        mensualite = 0.0

    # Principal mensuel fixe
    principal_mensuel = montant / duree_mois

    # Int√©r√™t mensuel
    interet_mensuel = mensualite - principal_mensuel

    # Total √† rembourser
    total_a_rembourser = mensualite * duree_mois

    # Int√©r√™ts totaux
    interets_totaux = interet_mensuel * duree_mois

    # Int√©r√™ts par ann√©e, limit√©s √† 12 mois maximum
    interets_annee1 = interet_mensuel * min(duree_mois, 12)
    interets_annee2 = interet_mensuel * min(max(duree_mois - 12, 0), 12)
    interets_annee3 = interet_mensuel * min(max(duree_mois - 24, 0), 12)

    return {
        "mensualite": round(mensualite, 2),
        "total_a_rembourser": round(total_a_rembourser, 2),
        "principal_mensuel": round(principal_mensuel, 2),
        "interet_mensuel": round(interet_mensuel, 2),
        "interets_totaux": round(interets_totaux, 2),
        "interets_annee1": round(interets_annee1, 2),
        "interets_annee2": round(interets_annee2, 2),
        "interets_annee3": round(interets_annee3, 2)
    }  

def page_financement():
    st.title("Financement des Besoins de D√©marrage")
    
    data = st.session_state.get("data", {})
    
    # Initialiser la section des financements
    if "financements" not in data:
        data["financements"] = {}
    
    financements_dict = data["financements"]
    
    total_financement = 0.0
    
    st.subheader("Apports")
    
    # Apport personnel ou familial
    apport_personnel = st.number_input(
        "Apport personnel ou familial ($)",
        min_value=0.0,
        key="financement_apport_personnel",
        value=financements_dict.get("Apport personnel ou familial", 0.00)
    )
    financements_dict["Apport personnel ou familial"] = apport_personnel
    total_financement += apport_personnel
    
    # Apports en nature (en valeur)
    apport_nature = st.number_input(
        "Apports en nature (en valeur) ($)",
        min_value=0.0,
        key="financement_apport_nature",
        value=financements_dict.get("Apports en nature (en valeur)", 0.00)
    )
    financements_dict["Apports en nature (en valeur)"] = apport_nature
    total_financement += apport_nature
    
    st.subheader("Pr√™ts")
    
    # Nombre de pr√™ts (maximum 3)
    num_prets = 3  # Limit√© √† 3 pr√™ts comme demand√©
    
    interets_prets = {
        "annee1": 0.0,
        "annee2": 0.0,
        "annee3": 0.0
    }
    
    for i in range(1, num_prets + 1):
        st.markdown(f"#### Pr√™t {i}")
        pret_name = st.text_input(
            f"Nom du pr√™t {i}",
            value=financements_dict.get(f"Pr√™t {i}", {}).get("nom", f"Pr√™t {i}"),
            key=f"pret_{i}_nom"
        )
        pret_montant = st.number_input(
            f"Montant du {pret_name} ($)",
            min_value=0.0,
            value=financements_dict.get(f"Pr√™t {i}", {}).get("montant", 0.0),
            key=f"pret_{i}_montant"
        )
        pret_taux = st.number_input(
            f"Taux du {pret_name} (%)",
            min_value=0.0,
            max_value=100.0,
            value=financements_dict.get(f"Pr√™t {i}", {}).get("taux", 0.0),
            key=f"pret_{i}_taux"
        )
        pret_duree = st.number_input(
            f"Dur√©e du {pret_name} (en mois)",
            min_value=1,
            value=financements_dict.get(f"Pr√™t {i}", {}).get("duree", 12),
            key=f"pret_{i}_duree"
        )
        
        # Stocker les d√©tails du pr√™t
        financements_dict[f"Pr√™t {i}"] = {
            "nom": pret_name,
            "montant": pret_montant,
            "taux": pret_taux,
            "duree": pret_duree
        }
        total_financement += pret_montant
        
        # Calculer les d√©tails du remboursement du pr√™t
        if pret_montant > 0 and pret_taux > 0 and pret_duree > 0:
            pret_info = calculer_pret_interet_fixe(pret_montant, pret_taux, pret_duree)
            # Stocker les r√©sultats du calcul
            financements_dict[f"Pr√™t {i}"].update(pret_info)
            # Ajouter les int√©r√™ts par ann√©e
            interets_prets["annee1"] += pret_info["interets_annee1"]
            interets_prets["annee2"] += pret_info["interets_annee2"]
            interets_prets["annee3"] += pret_info["interets_annee3"]
            
            # Afficher les d√©tails du pr√™t pour v√©rification
            st.write(f"**D√©tails du {pret_name}:**")
            st.write(f"Mensualit√© : {pret_info['mensualite']:.2f} $")
            st.write(f"Total √† rembourser : {pret_info['total_a_rembourser']:.2f} $")
            st.write(f"Principal mensuel : {pret_info['principal_mensuel']:.2f} $")
            st.write(f"Int√©r√™t mensuel : {pret_info['interet_mensuel']:.2f} $")
            st.write(f"Int√©r√™ts totaux : {pret_info['interets_totaux']:.2f} $")
            st.write(f"Int√©r√™ts Ann√©e 1 : {pret_info['interets_annee1']:.2f} $")
            st.write(f"Int√©r√™ts Ann√©e 2 : {pret_info['interets_annee2']:.2f} $")
            st.write(f"Int√©r√™ts Ann√©e 3 : {pret_info['interets_annee3']:.2f} $")
            st.write("---")
    
    st.subheader("Subventions")
    
    # Nombre de subventions (maximum 2)
    num_subventions = 2  # Limit√© √† 2 subventions comme demand√©
    
    for i in range(1, num_subventions + 1):
        st.markdown(f"#### Subvention {i}")
        subvention_name = st.text_input(
            f"Nom de la subvention {i}",
            value=financements_dict.get(f"Subvention {i}", {}).get("nom", f"Subvention {i}"),
            key=f"subvention_{i}_nom"
        )
        subvention_montant = st.number_input(
            f"Montant de {subvention_name} ($)",
            min_value=0.0,
            value=financements_dict.get(f"Subvention {i}", {}).get("montant", 0.0),
            key=f"subvention_{i}_montant"
        )
        # Stocker les d√©tails de la subvention
        financements_dict[f"Subvention {i}"] = {
            "nom": subvention_name,
            "montant": subvention_montant
        }
        total_financement += subvention_montant
    
    st.subheader("Autres Financements")
    
    # Autre financement
    autre_financement = st.number_input(
        "Autre financement ($)",
        min_value=0.0,
        key="financement_autre",
        value=financements_dict.get("Autre financement", 0.00)
    )
    financements_dict["Autre financement"] = autre_financement
    total_financement += autre_financement
    
    st.write("---")
    st.markdown(f"**Total des Financements :** {total_financement:,.2f} $")
    
    # Validation du total des financements
    besoin_total = data.get("besoins", 0.0)  # Assurez-vous que cette cl√© existe dans vos donn√©es
    if besoin_total > 0 and total_financement != besoin_total:
        st.error(f"Le total des financements ({total_financement:,.2f} $) ne correspond pas au besoin total ({besoin_total:,.2f} $). Veuillez ajuster les montants.")
    elif besoin_total > 0:
        st.success(f"Le total des financements correspond au besoin total ({besoin_total:,.2f} $).")
    
    # Stocker les donn√©es dans la session
    data["financements"] = financements_dict
    data["total_financement"] = total_financement
    data["interets_prets"] = interets_prets  # Stocker les int√©r√™ts des pr√™ts
    
    st.session_state["data"] = data



def page_charges_fixes():
    st.title("Charges Fixes sur 3 Ann√©es")
    
    data = st.session_state.get("data", {})
    
    charges_fixes = [
        "Assurances v√©hicule et RC pro", "T√©l√©phone, internet", "Autres abonnements",
        "Carburant", "Frais de d√©placement / h√©bergement", "Eau, √©lectricit√©, gaz",
        "Mutuelle", "Fournitures diverses", "Entretien Moto livraison et mat√©riel",
        "Nettoyage des locaux", "Budget publicit√© et communication", "Emplacements",
        "Expert comptable, avocats", "Frais bancaires et terminal carte bleue", "Taxes, CFE"
    ]
    
    # Initialisation des charges fixes si non pr√©sentes
    if "charges_fixes" not in data:
        data["charges_fixes"] = {"annee1": {}, "annee2": {}, "annee3": {}}
        for charge in charges_fixes:
            data["charges_fixes"]["annee1"][charge] = 0.0
            data["charges_fixes"]["annee2"][charge] = 0.0
            data["charges_fixes"]["annee3"][charge] = 0.0
    charges_fixes_dict = data["charges_fixes"]
    
    # Initialisation des charges suppl√©mentaires si non pr√©sentes
    if "charges_supplementaires" not in data:
        data["charges_supplementaires"] = []
    
    # Fonctions de mise √† jour
    def update_year1(charge):
        year1_key = f"charge_{charge}_annee1"
        year2_key = f"charge_{charge}_annee2"
        year3_key = f"charge_{charge}_annee3"
        
        year1_val = st.session_state.get(year1_key, 0.0)
        
        # Mettre √† jour ann√©e 2 et 3 seulement si l'utilisateur n'a pas d√©j√† modifi√© ces champs
        if st.session_state.get(f"updated_{year2_key}", False) == False:
            st.session_state[year2_key] = year1_val
            charges_fixes_dict["annee2"][charge] = year1_val
        if st.session_state.get(f"updated_{year3_key}", False) == False:
            st.session_state[year3_key] = year1_val
            charges_fixes_dict["annee3"][charge] = year1_val

    def update_year2(charge):
        year2_key = f"charge_{charge}_annee2"
        year3_key = f"charge_{charge}_annee3"
        
        year2_val = st.session_state.get(year2_key, 0.0)
        
        # Mettre √† jour ann√©e 3 seulement si l'utilisateur n'a pas d√©j√† modifi√© ce champ
        if st.session_state.get(f"updated_{year3_key}", False) == False:
            st.session_state[year3_key] = year2_val
            charges_fixes_dict["annee3"][charge] = year2_val

    def update_year3(charge):
        # Indiquer que l'ann√©e 3 a √©t√© mise √† jour manuellement
        year3_key = f"charge_{charge}_annee3"
        st.session_state[f"updated_{year3_key}"] = True

    st.subheader("Charges Fixes par D√©faut")
    for charge in charges_fixes:
        col1, col2, col3 = st.columns(3)
        with col1:
            year1_key = f"charge_{charge}_annee1"
            if year1_key not in st.session_state:
                st.session_state[year1_key] = charges_fixes_dict["annee1"].get(charge, 0.0)
            montant1 = st.number_input(
                f"{charge} - Ann√©e 1 ($)",
                min_value=0.0,
                key=year1_key,
                on_change=update_year1,
                args=(charge,),
                value=st.session_state[year1_key]
            )
            charges_fixes_dict["annee1"][charge] = montant1
        with col2:
            year2_key = f"charge_{charge}_annee2"
            if year2_key not in st.session_state:
                st.session_state[year2_key] = charges_fixes_dict["annee2"].get(charge, 0.0)
                st.session_state[f"updated_{year2_key}"] = False
            montant2 = st.number_input(
                f"{charge} - Ann√©e 2 ($)",
                min_value=0.0,
                key=year2_key,
                on_change=update_year2,
                args=(charge,),
                value=st.session_state[year2_key]
            )
            charges_fixes_dict["annee2"][charge] = montant2
        with col3:
            year3_key = f"charge_{charge}_annee3"
            if year3_key not in st.session_state:
                st.session_state[year3_key] = charges_fixes_dict["annee3"].get(charge, 0.0)
                st.session_state[f"updated_{year3_key}"] = False
            montant3 = st.number_input(
                f"{charge} - Ann√©e 3 ($)",
                min_value=0.0,
                key=year3_key,
                on_change=update_year3,
                args=(charge,),
                value=st.session_state[year3_key]
            )
            charges_fixes_dict["annee3"][charge] = montant3
        
    # Charges suppl√©mentaires
    st.write("---")
    st.subheader("Ajouter des Charges Suppl√©mentaires")
    
    nouvelle_charge = st.text_input("Nom de la nouvelle charge :", key="nouvelle_charge")
    
    if st.button("Ajouter la charge"):
        nouvelle_charge = nouvelle_charge.strip()
        if nouvelle_charge and nouvelle_charge not in data["charges_supplementaires"]:
            data["charges_supplementaires"].append(nouvelle_charge)
            charges_fixes_dict["annee1"][nouvelle_charge] = 0.0
            charges_fixes_dict["annee2"][nouvelle_charge] = 0.0
            charges_fixes_dict["annee3"][nouvelle_charge] = 0.0
            # R√©initialiser le champ de texte
            st.session_state["nouvelle_charge"] = ""
    
    for charge in data["charges_supplementaires"]:
        col1, col2, col3 = st.columns(3)
        with col1:
            year1_key = f"charge_{charge}_supp_annee1"
            if year1_key not in st.session_state:
                st.session_state[year1_key] = charges_fixes_dict["annee1"].get(charge, 0.0)
            montant1 = st.number_input(
                f"{charge} - Ann√©e 1 ($)",
                min_value=0.0,
                key=year1_key,
                on_change=update_year1,
                args=(charge,),
                value=st.session_state[year1_key]
            )
            charges_fixes_dict["annee1"][charge] = montant1
        with col2:
            year2_key = f"charge_{charge}_supp_annee2"
            if year2_key not in st.session_state:
                st.session_state[year2_key] = charges_fixes_dict["annee2"].get(charge, 0.0)
                st.session_state[f"updated_{year2_key}"] = False
            montant2 = st.number_input(
                f"{charge} - Ann√©e 2 ($)",
                min_value=0.0,
                key=year2_key,
                on_change=update_year2,
                args=(charge,),
                value=st.session_state[year2_key]
            )
            charges_fixes_dict["annee2"][charge] = montant2
        with col3:
            year3_key = f"charge_{charge}_supp_annee3"
            if year3_key not in st.session_state:
                st.session_state[year3_key] = charges_fixes_dict["annee3"].get(charge, 0.0)
                st.session_state[f"updated_{year3_key}"] = False
            montant3 = st.number_input(
                f"{charge} - Ann√©e 3 ($)",
                min_value=0.0,
                key=year3_key,
                on_change=update_year3,
                args=(charge,),
                value=st.session_state[year3_key]
            )
            charges_fixes_dict["annee3"][charge] = montant3
    
    # Calcul des totaux
    total_annee1 = sum(charges_fixes_dict["annee1"].values())
    total_annee2 = sum(charges_fixes_dict["annee2"].values())
    total_annee3 = sum(charges_fixes_dict["annee3"].values())
    
    data["total_charges_fixes_annee1"] = total_annee1
    data["total_charges_fixes_annee2"] = total_annee2
    data["total_charges_fixes_annee3"] = total_annee3
    
    st.write("---")
    st.markdown(f"**Total Charges Fixes Ann√©e 1 :** {total_annee1:.2f} $")
    st.markdown(f"**Total Charges Fixes Ann√©e 2 :** {total_annee2:.2f} $")
    st.markdown(f"**Total Charges Fixes Ann√©e 3 :** {total_annee3:.2f} $")
    
    st.session_state["data"] = data

def page_chiffre_affaires():
    st.title("Chiffre d'Affaires Pr√©visionnel")
    
    data = st.session_state.get("data", {})
    type_vente = data.get("informations_generales", {}).get("type_vente", "Marchandises")
    
    data["chiffre_affaires"] = data.get("chiffre_affaires", {})
    chiffre_affaires_dict = data["chiffre_affaires"]
    
    mois = [f"Mois {i}" for i in range(1, 13)]
    
    # Fonctions de mise √† jour
    def update_jours_travailles(nom_vente):
        key_jours_mois1 = f"{nom_vente}_Mois 1_jours"
        new_val = st.session_state.get(key_jours_mois1, 0)
        for mois_nom in mois[1:]:
            key = f"{nom_vente}_{mois_nom}_jours"
            if not st.session_state.get(f"updated_{key}", False):
                st.session_state[key] = new_val
                chiffre_affaires_dict[key] = new_val

    def update_ca_moyen_jour(nom_vente):
        key_ca_mois1 = f"{nom_vente}_Mois 1_ca_moyen"
        new_val = st.session_state.get(key_ca_mois1, 0.0)
        for mois_nom in mois[1:]:
            key = f"{nom_vente}_{mois_nom}_ca_moyen"
            if not st.session_state.get(f"updated_{key}", False):
                st.session_state[key] = new_val
                chiffre_affaires_dict[key] = new_val

    def mark_updated(key):
        st.session_state[f"updated_{key}"] = True

    def calcul_chiffre_affaires(nom_vente):
        mois_list = [f"Mois {i}" for i in range(1, 13)]
        data_ca = []
        
        st.subheader(f"Ann√©e 1 - {nom_vente}")
        for mois_nom in mois_list:
            col1, col2, col3 = st.columns(3)
            key_jours = f"{nom_vente}_{mois_nom}_jours"
            key_ca_moyen = f"{nom_vente}_{mois_nom}_ca_moyen"
            key_ca = f"{nom_vente}_{mois_nom}_ca"
            
            with col1:
                if mois_nom == "Mois 1":
                    montant_jours = st.number_input(
                        f"{mois_nom} - Nombre de jours travaill√©s",
                        min_value=0,
                        key=key_jours,
                        value=chiffre_affaires_dict.get(key_jours, 0),
                        on_change=update_jours_travailles,
                        args=(nom_vente,)
                    )
                else:
                    montant_jours = st.number_input(
                        f"{mois_nom} - Nombre de jours travaill√©s",
                        min_value=0,
                        key=key_jours,
                        value=chiffre_affaires_dict.get(key_jours, 0),
                        on_change=lambda key=key_jours: mark_updated(key)
                    )
                chiffre_affaires_dict[key_jours] = montant_jours
            
            with col2:
                if mois_nom == "Mois 1":
                    montant_ca_moyen = st.number_input(
                        f"{mois_nom} - Chiffre d'affaires moyen / jour ($)",
                        min_value=0.0,
                        key=key_ca_moyen,
                        value=chiffre_affaires_dict.get(key_ca_moyen, 0.0),
                        on_change=update_ca_moyen_jour,
                        args=(nom_vente,)
                    )
                else:
                    montant_ca_moyen = st.number_input(
                        f"{mois_nom} - Chiffre d'affaires moyen / jour ($)",
                        min_value=0.0,
                        key=key_ca_moyen,
                        value=chiffre_affaires_dict.get(key_ca_moyen, 0.0),
                        on_change=lambda key=key_ca_moyen: mark_updated(key)
                    )
                chiffre_affaires_dict[key_ca_moyen] = montant_ca_moyen
            
            ca_mensuel = montant_jours * montant_ca_moyen
            chiffre_affaires_dict[key_ca] = ca_mensuel
            data_ca.append({
                "mois": mois_nom,
                "jours_travailles": montant_jours,
                "ca_moyen_jour": montant_ca_moyen,
                "ca_mensuel": ca_mensuel
            })
            
            with col3:
                st.write(f"CA mensuel: {ca_mensuel:.2f} $")
        
        df_ca = pd.DataFrame(data_ca)
        total_ca_annee1 = df_ca["ca_mensuel"].sum()
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee1"] = total_ca_annee1
        
        st.write("---")
        st.markdown(f"**Total Chiffre d'Affaires Ann√©e 1 ({nom_vente}) :** {total_ca_annee1:.2f} $")
        
        # Pourcentages d'augmentation
        key_aug_annee2 = f"{nom_vente}_augmentation_annee2"
        key_aug_annee3 = f"{nom_vente}_augmentation_annee3"
        pourcentage_augmentation_annee2 = st.number_input(
            f"Pourcentage d'augmentation du CA entre l'ann√©e 1 et l'ann√©e 2 (%) ({nom_vente})",
            min_value=0.0,
            key=key_aug_annee2,
            value=chiffre_affaires_dict.get(key_aug_annee2, 0.0)
        )
        chiffre_affaires_dict[key_aug_annee2] = pourcentage_augmentation_annee2
        pourcentage_augmentation_annee3 = st.number_input(
            f"Pourcentage d'augmentation du CA entre l'ann√©e 2 et l'ann√©e 3 (%) ({nom_vente})",
            min_value=0.0,
            key=key_aug_annee3,
            value=chiffre_affaires_dict.get(key_aug_annee3, 0.0)
        )
        chiffre_affaires_dict[key_aug_annee3] = pourcentage_augmentation_annee3
        
        total_ca_annee2 = total_ca_annee1 * (1 + pourcentage_augmentation_annee2 / 100)
        total_ca_annee3 = total_ca_annee2 * (1 + pourcentage_augmentation_annee3 / 100)
        
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee2"] = total_ca_annee2
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee3"] = total_ca_annee3
        
        st.markdown(f"**Total Chiffre d'Affaires Ann√©e 2 ({nom_vente}) :** {total_ca_annee2:.2f} $")
        st.markdown(f"**Total Chiffre d'Affaires Ann√©e 3 ({nom_vente}) :** {total_ca_annee3:.2f} $")
    
    if type_vente in ["Marchandises", "Mixte"]:
        calcul_chiffre_affaires("Marchandises")
    if type_vente in ["Services", "Mixte"]:
        calcul_chiffre_affaires("Services")
    
    # Calcul du total CA toutes ventes
    total_ca_annee1 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee1", 0.0) for type in ["Marchandises", "Services"]
    )
    total_ca_annee2 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee2", 0.0) for type in ["Marchandises", "Services"]
    )
    total_ca_annee3 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee3", 0.0) for type in ["Marchandises", "Services"]
    )
    
    data["total_chiffre_affaires_annee1"] = total_ca_annee1
    data["total_chiffre_affaires_annee2"] = total_ca_annee2
    data["total_chiffre_affaires_annee3"] = total_ca_annee3
    
    st.write("---")
    st.markdown(f"**Total Chiffre d'Affaires Ann√©e 1 (toutes ventes) :** {total_ca_annee1:.2f} $")
    st.markdown(f"**Total Chiffre d'Affaires Ann√©e 2 (toutes ventes) :** {total_ca_annee2:.2f} $")
    st.markdown(f"**Total Chiffre d'Affaires Ann√©e 3 (toutes ventes) :** {total_ca_annee3:.2f} $")
    
    st.session_state["data"] = data
# Section 6 : Charges Variables
def page_charges_variables():
    st.title("Charges Variables")
    
    data = st.session_state["data"]
    type_vente = data["informations_generales"].get("type_vente", "Marchandises")
    
    if type_vente in ["Marchandises", "Mixte"]:
        st.markdown("""
        ### Vos charges variables
        Les charges variables sont li√©es au niveau d‚Äôactivit√© ou √† la production. 
        Il s‚Äôagit des achats de marchandises destin√©es √† √™tre revendues, des achats de mati√®res destin√©es √† √™tre transform√©es, 
        des commissions vers√©es √† des agents commerciaux.
        """)
        
        data["charges_variables"] = data.get("charges_variables", {})
        charges_variables = data["charges_variables"]
        
        # Co√ªt d'achat des marchandises en %
        cout_achat_marchandises_pct = st.number_input(
            "Quel est, en % du prix de vente, le co√ªt d'achat de vos marchandises ? (concerne uniquement le chiffre d'affaires vente de marchandises)",
            min_value=0.0,
            max_value=100.0,
            format="%.2f",
            key="cout_achat_marchandises_pct",
            value=charges_variables.get("cout_achat_marchandises_pct", 0.0)
        )
        charges_variables["cout_achat_marchandises_pct"] = cout_achat_marchandises_pct
        
        st.write(f"Co√ªt d'achat des marchandises : {cout_achat_marchandises_pct:.2f}% du prix de vente")
        
        total_ca_marchandises_annee1 = data["chiffre_affaires"].get("total_ca_Marchandises_annee1", 0.0)
        total_charges_variables = total_ca_marchandises_annee1 * cout_achat_marchandises_pct / 100.0
        
        data["total_charges_variables"] = total_charges_variables
        
        st.write(f"Total des Charges Variables Ann√©e 1 : {total_charges_variables:.2f} $")
        
    else:
        st.info("Cette section est uniquement applicable si vous vendez des marchandises ou des services mixtes.")
        data["total_charges_variables"] = 0.0
    
    st.session_state["data"] = data

# Section 7 : Fonds de Roulement
def page_fonds_roulement():
    st.title("Votre Besoin en Fonds de Roulement")
    
    data = st.session_state["data"]
    
    st.markdown("""
    ### D√©terminez votre besoin en fonds de roulement
    Le fonds de roulement repr√©sente le montant n√©cessaire pour financer le cycle d'exploitation de votre entreprise.
    """)
    
    data["fonds_roulement"] = data.get("fonds_roulement", {})
    fonds_roulement = data["fonds_roulement"]
    
    duree_credits_clients = st.number_input(
        "Dur√©e moyenne des cr√©dits accord√©s aux clients (en jours) :",
        min_value=0,
        help="Temps moyen qu'un client met pour vous payer.",
        key="duree_credits_clients",
        value=fonds_roulement.get("duree_credits_clients", 0)
    )
    fonds_roulement["duree_credits_clients"] = duree_credits_clients
    
    duree_dettes_fournisseurs = st.number_input(
        "Dur√©e moyenne des dettes fournisseurs (en jours) :",
        min_value=0,
        help="Temps moyen que vous mettez pour payer vos fournisseurs.",
        key="duree_dettes_fournisseurs",
        value=fonds_roulement.get("duree_dettes_fournisseurs", 0)
    )
    fonds_roulement["duree_dettes_fournisseurs"] = duree_dettes_fournisseurs
    
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    
    bfr = (total_ca_annee1 * duree_credits_clients / 360) - (total_charges_variables * duree_dettes_fournisseurs / 360)
    fonds_roulement["bfr"] = bfr
    
    st.write("---")
    st.markdown(f"**Dur√©e moyenne des cr√©dits clients :** {duree_credits_clients} jours")
    st.markdown(f"**Dur√©e moyenne des dettes fournisseurs :** {duree_dettes_fournisseurs} jours")
    st.markdown(f"**Besoin en Fonds de Roulement (BFR) Ann√©e 1 :** {bfr:.2f} $")
    
    st.session_state["data"] = data

# Section 8 : Salaires
def page_salaires():
    st.title("Salaires Employ√©s et R√©mun√©ration Chef d'Entreprise")
    
    data = st.session_state["data"]
    data["salaires"] = data.get("salaires", {"employes": {}, "dirigeants": {}})
    salaires = data["salaires"]
    
    st.markdown("""
    ### Saisissez les salaires et r√©mun√©rations pour les 3 ann√©es
    Veuillez entrer les chiffres annuels pour les salaires des employ√©s et la r√©mun√©ration nette des dirigeants.
    """)
    
    st.subheader("Salaires Employ√©s (Net)")
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires["employes"][key] = st.number_input(
            f"Salaires Employ√©s Ann√©e {annee} ($)",
            min_value=0.0,
            key=f"salaires_employes_annee_{annee}",
            value=salaires["employes"].get(key, 0.0)
        )
    
    st.subheader("R√©mun√©ration Nette Dirigeant(s)")
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires["dirigeants"][key] = st.number_input(
            f"R√©mun√©ration Dirigeant Ann√©e {annee} ($)",
            min_value=0.0,
            key=f"remuneration_dirigeant_annee_{annee}",
            value=salaires["dirigeants"].get(key, 0.0)
        )
    
    st.write("---")
    accre = st.selectbox(
        "Le(s) dirigeant(s) b√©n√©ficient-ils de l'ACRE ?",
        options=["Oui", "Non"],
        key="accre",
        index=["Oui", "Non"].index(data.get("accre", "Non")),
        help="Veuillez s√©lectionner 'Oui' si les dirigeants b√©n√©ficient de l'ACRE. Cette question est obligatoire."
    )
    data["accre"] = accre
    
    total_salaires_annee1 = salaires["employes"]["annee1"] + salaires["dirigeants"]["annee1"]
    total_salaires_annee2 = salaires["employes"]["annee2"] + salaires["dirigeants"]["annee2"]
    total_salaires_annee3 = salaires["employes"]["annee3"] + salaires["dirigeants"]["annee3"]
    
    data["total_salaires_annee1"] = total_salaires_annee1
    data["total_salaires_annee2"] = total_salaires_annee2
    data["total_salaires_annee3"] = total_salaires_annee3
    
    st.write("---")
    st.markdown(f"**Total Salaires et R√©mun√©ration Ann√©e 1 :** {total_salaires_annee1:.2f} $")
    st.markdown(f"**Total Salaires et R√©mun√©ration Ann√©e 2 :** {total_salaires_annee2:.2f} $")
    st.markdown(f"**Total Salaires et R√©mun√©ration Ann√©e 3 :** {total_salaires_annee3:.2f} $")
    
    st.session_state["data"] = data

# Section 9 : Contr√¥le de Rentabilit√©
def page_rentabilite():
    st.title("Contr√¥le de Rentabilit√©")
    
    data = st.session_state["data"]
    
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    total_chiffre_affaires = data.get("total_chiffre_affaires_annee1", 0.0)
    total_salaires_annee1 = data.get("total_salaires_annee1", 0.0)
    
    if total_chiffre_affaires > 0:
        marge_brute = ((total_chiffre_affaires - total_charges_variables) / total_chiffre_affaires) * 100.0
    else:
        marge_brute = 0.0
    
    charges_fixes_totales = total_charges_fixes_annee1 + total_salaires_annee1
    if marge_brute > 0:
        seuil_rentabilite = charges_fixes_totales / (marge_brute / 100.0)
    else:
        seuil_rentabilite = 0.0
    
    if total_chiffre_affaires >= seuil_rentabilite:
        rentabilite = "Rentable"
        message_rentabilite = "L'entreprise est rentable."
        couleur_rentabilite = "green"
    else:
        rentabilite = "Non rentable"
        message_rentabilite = "L'entreprise n'est pas rentable. Il faut augmenter le chiffre d'affaires ou r√©duire les charges."
        couleur_rentabilite = "red"
    
    data["marge_brute"] = marge_brute
    data["seuil_rentabilite"] = seuil_rentabilite
    data["rentabilite"] = rentabilite
    
    st.write("---")
    st.markdown(f"**Marge Brute :** {marge_brute:.2f} %")
    st.markdown(f"**Seuil de Rentabilit√© :** {seuil_rentabilite:.2f} $")
    st.markdown(f"<div style='background-color:{couleur_rentabilite}; color:white; padding:10px; border-radius:5px; text-align:center;'>"
                f"<strong>{rentabilite}</strong> - {message_rentabilite}</div>", unsafe_allow_html=True)
    
    st.session_state["data"] = data

# Section 10 : Tr√©sorerie de D√©part
def page_tresorerie():
    st.title("Contr√¥le du Niveau de votre Tr√©sorerie de D√©part")
    data = st.session_state["data"]
    besoins_demarrage=data.get("besoins_demarrage", 0.0)
    tresorerie_depart1 = besoins_demarrage.get("Tr√©sorerie de d√©part", 0.0)
    
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    st.markdown(f"**Tr√©sorerie de d√©part :** {tresorerie_depart1 :.2f} $")
    tresorerie_depart=tresorerie_depart1
    
    data["tresorerie_depart"] = tresorerie_depart
    
    seuil_tresorerie = total_charges_fixes_annee1 / 4.0  # 3 mois de charges fixes
    if tresorerie_depart >= seuil_tresorerie:
        niveau_tresorerie = "Ad√©quate"
        message_tresorerie = "Votre tr√©sorerie de d√©part est ad√©quate pour couvrir les charges initiales."
        couleur_tresorerie = "green"
    else:
        niveau_tresorerie = "Trop faible"
        message_tresorerie = "Votre tr√©sorerie de d√©part est trop faible. Pr√©voyez plus de tr√©sorerie pour couvrir les charges."
        couleur_tresorerie = "red"
    
    data["niveau_tresorerie"] = niveau_tresorerie
    
    st.write("---")
    st.markdown(f"### R√©sultat pour la 1√®re ann√©e :")
    st.markdown(f"<div style='background-color:{couleur_tresorerie}; color:white; padding:10px; border-radius:5px; text-align:center;'>"
                f"<strong>{niveau_tresorerie}</strong> - {message_tresorerie}</div>", unsafe_allow_html=True)
    
    st.session_state["data"] = data

# Section 11 : R√©capitulatif
def page_recapitulatif():
    st.title("R√©capitulatif Complet des Donn√©es")
    
    data = st.session_state["data"]
    
    st.subheader("1. Informations G√©n√©rales")
    info = data.get("informations_generales", {})
    st.write(f"Pr√©nom, nom : {info.get('prenom_nom', '')}")
    st.write(f"Intitul√© du projet : {info.get('intitule_projet', '')}")
    st.write(f"Statut juridique : {info.get('statut_juridique', '')}")
    st.write(f"T√©l√©phone : {info.get('telephone', '')}")
    st.write(f"Email : {info.get('email', '')}")
    st.write(f"Ville : {info.get('ville', '')}")
    st.write(f"Type de vente : {info.get('type_vente', '')}")
    
    st.subheader("2. Besoins de D√©marrage")
    besoins = data.get("besoins_demarrage", {})
    total_besoins = data.get("total_besoins", 0.0)
    for besoin, montant in besoins.items():
        st.write(f"{besoin} : {montant:.2f} $")
    st.write(f"**Total des Besoins de D√©marrage : {total_besoins:.2f} $**")
    
    st.title("R√©capitulatif des Financements")
    data = st.session_state.get("data", {})
    financements_dict = data.get("financements", {})
    total_financement = data.get("total_financement", 0.0)
    st.subheader("Financements")
    for financement, details in financements_dict.items():
        if isinstance(details, dict):
            montant = details.get("montant", 0.0)
            st.write(f"{details.get('nom', financement)} : {montant:.2f} $")
        else:
            montant = details
            st.write(f"{financement} : {montant:.2f} $")
    
    st.markdown(f"**Total des Financements :** {total_financement:.2f} $")
    
    st.subheader("4. Charges Fixes sur 3 Ann√©es")
    charges_fixes_dict = data.get("charges_fixes", {"annee1": {}, "annee2": {}, "annee3": {}})
    total_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_annee2 = data.get("total_charges_fixes_annee2", 0.0)
    total_annee3 = data.get("total_charges_fixes_annee3", 0.0)
    charges_supp = data.get("charges_supplementaires", [])
    
    for charge in charges_fixes_dict["annee1"].keys():
        montant1 = charges_fixes_dict["annee1"].get(charge, 0.0)
        montant2 = charges_fixes_dict["annee2"].get(charge, 0.0)
        montant3 = charges_fixes_dict["annee3"].get(charge, 0.0)
        st.write(f"{charge} - Ann√©e 1 : {montant1:.2f} $, Ann√©e 2 : {montant2:.2f} $, Ann√©e 3 : {montant3:.2f} $")
    
    st.write(f"**Total Charges Fixes Ann√©e 1 : {total_annee1:.2f} $**")
    st.write(f"**Total Charges Fixes Ann√©e 2 : {total_annee2:.2f} $**")
    st.write(f"**Total Charges Fixes Ann√©e 3 : {total_annee3:.2f} $**")
    
    st.subheader("5. Chiffre d'Affaires Pr√©visionnel")
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_ca_annee2 = data.get("total_chiffre_affaires_annee2", 0.0)
    total_ca_annee3 = data.get("total_chiffre_affaires_annee3", 0.0)
    
    st.write(f"Total Chiffre d'Affaires Ann√©e 1 : {total_ca_annee1:.2f} $")
    st.write(f"Total Chiffre d'Affaires Ann√©e 2 : {total_ca_annee2:.2f} $")
    st.write(f"Total Chiffre d'Affaires Ann√©e 3 : {total_ca_annee3:.2f} $")
    
    st.subheader("6. Charges Variables")
    cout_achat_marchandises_pct = data.get("charges_variables", {}).get("cout_achat_marchandises_pct", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    st.write(f"Co√ªt d'achat des marchandises (% du CA) : {cout_achat_marchandises_pct:.2f} %")
    st.write(f"Total Charges Variables Ann√©e 1 : {total_charges_variables:.2f} $")
    
    st.subheader("7. Fonds de Roulement")
    fonds_roulement = data.get("fonds_roulement", {})
    duree_credits_clients = fonds_roulement.get("duree_credits_clients", 0)
    duree_dettes_fournisseurs = fonds_roulement.get("duree_dettes_fournisseurs", 0)
    bfr = fonds_roulement.get("bfr", 0.0)
    st.write(f"Dur√©e moyenne des cr√©dits clients : {duree_credits_clients} jours")
    st.write(f"Dur√©e moyenne des dettes fournisseurs : {duree_dettes_fournisseurs} jours")
    st.write(f"Besoin en Fonds de Roulement (BFR) Ann√©e 1 : {bfr:.2f} $")
    
    st.subheader("8. Salaires et R√©mun√©ration")
    salaires = data.get("salaires", {})
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires_employes = salaires.get("employes", {}).get(key, 0.0)
        remuneration_dirigeants = salaires.get("dirigeants", {}).get(key, 0.0)
        st.write(f"Ann√©e {annee} : Salaires employ√©s : {salaires_employes:.2f} $, R√©mun√©ration dirigeants : {remuneration_dirigeants:.2f} $")
        st.write(f"Total Salaires Ann√©e {annee} : {(salaires_employes + remuneration_dirigeants):.2f} $")
    
    st.subheader("9. Rentabilit√©")
    marge_brute = data.get("marge_brute", 0.0)
    seuil_rentabilite = data.get("seuil_rentabilite", 0.0)
    rentabilite = data.get("rentabilite", "Non rentable")
    st.write(f"Marge Brute : {marge_brute:.2f} %")
    st.write(f"Seuil de Rentabilit√© : {seuil_rentabilite:.2f} $")
    st.write(f"Rentabilit√© : {rentabilite}")
    
    st.subheader("10. Tr√©sorerie de D√©part")
    tresorerie_depart = data.get("tresorerie_depart", 0.0)
    niveau_tresorerie = data.get("niveau_tresorerie", "Trop faible")
    st.write(f"Montant de la Tr√©sorerie Initiale : {tresorerie_depart:.2f} $")
    st.write(f"Niveau de Tr√©sorerie : {niveau_tresorerie}")
    
    st.session_state["data"] = data
    
    
    
    
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def page_investissements_et_financements(): 
    st.title("Investissements et Financements")
    
    # Initialiser la cl√© 'export_data' dans session_state si elle n'existe pas
    if 'export_data' not in st.session_state:
        st.session_state['export_data'] = {}
    
    # R√©cup√©rer les donn√©es de la session
    data = st.session_state.get("data", {})
    
    # R√©cup√©rer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "N/A")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "N/A")
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Initialiser une liste pour stocker toutes les lignes du tableau
    table_data = []
    
    # Immobilisations Incorporelles
    immobilisations_incorporelles = {
        "Frais d‚Äô√©tablissement": data.get("besoins_demarrage", {}).get("Frais d‚Äô√©tablissement", 0.0),
        "Frais d‚Äôouverture de compteurs": data.get("besoins_demarrage", {}).get("Frais d‚Äôouverture de compteurs", 0.0),
        "Logiciels, formations": data.get("besoins_demarrage", {}).get("Logiciels, formations", 0.0),
        "D√©p√¥t de marque": data.get("besoins_demarrage", {}).get("D√©p√¥t de marque", 0.0),
        "Droits d‚Äôentr√©e": data.get("besoins_demarrage", {}).get("Droits d‚Äôentr√©e", 0.0),
        "Achat fonds de commerce ou parts": data.get("besoins_demarrage", {}).get("Achat fonds de commerce ou parts", 0.0),
        "Droit au bail": data.get("besoins_demarrage", {}).get("Droit au bail", 0.0),
        "Caution ou d√©p√¥t de garantie": data.get("besoins_demarrage", {}).get("Caution ou d√©p√¥t de garantie", 0.0),
        "Frais de dossier": data.get("besoins_demarrage", {}).get("Frais de dossier", 0.0),
        "Frais de notaire": data.get("besoins_demarrage", {}).get("Frais de notaire", 0.0),
    }
    total_incorporelles = sum(immobilisations_incorporelles.values())
    table_data.append({
        "Investissements": "Immobilisations incorporelles",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_incorporelles:.2f}"
    })
    for desc, montant in immobilisations_incorporelles.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Dur√©e (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # Immobilisations Corporelles
    immobilisations_corporelles = {
        "Enseigne et √©l√©ments de communication": data.get("besoins_demarrage", {}).get("Enseigne et √©l√©ments de communication", 0.0),
        "V√©hicule": data.get("besoins_demarrage", {}).get("V√©hicule", 0.0),
        "Mat√©riel professionnel": data.get("besoins_demarrage", {}).get("Mat√©riel professionnel", 0.0),
        "Mat√©riel autre": data.get("besoins_demarrage", {}).get("Mat√©riel autre", 0.0),
        "Mat√©riel de bureau": data.get("besoins_demarrage", {}).get("Mat√©riel de bureau", 0.0),
    }
    total_corporelles = sum(immobilisations_corporelles.values())
    table_data.append({
        "Investissements": "Immobilisations corporelles",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_corporelles:.2f}"
    })
    for desc, montant in immobilisations_corporelles.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Dur√©e (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # Autres Investissements
    autres_investissements = {
        "Stock de mati√®res et produits": data.get("besoins_demarrage", {}).get("Stock de mati√®res et produits", 0.0),
        "Tr√©sorerie de d√©part": data.get("besoins_demarrage", {}).get("Tr√©sorerie de d√©part", 0.0)
    }
    total_autres = sum(autres_investissements.values())
    table_data.append({
        "Investissements": "Autres investissements",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_autres:.2f}"
    })
    for desc, montant in autres_investissements.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Dur√©e (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # TOTAL BESOINS
    total_besoins = total_incorporelles + total_corporelles + total_autres
    table_data.append({
        "Investissements": "TOTAL BESOINS",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_besoins:.2f}"
    })
    
    # Section FINANCEMENT DES INVESTISSEMENTS
    table_data.append({
        "Investissements": "FINANCEMENT DES INVESTISSEMENTS",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": ""
    })
    table_data.append({
        "Investissements": "Montant $ hors taxes",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": ""
    })
    
    # Apport Personnel
    financements = data.get("financements", {})
    apport_personnel = {
        "Apport personnel ou familial": financements.get("Apport personnel ou familial", 0.0),
        "Apports en nature (en valeur)": financements.get("Apports en nature (en valeur)", 0.0),
    }
    total_apport_personnel = sum(apport_personnel.values())
    table_data.append({
        "Investissements": "Apport personnel",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_apport_personnel:.2f}"
    })
    for desc, montant in apport_personnel.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Dur√©e (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # Emprunts Dynamiques
    emprunts_keys = ["Pr√™t 1", "Pr√™t 2", "Pr√™t 3"]
    emprunts_list = []
    total_emprunts = 0.0

    # Collecter les d√©tails des emprunts
    for i, key in enumerate(emprunts_keys, start=1):
        pret = financements.get(key, {})
        nom = pret.get("nom", "")
        taux = pret.get("taux", 0.0)
        duree = pret.get("duree", 0)
        montant = pret.get("montant", 0.0)
        
        # D√©finir le nom de l'emprunt
        emprunt_nom = nom if nom else f"Pr√™t {i}"
        
        # Ajouter les d√©tails du pr√™t
        if montant > 0:
            emprunts_list.append({
                "Investissements": emprunt_nom,
                "Taux (%)": f"{taux:.2f}%",
                "Dur√©e (mois)": duree,
                "Montant ($)": f"{montant:.2f}"
            })
            total_emprunts += montant
        else:
            emprunts_list.append({
                "Investissements": emprunt_nom,
                "Taux (%)": "-",
                "Dur√©e (mois)": "-",
                "Montant ($)": "0.00"
            })

    # TOTAL EMPRUNTS - plac√© avant les emprunts individuels
    table_data.append({
        "Investissements": "TOTAL EMPRUNTS",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_emprunts:.2f}"
    })

    # Ajouter les emprunts individuels apr√®s le total
    for emprunt in emprunts_list:
        table_data.append(emprunt)

     # Subventions Dynamiques
    subventions_keys = ["Subvention 1", "Subvention 2"]
    subventions_list = []
    total_subventions = 0.0
    
    # Calculer le total des subventions d'abord
    for i, key in enumerate(subventions_keys, start=1):
        subv = financements.get(key, {})
        nom = subv.get("nom", "")
        montant = subv.get("montant", 0.0)
        
        # D√©finir le nom de la subvention
        subvention_nom = nom if nom else f"Subvention {i}"
        
        # Ajouter les d√©tails de la subvention
        if montant > 0:
            subventions_list.append({
                "Investissements": subvention_nom,
                "Taux (%)": "",
                "Dur√©e (mois)": "",
                "Montant ($)": f"{montant:.2f}"
            })
            total_subventions += montant
        else:
            subventions_list.append({
                "Investissements": subvention_nom,
                "Taux (%)": "",
                "Dur√©e (mois)": "",
                "Montant ($)": "0.00"
            })
    
    # TOTAL SUBVENTIONS - plac√© avant les subventions individuelles
    table_data.append({
        "Investissements": "TOTAL SUBVENTIONS",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_subventions:.2f}"
    })
    
    # Ajouter les subventions individuelles apr√®s le total
    for subv in subventions_list:
        table_data.append(subv)
    
    # Autre Financement
    autre_financement = financements.get("Autre financement", 0.0)
    table_data.append({
        "Investissements": "Autre financement",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{autre_financement:.2f}"
    })
    
    # TOTAL RESSOURCES
    total_ressources = total_apport_personnel + total_emprunts + total_subventions + autre_financement
    table_data.append({
        "Investissements": "TOTAL RESSOURCES",
        "Taux (%)": "",
        "Dur√©e (mois)": "",
        "Montant ($)": f"{total_ressources:.2f}"
    })
    
    # V√©rification de l'√©quilibre
    if total_ressources == total_besoins:
        equilibrium_message = "Le total des ressources couvre exactement les besoins."
        equilibrium_type = "success"
    elif total_ressources > total_besoins:
        surplus = total_ressources - total_besoins
        equilibrium_message = f"Les ressources d√©passent les besoins de {surplus:.2f} $."
        equilibrium_type = "info"
    else:
        deficit = total_besoins - total_ressources
        equilibrium_message = f"Il manque {deficit:.2f} $ pour couvrir les besoins."
        equilibrium_type = "warning"
    
    if equilibrium_type == "success":
        st.success(equilibrium_message)
    elif equilibrium_type == "info":
        st.info(equilibrium_message)
    else:
        st.warning(equilibrium_message)
    
    st.write("---")
    
    # Cr√©er le DataFrame unique avec les quatre colonnes
    df_unique = pd.DataFrame(table_data, columns=["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"])
    
    # Afficher le tableau dans Streamlit
    st.dataframe(df_unique.style.apply(lambda x: ['background-color: #f0f0f0' if pd.isna(v) else '' for v in x], axis=1))
    
    # Stocker les totaux dans la session
    data["total_investissements"] = total_besoins
    data["total_financements"] = total_ressources
    
    st.session_state["data"] = data
    
    # Stocker les donn√©es d'exportation dans la nouvelle session
    st.session_state['export_data_investissements'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": table_data,
        "equilibre": {
            "type": equilibrium_type,
            "message": equilibrium_message
        }
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le Markdown
    if st.button("T√©l√©charger en Markdown"):
        markdown_content = f"# Investissements et Financements\n\n**Projet :** {projet}\n\n**Porteur de projet :** {porteur_projet}\n\n"
        
        # Convertir le DataFrame en Markdown
        markdown_content += df_unique.to_markdown(index=False)
        markdown_content += f"\n\n---\n\n{equilibrium_message}\n"
        
        markdown_bytes = markdown_content.encode('utf-8')
        st.download_button(
            label="T√©l√©charger le Markdown",
            data=markdown_bytes,
            file_name="investissements_et_financements.md",
            mime="text/markdown"
        )
    
    # Bouton pour t√©l√©charger le fichier Word
    if st.button("T√©l√©charger en Word"):
        export_data = st.session_state.get('export_data', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donn√©e disponible pour l'exportation.")
        else:
            doc = Document()
            doc.add_heading('Investissements et Financements', level=1)
            doc.add_paragraph(f"**Projet :** {export_data['projet']}")
            doc.add_paragraph(f"**Porteur de projet :** {export_data['porteur_projet']}")
            doc.add_page_break()
            
            # Cr√©er le DataFrame pour Word
            df_word = pd.DataFrame(export_data['table_data'], columns=["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"])
            
            # Ajouter le tableau au document Word
            table = doc.add_table(rows=1, cols=len(df_word.columns))
            table.style = 'Light List Accent 1'  # Choisissez un style appropri√©
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df_word.columns):
                hdr_cells[i].text = column
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for index, row in df_word.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
                    # Mettre en gras les cat√©gories principales et les totaux
                    if row["Investissements"] in ["INVESTISSEMENTS", "Montant $ hors taxes",
                                                 "Immobilisations incorporelles", "Immobilisations corporelles",
                                                 "Autres investissements", "TOTAL BESOINS",
                                                 "FINANCEMENT DES INVESTISSEMENTS", "Apport personnel",
                                                 "Emprunt", "TOTAL EMPRUNTS", "Subvention",
                                                 "TOTAL SUBVENTIONS", "Autre financement", "TOTAL RESSOURCES"]:
                        run = row_cells[i].paragraphs[0].runs
                        if run:
                            run[0].font.bold = True
            doc.add_paragraph()
            doc.add_paragraph(export_data['equilibre']['message'])
            
            # Enregistrer le document dans un buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label="T√©l√©charger le fichier Word",
                data=buffer,
                file_name="investissements_et_financements.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )



import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_salaires_charges_sociales():
    st.title("Salaires et Charges Sociales")
    
    # Initialiser la cl√© 'export_data_salaires_charges_sociales' dans session_state si elle n'existe pas
    if 'export_data_salaires_charges_sociales' not in st.session_state:
        st.session_state['export_data_salaires_charges_sociales'] = {}
    
    data = st.session_state.get("data", {})
    
    # R√©cup√©rer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "")
    statut_juridique = data.get("informations_generales", {}).get("statut_juridique", "")
    benefice_accre = data.get("accre", "Non")  # Assurez-vous que cette information est bien stock√©e dans data
    
    # D√©terminer le statut social du dirigeant en fonction du statut juridique
    if statut_juridique in ["Entreprise individuelle", "EURL (IS)", "EIRL (IS)", "Micro-entreprise"]:
        statut_social_dirigeant = "Travailleur Non Salari√© (TNS)"
    elif statut_juridique in ["SARL (IS)", "SAS (IS)", "SASU (IS)"]:
        statut_social_dirigeant = "Assimil√© Salari√©"
    else:
        statut_social_dirigeant = "Autre"
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    st.write(f"**Statut juridique :** {statut_juridique}")
    st.write(f"**B√©n√©fice de l'ACRE :** {benefice_accre}")
    st.write(f"**Statut social du (des) dirigeant(s) :** {statut_social_dirigeant}")
    
    st.write("---")
    
    # R√©cup√©rer les donn√©es de salaires
    salaires = data.get("salaires", {})
    salaires_dirigeant = salaires.get("dirigeants", {})
    salaires_employes = salaires.get("employes", {})
    
    # D√©finir les taux de charges sociales en fonction du statut juridique et de l'ACCRE
    taux_charges_dirigeant = {
        # Sans ACCRE
        "Sans ACCRE": {
            "Travailleur Non Salari√© (TNS)": 0.45,
            "Assimil√© Salari√©": 0.80,  # Taux approximatif pour les assimil√©s salari√©s
        },
        # Avec ACCRE
        "Avec ACCRE": {
            "Travailleur Non Salari√© (TNS)": 0.22,
            "Assimil√© Salari√©": 0.50,  # Taux r√©duit pour les assimil√©s salari√©s avec ACRE
        }
    }
    
    # S√©lection du taux appropri√© pour le dirigeant
    if benefice_accre.lower() == "oui":
        taux_dirigeant = taux_charges_dirigeant["Avec ACCRE"].get(statut_social_dirigeant, 0.45)
    else:
        taux_dirigeant = taux_charges_dirigeant["Sans ACCRE"].get(statut_social_dirigeant, 0.45)
    
    # Taux de charges sociales pour les employ√©s
    taux_charges_employe = 0.72  # Comme indiqu√©, multiplier par 0.72 qu'il ait ACCRE ou pas
    
    # Pr√©paration des donn√©es pour le tableau
    annees = ["Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    remuneration_dirigeant = []
    augmentation_dirigeant = []
    charges_sociales_dirigeant = []
    remuneration_employes = []
    augmentation_employes = []
    charges_sociales_employes = []
    
    for i, annee in enumerate(annees):
        annee_key = f"annee{i+1}"
        # R√©mun√©ration du (des) dirigeants
        remu_dirigeant = salaires_dirigeant.get(annee_key, 0.0)
        remuneration_dirigeant.append(remu_dirigeant)
        # % augmentation dirigeant
        if i == 0:
            aug_dirigeant = "-"
        else:
            previous_remu_dirigeant = remuneration_dirigeant[i-1]
            if previous_remu_dirigeant != 0:
                aug_dirigeant_value = ((remu_dirigeant - previous_remu_dirigeant) / previous_remu_dirigeant) * 100
                aug_dirigeant = f"{aug_dirigeant_value:.2f}%"
            else:
                aug_dirigeant = "-"
        augmentation_dirigeant.append(aug_dirigeant)
        # Charges sociales du (des) dirigeant(s)
        charge_sociale_dirigeant = remu_dirigeant * taux_dirigeant
        charges_sociales_dirigeant.append(charge_sociale_dirigeant)
        
        # Salaires des employ√©s
        remu_employes = salaires_employes.get(annee_key, 0.0)
        remuneration_employes.append(remu_employes)
        # % augmentation employ√©s
        if i == 0:
            aug_employes = "-"
        else:
            previous_remu_employes = remuneration_employes[i-1]
            if previous_remu_employes != 0:
                aug_employes_value = ((remu_employes - previous_remu_employes) / previous_remu_employes) * 100
                aug_employes = f"{aug_employes_value:.2f}%"
            else:
                aug_employes = "-"
        augmentation_employes.append(aug_employes)
        # Charges sociales employ√©s
        charge_sociale_employes = remu_employes * taux_charges_employe
        charges_sociales_employes.append(charge_sociale_employes)
    
    # Cr√©ation du DataFrame pour l'affichage
    df = pd.DataFrame({
        "": ["R√©mun√©ration du (des) dirigeants", "% augmentation", "Charges sociales du (des) dirigeant(s)",
             "Salaires des employ√©s", "% augmentation", "Charges sociales employ√©s"],
        "Ann√©e 1": [f"{remuneration_dirigeant[0]:.2f} $", augmentation_dirigeant[0], f"{charges_sociales_dirigeant[0]:.2f} $",
                    f"{remuneration_employes[0]:.2f} $", augmentation_employes[0], f"{charges_sociales_employes[0]:.2f} $"],
        "Ann√©e 2": [f"{remuneration_dirigeant[1]:.2f} $", augmentation_dirigeant[1], f"{charges_sociales_dirigeant[1]:.2f} $",
                    f"{remuneration_employes[1]:.2f} $", augmentation_employes[1], f"{charges_sociales_employes[1]:.2f} $"],
        "Ann√©e 3": [f"{remuneration_dirigeant[2]:.2f} $", augmentation_dirigeant[2], f"{charges_sociales_dirigeant[2]:.2f} $",
                    f"{remuneration_employes[2]:.2f} $", augmentation_employes[2], f"{charges_sociales_employes[2]:.2f} $"]
    })
    
    st.table(df)
    
    # Stocker les charges sociales dans les donn√©es pour exportation
    data["charges_sociales"] = {
        "dirigeants": {
            "annee1": charges_sociales_dirigeant[0],
            "annee2": charges_sociales_dirigeant[1],
            "annee3": charges_sociales_dirigeant[2]
        },
        "employes": {
            "annee1": charges_sociales_employes[0],
            "annee2": charges_sociales_employes[1],
            "annee3": charges_sociales_employes[2]
        }
    }
    
    st.session_state["data"] = data
    
    # Pr√©parer les donn√©es d'exportation pour Salaires et Charges Sociales
    export_table_data = []
    
    # Ajouter les lignes du tableau
    for index, row in df.iterrows():
        export_table_data.append({
            "Description": row[""],
            "Ann√©e 1": row["Ann√©e 1"],
            "Ann√©e 2": row["Ann√©e 2"],
            "Ann√©e 3": row["Ann√©e 3"]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_salaires_charges_sociales'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "statut_juridique": statut_juridique,
        "benefice_accre": benefice_accre,
        "statut_social_dirigeant": statut_social_dirigeant,
        "table_data": export_table_data
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le Markdown
    if st.button("T√©l√©charger Salaires en Markdown"):
        export_data = st.session_state.get('export_data_salaires_charges_sociales', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donn√©e disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Salaires et Charges Sociales\n\n**Projet :** {export_data['projet']}\n\n"
            markdown_content += f"**Porteur de projet :** {export_data['porteur_projet']}\n\n"
            markdown_content += f"**Statut juridique :** {export_data['statut_juridique']}\n\n"
            markdown_content += f"**B√©n√©fice de l'ACRE :** {export_data['benefice_accre']}\n\n"
            markdown_content += f"**Statut social du (des) dirigeant(s) :** {export_data['statut_social_dirigeant']}\n\n"
            markdown_content += "---\n\n"
            
            # Cr√©er un DataFrame pour Markdown
            df_markdown = pd.DataFrame(export_data['table_data'])
            markdown_content += df_markdown.to_markdown(index=False)
            
            markdown_content += f"\n\n---\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="T√©l√©charger le Markdown",
                data=markdown_bytes,
                file_name="salaires_charges_sociales.md",
                mime="text/markdown"
            )
    
    # Bouton pour t√©l√©charger le fichier Word
    if st.button("T√©l√©charger Salaires en Word"):
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        
        if not export_data_salaires or "table_data" not in export_data_salaires:
            st.error("Aucune donn√©e disponible pour l'exportation des Salaires et Charges Sociales.")
            return
        
        if not export_data_investissements or "table_data" not in export_data_investissements:
            st.error("Aucune donn√©e disponible pour l'exportation des Investissements et Financements.")
            return
        
        doc = Document()
        
        # Ajouter la premi√®re table : Investissements et Financements
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements['porteur_projet']}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements['equilibre']['message']}")
        doc.add_page_break()
        
        # Cr√©er le tableau Investissements et Financements dans Word
        table_word_inv = doc.add_table(rows=1, cols=4)
        table_word_inv.style = 'Light List Accent 1'
        table_word_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_word_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            # Mettre en gras les en-t√™tes
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les donn√©es Investissements et Financements au tableau
        for row in export_data_investissements['table_data']:
            row_cells_inv = table_word_inv.add_row().cells
            row_cells_inv[0].text = row["Investissements"]
            row_cells_inv[1].text = row["Taux (%)"]
            row_cells_inv[2].text = str(row["Dur√©e (mois)"]) if row["Dur√©e (mois)"] != "-" else "-"
            row_cells_inv[3].text = row["Montant ($)"]
            
            # Mise en forme des lignes sp√©cifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                pass  # Aucune mise en forme suppl√©mentaire
            
            # Alignement des cellules
            row_cells_inv[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la deuxi√®me table : Salaires et Charges Sociales
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires['porteur_projet']}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires['statut_juridique']}")
        doc.add_paragraph(f"**B√©n√©fice de l'ACRE :** {export_data_salaires['benefice_accre']}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires['statut_social_dirigeant']}")
        doc.add_paragraph("---")
        
        # Cr√©er le tableau Salaires et Charges Sociales dans Word
        table_word_sal = doc.add_table(rows=1, cols=4)
        table_word_sal.style = 'Light List Accent 1'
        table_word_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_word_sal.rows[0].cells
        headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            # Mettre en gras les en-t√™tes
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les donn√©es Salaires et Charges Sociales au tableau
        for row in export_data_salaires['table_data']:
            row_cells_sal = table_word_sal.add_row().cells
            row_cells_sal[0].text = row["Description"]
            row_cells_sal[1].text = row["Ann√©e 1"]
            row_cells_sal[2].text = row["Ann√©e 2"]
            row_cells_sal[3].text = row["Ann√©e 3"]
            
            # Mise en forme des lignes sp√©cifiques
            # Vous pouvez ajouter des conditions ici pour mettre en forme certaines lignes si n√©cessaire
            
            # Alignement des cellules
            row_cells_sal[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter des informations suppl√©mentaires si n√©cessaire
        doc.add_paragraph()
        doc.add_paragraph("Les charges sociales sont calcul√©es en fonction des taux applicables.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="T√©l√©charger le fichier Word Complet",
            data=buffer,
            file_name="investissements_et_salaires_charges_sociales.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def page_detail_amortissements():
    st.title("D√©tail des Amortissements")
    
    # Initialiser la cl√© 'export_data_detail_amortissements' dans session_state si elle n'existe pas
    if 'export_data_detail_amortissements' not in st.session_state:
        st.session_state['export_data_detail_amortissements'] = {}
    
    data = st.session_state.get("data", {})
    
    st.write("---")
    
    # R√©cup√©rer la dur√©e d'amortissement
    duree_amortissement = data.get("duree_amortissement", 3)  # Par d√©faut √† 3 ans si non d√©fini
    if duree_amortissement <= 0:
        st.warning("La dur√©e d'amortissement doit √™tre sup√©rieure √† z√©ro.")
        return
    
    # R√©cup√©rer les montants des investissements
    besoins_demarrage = data.get("besoins_demarrage", {})
    
    # Fonction pour calculer les amortissements
    def calcul_amortissements(items):
        amortissements = {}
        total_amort = [0.0, 0.0, 0.0]
        for item in items:
            amount = besoins_demarrage.get(item, 0.0)
            annual_depreciation = amount / duree_amortissement if duree_amortissement > 0 else 0.0
            amortization_years = [0.0, 0.0, 0.0]
            for year in range(3):
                if year < duree_amortissement:
                    amortization_years[year] = annual_depreciation
                    total_amort[year] += annual_depreciation
            amortissements[item] = amortization_years
        return amortissements, total_amort
    
    # Incorporels
    incorporels_items = [
        "Frais d‚Äô√©tablissement",
        "Logiciels, formations",
        "Droits d‚Äôentr√©e",
        "Frais de dossier",
        "Frais de notaire"
    ]
    incorporels_amortissements, total_incorporels_amort = calcul_amortissements(incorporels_items)
    
    # Corporels
    corporels_items = [
        "Enseigne et √©l√©ments de communication",
        "V√©hicule",
        "Mat√©riel professionnel",
        "Mat√©riel autre",
        "Mat√©riel de bureau"
    ]
    corporels_amortissements, total_corporels_amort = calcul_amortissements(corporels_items)
    
    # Total amortissements par ann√©e
    total_amortissements = [
        total_incorporels_amort[year] + total_corporels_amort[year] for year in range(3)
    ]
    
    # Cr√©ation d'un tableau unique
    st.subheader("Amortissements")
    amortissements_data = []
    
    # Ajout des totaux des cat√©gories
    amortissements_data.append({
        "Amortissement": "Amortissements incorporels",
        "Ann√©e 1": f"{total_incorporels_amort[0]:.2f}",
        "Ann√©e 2": f"{total_incorporels_amort[1]:.2f}",
        "Ann√©e 3": f"{total_incorporels_amort[2]:.2f}"
    })
    
    # Ajout d'une ligne vide pour la lisibilit√©
    amortissements_data.append({
        "Amortissement": "",
        "Ann√©e 1": "",
        "Ann√©e 2": "",
        "Ann√©e 3": ""
    })
    
    # Ajout des d√©tails des amortissements incorporels
    for item in incorporels_items:
        amortization_years = incorporels_amortissements.get(item, [0.0, 0.0, 0.0])
        amortissements_data.append({
            "Amortissement": item,
            "Ann√©e 1": f"{amortization_years[0]:.2f}",
            "Ann√©e 2": f"{amortization_years[1]:.2f}",
            "Ann√©e 3": f"{amortization_years[2]:.2f}"
        })
    
    # Ajout d'une ligne vide pour la lisibilit√©
    amortissements_data.append({
        "Amortissement": "",
        "Ann√©e 1": "",
        "Ann√©e 2": "",
        "Ann√©e 3": ""
    })
    amortissements_data.append({
        "Amortissement": "Amortissements corporels",
        "Ann√©e 1": f"{total_corporels_amort[0]:.2f}",
        "Ann√©e 2": f"{total_corporels_amort[1]:.2f}",
        "Ann√©e 3": f"{total_corporels_amort[2]:.2f}"
    })
        # Ajout d'une ligne vide pour la lisibilit√©
    amortissements_data.append({
        "Amortissement": "",
        "Ann√©e 1": "",
        "Ann√©e 2": "",
        "Ann√©e 3": ""
    })
        
    # Ajout des d√©tails des amortissements corporels
    for item in corporels_items:
        amortization_years = corporels_amortissements.get(item, [0.0, 0.0, 0.0])
        amortissements_data.append({
            "Amortissement": item,
            "Ann√©e 1": f"{amortization_years[0]:.2f}",
            "Ann√©e 2": f"{amortization_years[1]:.2f}",
            "Ann√©e 3": f"{amortization_years[2]:.2f}"
        })
    
    # Ajout d'une ligne vide pour la lisibilit√©
    amortissements_data.append({
        "Amortissement": "",
        "Ann√©e 1": "",
        "Ann√©e 2": "",
        "Ann√©e 3": ""
    })
    
    # Total amortissements
    amortissements_data.append({
        "Amortissement": "Total Amortissements",
        "Ann√©e 1": f"{total_amortissements[0]:.2f}",
        "Ann√©e 2": f"{total_amortissements[1]:.2f}",
        "Ann√©e 3": f"{total_amortissements[2]:.2f}"
    })
    
    # Cr√©ation du DataFrame
    df_amortissements = pd.DataFrame(amortissements_data)
    
    # Affichage du tableau avec des bordures pour plus de clart√©
    st.table(df_amortissements.style.set_properties(**{
        'text-align': 'right'
    }).set_table_styles([
        {'selector': 'th', 'props': [('text-align', 'center')]}
    ]))

    
    # Stocker les amortissements dans les donn√©es pour exportation
    data["amortissements"] = {
        "incorporels": {
            "annee1": total_incorporels_amort[0],
            "annee2": total_incorporels_amort[1],
            "annee3": total_incorporels_amort[2]
        },
        "corporels": {
            "annee1": total_corporels_amort[0],
            "annee2": total_corporels_amort[1],
            "annee3": total_corporels_amort[2]
        },
        "total": {
            "annee1": total_amortissements[0],
            "annee2": total_amortissements[1],
            "annee3": total_amortissements[2]
        }
    }
    
    st.session_state["data"] = data
    
    # Pr√©parer les donn√©es d'exportation pour D√©tail des Amortissements
    export_table_amortissements = []
    for row in amortissements_data:
        export_table_amortissements.append({
            "Amortissement": row["Amortissement"],
            "Ann√©e 1": row["Ann√©e 1"],
            "Ann√©e 2": row["Ann√©e 2"],
            "Ann√©e 3": row["Ann√©e 3"]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_detail_amortissements'] = {
        "amortissements": export_table_amortissements
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le Markdown
    if st.button("T√©l√©charger Amortissements en Markdown"):
        export_data = st.session_state.get('export_data_detail_amortissements', {})
        if not export_data:
            st.error("Aucune donn√©e disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# D√©tail des Amortissements\n\n"
            markdown_content += "---\n\n"
            
            # Amortissements
            markdown_content += "## Amortissements\n\n"
            df_amortissements_md = pd.DataFrame(export_data['amortissements'])
            markdown_content += df_amortissements_md.to_markdown(index=False)
            markdown_content += "\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="T√©l√©charger le Markdown",
                data=markdown_bytes,
                file_name="detail_amortissements.md",
                mime="text/markdown"
            )
    
    # Bouton pour t√©l√©charger le fichier Word
    if st.button("T√©l√©charger Amortissements en Word"):
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        
        if not export_data_amortissements or "amortissements" not in export_data_amortissements:
            st.error("Aucune donn√©e disponible pour l'exportation des Amortissements.")
            return
        
        if not export_data_investissements or "table_data" not in export_data_investissements:
            st.error("Aucune donn√©e disponible pour l'exportation des Investissements et Financements.")
            return
        
        if not export_data_salaires or "table_data" not in export_data_salaires:
            st.error("Aucune donn√©e disponible pour l'exportation des Salaires et Charges Sociales.")
            return
        
        doc = Document()
        
        # Ajouter la premi√®re table : Investissements et Financements
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements['porteur_projet']}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements['equilibre']['message']}")
        doc.add_page_break()
        
        # Cr√©er le tableau Investissements et Financements dans Word
        table_word_inv = doc.add_table(rows=1, cols=4)
        table_word_inv.style = 'Light List Accent 1'
        table_word_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_word_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            # Mettre en gras les en-t√™tes
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les donn√©es Investissements et Financements au tableau
        for row in export_data_investissements['table_data']:
            row_cells_inv = table_word_inv.add_row().cells
            row_cells_inv[0].text = row["Investissements"]
            row_cells_inv[1].text = row["Taux (%)"]
            row_cells_inv[2].text = str(row["Dur√©e (mois)"]) if row["Dur√©e (mois)"] != "-" else "-"
            row_cells_inv[3].text = row["Montant ($)"]
            
            # Mise en forme des lignes sp√©cifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                pass  # Aucune mise en forme suppl√©mentaire
            
            # Alignement des cellules
            row_cells_inv[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la deuxi√®me table : Salaires et Charges Sociales
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires['porteur_projet']}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires['statut_juridique']}")
        doc.add_paragraph(f"**B√©n√©fice de l'ACRE :** {export_data_salaires['benefice_accre']}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires['statut_social_dirigeant']}")
        doc.add_paragraph("---")
        
        # Cr√©er le tableau Salaires et Charges Sociales dans Word
        table_word_sal = doc.add_table(rows=1, cols=4)
        table_word_sal.style = 'Light List Accent 1'
        table_word_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_word_sal.rows[0].cells
        headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            # Mettre en gras les en-t√™tes
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les donn√©es Salaires et Charges Sociales au tableau
        for row in export_data_salaires['table_data']:
            row_cells_sal = table_word_sal.add_row().cells
            row_cells_sal[0].text = row["Description"]
            row_cells_sal[1].text = row["Ann√©e 1"]
            row_cells_sal[2].text = row["Ann√©e 2"]
            row_cells_sal[3].text = row["Ann√©e 3"]
            
            # Mise en forme des lignes sp√©cifiques
            # Vous pouvez ajouter des conditions ici pour mettre en forme certaines lignes si n√©cessaire
            
            # Alignement des cellules
            row_cells_sal[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la troisi√®me table : D√©tail des Amortissements
        doc.add_heading('D√©tail des Amortissements', level=1)
        
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        
        # V√©rifier si les donn√©es d'amortissements sont disponibles
        if not export_data_amortissements or "amortissements" not in export_data_amortissements:
            st.error("Aucune donn√©e disponible pour l'exportation des Amortissements.")
            return
        
        # Cr√©er le tableau Amortissements dans Word
        doc.add_heading('Amortissements', level=2)
        table_word_amort = doc.add_table(rows=1, cols=4)
        table_word_amort.style = 'Light List Accent 1'
        table_word_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_amort = table_word_amort.rows[0].cells
        headers_amort = ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
        for i, header in enumerate(headers_amort):
            hdr_cells_amort[i].text = header
            # Mettre en gras les en-t√™tes
            for paragraph in hdr_cells_amort[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les donn√©es Amortissements au tableau
        for row in export_data_amortissements['amortissements']:
            row_cells_amort = table_word_amort.add_row().cells
            row_cells_amort[0].text = row["Amortissement"]
            row_cells_amort[1].text = row["Ann√©e 1"]
            row_cells_amort[2].text = row["Ann√©e 2"]
            row_cells_amort[3].text = row["Ann√©e 3"]
        
        # Ajouter des informations suppl√©mentaires si n√©cessaire
        doc.add_paragraph()
        doc.add_paragraph("Les amortissements sont calcul√©s en fonction de la dur√©e d'amortissement sp√©cifi√©e.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="T√©l√©charger le fichier Word Complet",
            data=buffer,
            file_name="document_complet_financier.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def telecharger_document_complet():
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    
    # V√©rifiez que toutes les donn√©es sont pr√©sentes
    if not all([export_data_investissements.get("table_data"),
                export_data_salaires.get("table_data"),
                export_data_amortissements.get("amortissements"),
                export_data_compte.get("table_data")]):
        st.error("Toutes les sections doivent √™tre remplies avant de t√©l√©charger le document complet.")
        return
    
    # Cr√©er un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_paragraph(f"**Equilibre :** {export_data_investissements.get('equilibre', {}).get('message', '')}")
    doc.add_page_break()
    
    # Cr√©er le tableau Investissements et Financements
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Dur√©e (mois)", "")) if row.get("Dur√©e (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes sp√©cifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph(f"**Statut juridique :** {export_data_salaires.get('statut_juridique', 'N/A')}")
    doc.add_paragraph(f"**B√©n√©fice de l'ACRE :** {export_data_salaires.get('benefice_accre', 'N/A')}")
    doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires.get('statut_social_dirigeant', 'N/A')}")
    doc.add_paragraph("---")
    
    # Cr√©er le tableau Salaires et Charges Sociales
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Ann√©e 1", "")
        row_cells[2].text = row.get("Ann√©e 2", "")
        row_cells[3].text = row.get("Ann√©e 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section D√©tail des Amortissements ###
    doc.add_heading('D√©tail des Amortissements', level=1)
    
    # Cr√©er le tableau D√©tail des Amortissements
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_amortissements['amortissements']:
        row_cells = table_amort.add_row().cells
        row_cells[0].text = row.get("Amortissement", "")
        row_cells[1].text = row.get("Ann√©e 1", "")
        row_cells[2].text = row.get("Ann√©e 2", "")
        row_cells[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calcul√©s en fonction de la dur√©e d'amortissement sp√©cifi√©e.")
    
    ### 4. Ajouter la section Compte de R√©sultats Pr√©visionnel ###
    doc.add_heading('Compte de R√©sultats Pr√©visionnel', level=1)
    
    # Cr√©er le tableau Compte de R√©sultats Pr√©visionnel
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells = table_compte.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Ann√©e 1", "")
        row_cells[2].text = row.get("Ann√©e 2", "")
        row_cells[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de t√©l√©chargement
    st.download_button(
        label="T√©l√©charger le fichier Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



def calculer_impot_societes(resultat_avant_impots):
    """
    Calcule l'Imp√¥t sur les Soci√©t√©s (IS) selon la formule progressive.

    Args:
        resultat_avant_impots (float): R√©sultat avant imp√¥ts.

    Returns:
        float: Montant de l'IS.
    """
    if resultat_avant_impots < 0:
        return 0.0
    elif resultat_avant_impots > 38120:
        return 38120 * 0.15 + (resultat_avant_impots - 38120) * 0.28
    else:
        return resultat_avant_impots * 0.15

def page_compte_resultats_previsionnel():
    st.title("Compte de r√©sultats pr√©visionnel sur 3 ans")
    
    # Initialiser la cl√© 'export_data_compte_resultats_previsionnel' dans session_state si elle n'existe pas
    if 'export_data_compte_resultats_previsionnel' not in st.session_state:
        st.session_state['export_data_compte_resultats_previsionnel'] = {}
    
    data = st.session_state.get("data", {})
    
    # R√©cup√©rer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "")
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Pr√©paration des donn√©es
    # Chiffre d'affaires
    ca_marchandises = [
        data["chiffre_affaires"].get("total_ca_Marchandises_annee1", 0.0),
        data["chiffre_affaires"].get("total_ca_Marchandises_annee2", 0.0),
        data["chiffre_affaires"].get("total_ca_Marchandises_annee3", 0.0)
    ]
    ca_services = [
        data["chiffre_affaires"].get("total_ca_Services_annee1", 0.0),
        data["chiffre_affaires"].get("total_ca_Services_annee2", 0.0),
        data["chiffre_affaires"].get("total_ca_Services_annee3", 0.0)
    ]
    total_ca = [
        ca_marchandises[0] + ca_services[0],
        ca_marchandises[1] + ca_services[1],
        ca_marchandises[2] + ca_services[2]
    ]
    
    # Achats consomm√©s (charges variables) - Suppos√©s nuls si pas de marchandises vendues
    data["charges_variables"] = data.get("charges_variables", {})
    charges_variables = data["charges_variables"]
    cout_achat_marchandises_pct=charges_variables.get("cout_achat_marchandises_pct", 0.0)
    
    charges_variables = [ca_marchandises[0]* cout_achat_marchandises_pct / 100.0,
                         ca_marchandises[1]* cout_achat_marchandises_pct / 100.0,
                         ca_marchandises[2]* cout_achat_marchandises_pct / 100.0
                         ]
    
    # charges exploitations (charges exploitations) - Suppos√©s nuls si pas de marchandises vendues
    
    charges_exploitations = charges_variables 

    
    # Marge brute = Total CA - Achats consomm√©s
    marge_brute = [
        total_ca[0] - charges_variables[0],
        total_ca[1] - charges_variables[1],
        total_ca[2] - charges_variables[2]
    ]
    
    # Charges externes (charges fixes)
    charges_fixes_data = data.get("charges_fixes", {})
    charges_fixes_annee1 = charges_fixes_data.get("annee1", {})
    charges_fixes_annee2 = charges_fixes_data.get("annee2", {})
    charges_fixes_annee3 = charges_fixes_data.get("annee3", {})
    
    # Liste des charges externes d√©taill√©es
    liste_charges = [
        "Assurances v√©hicule et RC pro", "T√©l√©phone, internet", "Autres abonnements",
        "Carburant", "Frais de d√©placement / h√©bergement", "Eau, √©lectricit√©, gaz",
        "Mutuelle", "Fournitures diverses", "Entretien Moto livraison et mat√©riel",
        "Nettoyage des locaux", "Budget publicit√© et communication", "Emplacements",
        "Expert comptable, avocats", "Markting"
    ]
    
    # R√©cup√©ration des montants pour chaque charge
    charges_detaillees = {}
    total_charges_fixes = [0.0, 0.0, 0.0]
    for charge in liste_charges:
        montant_annee1 = charges_fixes_annee1.get(charge, 0.0)
        montant_annee2 = charges_fixes_annee2.get(charge, 0.0)
        montant_annee3 = charges_fixes_annee3.get(charge, 0.0)
        charges_detaillees[charge] = [montant_annee1, montant_annee2, montant_annee3]
        total_charges_fixes[0] += montant_annee1
        total_charges_fixes[1] += montant_annee2
        total_charges_fixes[2] += montant_annee3
    
    # Valeur ajout√©e = Marge brute - Charges externes
    valeur_ajoutee = [
        marge_brute[0] - total_charges_fixes[0],
        marge_brute[1] - total_charges_fixes[1],
        marge_brute[2] - total_charges_fixes[2]
    ]
    
    # Imp√¥ts et taxes (ajouter d'autres imp√¥ts si n√©cessaire)
    impots_et_taxes = [
        charges_fixes_annee1.get("Taxes, CFE", 0.0),
        charges_fixes_annee2.get("Taxes, CFE", 0.0),
        charges_fixes_annee3.get("Taxes, CFE", 0.0)
    ]
    
    # Salaires employ√©s
    salaires_employes = [
        data["salaires"]["employes"].get("annee1", 0.0),
        data["salaires"]["employes"].get("annee2", 0.0),
        data["salaires"]["employes"].get("annee3", 0.0)
    ]
    
    # Charges sociales employ√©s
    charges_sociales_employes = [
        data["charges_sociales"]["employes"].get("annee1", 0.0),
        data["charges_sociales"]["employes"].get("annee2", 0.0),
        data["charges_sociales"]["employes"].get("annee3", 0.0)
    ]
    
    # Pr√©l√®vement dirigeant(s)
    salaires_dirigeants = [
        data["salaires"]["dirigeants"].get("annee1", 0.0),
        data["salaires"]["dirigeants"].get("annee2", 0.0),
        data["salaires"]["dirigeants"].get("annee3", 0.0)
    ]
    
    # Charges sociales dirigeant(s)
    charges_sociales_dirigeants = [
        data["charges_sociales"]["dirigeants"].get("annee1", 0.0),
        data["charges_sociales"]["dirigeants"].get("annee2", 0.0),
        data["charges_sociales"]["dirigeants"].get("annee3", 0.0)
    ]
    
    # Exc√©dent brut d'exploitation = Valeur ajout√©e - Imp√¥ts et taxes - Salaires - Charges sociales
    ebe = [
        valeur_ajoutee[0] - impots_et_taxes[0] - salaires_employes[0] - charges_sociales_employes[0] - salaires_dirigeants[0] - charges_sociales_dirigeants[0],
        valeur_ajoutee[1] - impots_et_taxes[1] - salaires_employes[1] - charges_sociales_employes[1] - salaires_dirigeants[1] - charges_sociales_dirigeants[1],
        valeur_ajoutee[2] - impots_et_taxes[2] - salaires_employes[2] - charges_sociales_employes[2] - salaires_dirigeants[2] - charges_sociales_dirigeants[2]
    ]
    
    # Frais bancaires, charges financi√®res
    frais_bancaires = [
        charges_fixes_annee1.get("Frais bancaires et terminal carte bleue", 0.0),
        charges_fixes_annee2.get("Frais bancaires et terminal carte bleue", 0.0),
        charges_fixes_annee3.get("Frais bancaires et terminal carte bleue", 0.0)
    ]
    
    # Int√©r√™ts des pr√™ts
    interets_prets = data.get("interets_prets", {
        "annee1": 0.0,
        "annee2": 0.0,
        "annee3": 0.0
    })
    
    # Ajouter les int√©r√™ts des pr√™ts aux autres frais financiers
    frais_financiers = [
        interets_prets.get("annee1", 0.0),
        interets_prets.get("annee2", 0.0),
        interets_prets.get("annee3", 0.0)
    ]
    
    # Total des frais bancaires et charges financi√®res
    total_frais_financiers = [
        frais_bancaires[0] + frais_financiers[0],
        frais_bancaires[1] + frais_financiers[1],
        frais_bancaires[2] + frais_financiers[2]
    ]
    
    # Dotations aux amortissements (suppos√©es nulles si non fournies)
    amortissements = [0.0, 0.0, 0.0]
    
    # R√©sultat avant imp√¥ts = EBE - Frais bancaires - Amortissements
    resultat_avant_impots = [
        ebe[0] - total_frais_financiers[0] - amortissements[0],
        ebe[1] - total_frais_financiers[1] - amortissements[1],
        ebe[2] - total_frais_financiers[2] - amortissements[2]
    ]
    
    # Imp√¥t sur les soci√©t√©s (selon la formule progressive)
    impot_societes = [
        calculer_impot_societes(resultat_avant_impots[0]),
        calculer_impot_societes(resultat_avant_impots[1]),
        calculer_impot_societes(resultat_avant_impots[2])
    ]
    
    # R√©sultat net comptable (r√©sultat de l'exercice)
    resultat_net = [
        resultat_avant_impots[0] - impot_societes[0],
        resultat_avant_impots[1] - impot_societes[1],
        resultat_avant_impots[2] - impot_societes[2]
    ]
    
    # Pr√©paration des donn√©es pour le tableau
    tableau = {
        "": [
            "Produits d'exploitation",
            "Chiffre d'affaires HT vente de marchandises",
            "Chiffre d'affaires HT services",
            "",
            "Charges d'exploitation(charge variable)",
            "Achats consomm√©s",
            "",
            "Marge brute",
            "Charges externes",
            ""
        ],
        "Ann√©e 1": [
            f"{total_ca[0]:,.2f} $",
            f"{ca_marchandises[0]:,.2f} $",
            f"{ca_services[0]:,.2f} $",
            "",
            f"{charges_exploitations[0]:,.2f} $",
            f"{charges_variables[0]:,.2f} $",
            "",
            f"{marge_brute[0]:,.2f} $",
            "",
            ""
        ],
        "Ann√©e 2": [
            f"{total_ca[1]:,.2f} $",
            f"{ca_marchandises[1]:,.2f} $",
            f"{ca_services[1]:,.2f} $",
            "",
            f"{charges_exploitations[1]:,.2f} $",
            f"{charges_variables[1]:,.2f} $",
            "",
            f"{marge_brute[1]:,.2f} $",
            "",
            ""
        ],
        "Ann√©e 3": [
            f"{total_ca[2]:,.2f} $",
            f"{ca_marchandises[2]:,.2f} $",
            f"{ca_services[2]:,.2f} $",
            "",
            f"{charges_exploitations[2]:,.2f} $",
            f"{charges_variables[2]:,.2f} $",
            "",
            f"{marge_brute[2]:,.2f} $",
            "",
            ""
        ]
    }
    
    # Ajouter les charges d√©taill√©es au tableau
    for charge in liste_charges:
        tableau[""].append(charge)
        tableau["Ann√©e 1"].append(f"{charges_detaillees[charge][0]:,.2f} $")
        tableau["Ann√©e 2"].append(f"{charges_detaillees[charge][1]:,.2f} $")
        tableau["Ann√©e 3"].append(f"{charges_detaillees[charge][2]:,.2f} $")
    
    # Ajouter le total des charges externes
    tableau[""].append("Total Charges externes")
    tableau["Ann√©e 1"].append(f"{total_charges_fixes[0]:,.2f} $")
    tableau["Ann√©e 2"].append(f"{total_charges_fixes[1]:,.2f} $")
    tableau["Ann√©e 3"].append(f"{total_charges_fixes[2]:,.2f} $")
    
    # Continuer √† remplir le tableau
    additional_rows = {
        "Valeur ajout√©e": valeur_ajoutee,
        "Imp√¥ts et taxes": impots_et_taxes,
        "Salaires employ√©s": salaires_employes,
        "Charges sociales employ√©s": charges_sociales_employes,
        "Pr√©l√®vement dirigeant(s)": salaires_dirigeants,
        "Charges sociales dirigeant(s)": charges_sociales_dirigeants,
        "Exc√©dent brut d'exploitation": ebe,
        "Frais bancaires, charges financi√®res": total_frais_financiers,
        "Dotations aux amortissements": amortissements,
        "R√©sultat avant imp√¥ts": resultat_avant_impots,
        "Imp√¥t sur les soci√©t√©s": impot_societes,
        "R√©sultat net comptable (r√©sultat de l'exercice)": resultat_net
    }
    
    for key, values in additional_rows.items():
        tableau[""].append(key)
        tableau["Ann√©e 1"].append(f"{values[0]:,.2f} $")
        tableau["Ann√©e 2"].append(f"{values[1]:,.2f} $")
        tableau["Ann√©e 3"].append(f"{values[2]:,.2f} $")
    
    # Cr√©er le DataFrame
    df_resultats = pd.DataFrame(tableau)
    
    # Afficher le tableau
    st.table(df_resultats)
    
    # Ajouter les variables calcul√©es au dictionnaire 'data'
    data["compte_de_resultat"] = {
        "total_ca": total_ca,
        "ca_marchandises": ca_marchandises,
        "ca_services": ca_services,
        "charges_exploitations":charges_exploitations,
        "charges_variables": charges_variables,
        "marge_brute": marge_brute,
        "charges_fixes": total_charges_fixes,
        "valeur_ajoutee": valeur_ajoutee,
        "impots_et_taxes": impots_et_taxes,
        "salaires_employes": salaires_employes,
        "charges_sociales_employes": charges_sociales_employes,
        "salaires_dirigeants": salaires_dirigeants,
        "charges_sociales_dirigeants": charges_sociales_dirigeants,
        "ebe": ebe,
        "frais_bancaires": frais_bancaires,
        "frais_financiers": frais_financiers,
        "total_frais_financiers": total_frais_financiers,
        "amortissements": amortissements,
        "resultat_avant_impots": resultat_avant_impots,
        "impot_societes": impot_societes,
        "resultat_net": resultat_net
    }
    
    # Enregistrer les donn√©es dans la session
    st.session_state["data"] = data
    
    # Pr√©parer les donn√©es d'exportation pour Compte de r√©sultats pr√©visionnel
    export_table_compte = []
    for index, row in df_resultats.iterrows():
        export_table_compte.append({
            "Description": row[""],
            "Ann√©e 1": row["Ann√©e 1"],
            "Ann√©e 2": row["Ann√©e 2"],
            "Ann√©e 3": row["Ann√©e 3"]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_compte_resultats_previsionnel'] = {
        "table_data": export_table_compte
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    # Bouton pour t√©l√©charger le document complet
    st.button("T√©l√©charger le Document Word Complet", on_click=telecharger_document_complet)
    


import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_soldes_intermediaires_de_gestion():
    st.title("Soldes interm√©diaires de gestion")
    
    # R√©cup√©rer les donn√©es de la session
    data = st.session_state.get("data", {})
    
    # R√©cup√©rer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # R√©cup√©rer les donn√©es n√©cessaires pour les calculs
    compte_resultat = data.get("compte_de_resultat", {})
    
    # Fonction pour assurer que les listes contiennent trois √©l√©ments et convertir en float
    def get_three_years_data(key):
        values = compte_resultat.get(key, [])
        processed_values = []
        for v in values:
            try:
                processed_values.append(float(v))
            except (ValueError, TypeError):
                processed_values.append(0.0)
        # Compl√©ter avec 0.0 si moins de 3 √©l√©ments
        while len(processed_values) < 3:
            processed_values.append(0.0)
        return processed_values[:3]
    
    # R√©cup√©ration des donn√©es avec validation
    total_ca = get_three_years_data("total_ca")
    ca_marchandises = get_three_years_data("ca_marchandises")
    ca_services = get_three_years_data("ca_services")
    achats_consommes = get_three_years_data("charges_variables")  # Actuellement d√©fini √† [0.0, 0.0, 0.0]
    charges_fixes = get_three_years_data("charges_fixes")
    impot_societes = get_three_years_data("impot_societes")
    impots_et_taxes = get_three_years_data("impots_et_taxes")
    salaires_employes = get_three_years_data("salaires_employes")
    charges_sociales_employes = get_three_years_data("charges_sociales_employes")
    salaires_dirigeants = get_three_years_data("salaires_dirigeants")
    charges_sociales_dirigeants = get_three_years_data("charges_sociales_dirigeants")
    amortissements = get_three_years_data("amortissements")
    total_frais_financiers = get_three_years_data("total_frais_financiers")
    
    # Calcul des diff√©rents soldes interm√©diaires
    ventes_production_reelle = [ca_marchandises[i] + ca_services[i] for i in range(3)]
    marge_globale = [ventes_production_reelle[i] - achats_consommes[i] for i in range(3)]
    valeur_ajoutee = [marge_globale[i] - charges_fixes[i] for i in range(3)]
    charges_personnel = [
        salaires_employes[i] + charges_sociales_employes[i] + salaires_dirigeants[i] + charges_sociales_dirigeants[i]
        for i in range(3)
    ]
    ebe = [valeur_ajoutee[i] - impots_et_taxes[i] - charges_personnel[i] for i in range(3)]
    resultat_exploitation = [ebe[i] - amortissements[i] for i in range(3)]
    resultat_financier = [-total_frais_financiers[i] for i in range(3)]
    resultat_courant = [resultat_exploitation[i] + resultat_financier[i] for i in range(3)]
    resultat_exercice = [resultat_courant[i] - impot_societes[i] for i in range(3)]
    capacite_autofinancement = [resultat_exercice[i] + amortissements[i] for i in range(3)]
    
    # Fonction de calcul des pourcentages avec gestion de la division par z√©ro
    def calculate_percentage(value, ca):
        return (value / ca * 100) if ca != 0 else 0.0
    
    # Pr√©paration des donn√©es pour le tableau
    soldes = [
        "Chiffre d'affaires",
        "Ventes + production r√©elle",
        "Achats consomm√©s",
        "Marge globale",
        "Charges externes",
        "Valeur ajout√©e",
        "Imp√¥ts et taxes",
        "Charges de personnel",
        "Exc√©dent brut d'exploitation (EBE)",
        "Dotations aux amortissements",
        "R√©sultat d'exploitation",
        "Charges financi√®res",
        "R√©sultat financier",
        "R√©sultat courant",
        "R√©sultat de l'exercice",
        "Capacit√© d'autofinancement"
    ]
    
    # Initialiser le data_table avec les soldes
    data_table = {"Soldes interm√©diaires de gestion": soldes}
    
    # Ajouter les donn√©es pour chaque ann√©e et leurs pourcentages
    for year in range(3):
        data_table[f"Ann√©e {year+1}"] = [
            total_ca[year],
            ventes_production_reelle[year],
            achats_consommes[year],
            marge_globale[year],
            charges_fixes[year],
            valeur_ajoutee[year],
            impots_et_taxes[year],
            charges_personnel[year],
            ebe[year],
            amortissements[year],
            resultat_exploitation[year],
            total_frais_financiers[year],
            resultat_financier[year],
            resultat_courant[year],
            resultat_exercice[year],
            capacite_autofinancement[year]
        ]
        
        data_table[f"% Ann√©e {year+1}"] = [
            100.0,  # Chiffre d'affaires
            100.0,  # Ventes + production r√©elle
            calculate_percentage(achats_consommes[year], total_ca[year]),
            calculate_percentage(marge_globale[year], total_ca[year]),
            calculate_percentage(charges_fixes[year], total_ca[year]),
            calculate_percentage(valeur_ajoutee[year], total_ca[year]),
            calculate_percentage(impots_et_taxes[year], total_ca[year]),
            calculate_percentage(charges_personnel[year], total_ca[year]),
            calculate_percentage(ebe[year], total_ca[year]),
            calculate_percentage(amortissements[year], total_ca[year]),
            calculate_percentage(resultat_exploitation[year], total_ca[year]),
            calculate_percentage(total_frais_financiers[year], total_ca[year]),
            calculate_percentage(resultat_financier[year], total_ca[year]),
            calculate_percentage(resultat_courant[year], total_ca[year]),
            calculate_percentage(resultat_exercice[year], total_ca[year]),
            calculate_percentage(capacite_autofinancement[year], total_ca[year])
        ]
    
    # Cr√©er le DataFrame avec les donn√©es
    df = pd.DataFrame(data_table)
    
    # D√©finir l'ordre des colonnes altern√©es entre "Ann√©e x" et "%"
    columns_order = ["Soldes interm√©diaires de gestion"]
    for year in range(3):
        columns_order.append(f"Ann√©e {year+1}")
        columns_order.append(f"% Ann√©e {year+1}")
    df = df[columns_order]
    
    # Afficher le tableau avec une mise en forme am√©lior√©e
    st.dataframe(
        df.style.format({
            "Ann√©e 1": "{:,.2f} $",
            "Ann√©e 2": "{:,.2f} $",
            "Ann√©e 3": "{:,.2f} $",
            "% Ann√©e 1": "{:.2f}%",
            "% Ann√©e 2": "{:.2f}%",
            "% Ann√©e 3": "{:.2f}%"
        }).set_properties(**{
            'text-align': 'right'
        }).set_table_styles([{
            'selector': 'th',
            'props': [('text-align', 'center')]
        }])
    )
    
    # Stocker les r√©sultats dans les donn√©es pour exportation
    data["soldes_intermediaires_de_gestion"] = {
        "ca": total_ca,
        "ventes_production_reelle": ventes_production_reelle,
        "achats_consommes": achats_consommes,
        "marge_globale": marge_globale,
        "charges_externes": charges_fixes,
        "valeur_ajoutee": valeur_ajoutee,
        "impots_et_taxes": impots_et_taxes,
        "charges_personnel": charges_personnel,
        "ebe": ebe,
        "dotations_aux_amortissements": amortissements,
        "resultat_exploitation": resultat_exploitation,
        "charges_financieres": total_frais_financiers,
        "resultat_financier": resultat_financier,
        "resultat_courant": resultat_courant,
        "resultat_exercice": resultat_exercice,
        "capacite_autofinancement": capacite_autofinancement
    }
    
    # Enregistrer les donn√©es mises √† jour dans la session
    st.session_state["data"] = data
    
    # Pr√©parer les donn√©es d'exportation pour Soldes Interm√©diaires de Gestion avec % colonnes
    export_table_soldes = []
    for idx, solde in enumerate(soldes):
        export_table_soldes.append({
            "Description": solde,
            "Ann√©e 1": data_table["Ann√©e 1"][idx],
            "% Ann√©e 1": data_table["% Ann√©e 1"][idx],
            "Ann√©e 2": data_table["Ann√©e 2"][idx],
            "% Ann√©e 2": data_table["% Ann√©e 2"][idx],
            "Ann√©e 3": data_table["Ann√©e 3"][idx],
            "% Ann√©e 3": data_table["% Ann√©e 3"][idx]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_soldes_intermediaires_de_gestion'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_soldes
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le Markdown
    if st.button("T√©l√©charger Soldes Interm√©diaires en Markdown"):
        export_data = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donn√©e disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Soldes interm√©diaires de gestion\n\n**Projet :** {export_data['projet']}\n\n**Porteur de projet :** {export_data['porteur_projet']}\n\n"
            markdown_content += "---\n\n"
            
            # Cr√©er un DataFrame pour Markdown
            df_markdown = pd.DataFrame(export_data['table_data'])
            markdown_content += df_markdown.to_markdown(index=False)
            
            markdown_content += "\n\n---\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="T√©l√©charger le Markdown",
                data=markdown_bytes,
                file_name="soldes_intermediaires_gestion.md",
                mime="text/markdown"
            )
    
    # Bouton pour t√©l√©charger le fichier Word
    if st.button("T√©l√©charger Soldes Interm√©diaires en Word"):
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        
        if not export_data_soldes or "table_data" not in export_data_soldes:
            st.error("Aucune donn√©e disponible pour l'exportation des Soldes interm√©diaires de gestion.")
            return
        
        # V√©rifiez que toutes les autres sections sont √©galement export√©es
        if not all([
            export_data_investissements.get("table_data"),
            export_data_salaires.get("table_data"),
            export_data_amortissements.get("amortissements"),
            export_data_compte.get("table_data")
        ]):
            st.error("Toutes les sections doivent √™tre remplies avant de t√©l√©charger le document complet.")
            return
        
        # Cr√©er un document Word
        doc = Document()
        
        ### 1. Ajouter la section Investissements et Financements ###
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements.get('equilibre', {}).get('message', '')}")
        doc.add_page_break()
        
        # Cr√©er le tableau Investissements et Financements
        table_inv = doc.add_table(rows=1, cols=4)
        table_inv.style = 'Light List Accent 1'
        table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_investissements['table_data']:
            row_cells = table_inv.add_row().cells
            row_cells[0].text = row.get("Investissements", "")
            row_cells[1].text = row.get("Taux (%)", "")
            row_cells[2].text = str(row.get("Dur√©e (mois)", "")) if row.get("Dur√©e (mois)", "") != "-" else "-"
            row_cells[3].text = row.get("Montant ($)", "")
            
            # Mise en forme des lignes sp√©cifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Alignement des cellules
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        ### 2. Ajouter la section Salaires et Charges Sociales ###
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires.get('statut_juridique', 'N/A')}")
        doc.add_paragraph(f"**B√©n√©fice de l'ACRE :** {export_data_salaires.get('benefice_accre', 'N/A')}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires.get('statut_social_dirigeant', 'N/A')}")
        doc.add_paragraph("---")
        
        # Cr√©er le tableau Salaires et Charges Sociales
        table_sal = doc.add_table(rows=1, cols=4)
        table_sal.style = 'Light List Accent 1'
        table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_sal.rows[0].cells
        headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_salaires['table_data']:
            row_cells = table_sal.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = row.get("Ann√©e 1", "")
            row_cells[2].text = row.get("Ann√©e 2", "")
            row_cells[3].text = row.get("Ann√©e 3", "")
            
            # Alignement des cellules
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        ### 3. Ajouter la section D√©tail des Amortissements ###
        doc.add_heading('D√©tail des Amortissements', level=1)
        
        # Cr√©er le tableau D√©tail des Amortissements
        table_amort = doc.add_table(rows=1, cols=4)
        table_amort.style = 'Light List Accent 1'
        table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_amort = table_amort.rows[0].cells
        headers_amort = ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
        for i, header in enumerate(headers_amort):
            hdr_cells_amort[i].text = header
            for paragraph in hdr_cells_amort[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_amortissements['amortissements']:
            row_cells = table_amort.add_row().cells
            row_cells[0].text = row.get("Amortissement", "")
            row_cells[1].text = row.get("Ann√©e 1", "")
            row_cells[2].text = row.get("Ann√©e 2", "")
            row_cells[3].text = row.get("Ann√©e 3", "")
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les amortissements sont calcul√©s en fonction de la dur√©e d'amortissement sp√©cifi√©e.")
        
        ### 4. Ajouter la section Compte de R√©sultats Pr√©visionnel ###
        doc.add_heading('Compte de R√©sultats Pr√©visionnel', level=1)
        
        # Cr√©er le tableau Compte de R√©sultats Pr√©visionnel
        table_compte = doc.add_table(rows=1, cols=4)
        table_compte.style = 'Light List Accent 1'
        table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_compte = table_compte.rows[0].cells
        headers_compte = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
        for i, header in enumerate(headers_compte):
            hdr_cells_compte[i].text = header
            for paragraph in hdr_cells_compte[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_compte['table_data']:
            row_cells = table_compte.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = row.get("Ann√©e 1", "")
            row_cells[2].text = row.get("Ann√©e 2", "")
            row_cells[3].text = row.get("Ann√©e 3", "")
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
        
        ### 5. Ajouter la section Soldes Interm√©diaires de Gestion ###
        doc.add_heading('Soldes interm√©diaires de gestion', level=1)
        
        # Cr√©er le tableau Soldes interm√©diaires de gestion avec 7 colonnes
        table_soldes = doc.add_table(rows=1, cols=7)
        table_soldes.style = 'Light List Accent 1'
        table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_soldes = table_soldes.rows[0].cells
        headers_soldes = ["Description", "Ann√©e 1", "% Ann√©e 1", "Ann√©e 2", "% Ann√©e 2", "Ann√©e 3", "% Ann√©e 3"]
        for i, header in enumerate(headers_soldes):
            hdr_cells_soldes[i].text = header
            for paragraph in hdr_cells_soldes[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_soldes['table_data']:
            row_cells = table_soldes.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = f"{row.get('Ann√©e 1', 0.0):,.2f} $"
            row_cells[2].text = f"{row.get('% Ann√©e 1', 0.0):.2f}%"
            row_cells[3].text = f"{row.get('Ann√©e 2', 0.0):,.2f} $"
            row_cells[4].text = f"{row.get('% Ann√©e 2', 0.0):.2f}%"
            row_cells[5].text = f"{row.get('Ann√©e 3', 0.0):,.2f} $"
            row_cells[6].text = f"{row.get('% Ann√©e 3', 0.0):.2f}%"
            
            # Alignement des cellules de pourcentage
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Bouton de t√©l√©chargement
        st.download_button(
            label="T√©l√©charger le Document Word Complet",
            data=buffer,
            file_name="document_complet_financier.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Message de confirmation
        st.success("Le document Word complet a √©t√© g√©n√©r√© avec succ√®s !")



    # Enregistrer les donn√©es mises √† jour dans la session
    st.session_state["data"] = data
    
    
    
    
    

def telecharger_document_complets():
    # R√©cup√©rer les donn√©es export√©es de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    
    # V√©rifiez que toutes les donn√©es sont pr√©sentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data")
    ]):
        st.error("Toutes les sections doivent √™tre remplies avant de t√©l√©charger le document complet.")
        return
    
    # Cr√©er un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Cr√©er le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Dur√©e (mois)", "")) if row.get("Dur√©e (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes sp√©cifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Cr√©er le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Ann√©e 1", "")
        row_cells[2].text = row.get("Ann√©e 2", "")
        row_cells[3].text = row.get("Ann√©e 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section D√©tail des Amortissements ###
    doc.add_heading('D√©tail des Amortissements', level=1)
    
    # Cr√©er le tableau D√©tail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les donn√©es Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Ann√©e 1", "")
        row_cells_amort[2].text = row.get("Ann√©e 2", "")
        row_cells_amort[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calcul√©s en fonction de la dur√©e d'amortissement sp√©cifi√©e.")
    
    ### 4. Ajouter la section Compte de R√©sultats Pr√©visionnel ###
    doc.add_heading('Compte de R√©sultats Pr√©visionnel', level=1)
    
    # Cr√©er le tableau Compte de R√©sultats Pr√©visionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Ann√©e 1", "")
        row_cells_compte[2].text = row.get("Ann√©e 2", "")
        row_cells_compte[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 5. Ajouter la section Soldes Interm√©diaires de Gestion ###
    doc.add_heading('Soldes interm√©diaires de gestion', level=1)
    
    # Cr√©er le tableau Soldes interm√©diaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 6. Ajouter la section Capacit√© d'Autofinancement ###
    doc.add_heading('Capacit√© d\'autofinancement', level=1)
    
    # Cr√©er le tableau Capacit√© d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilit√© √âconomique ###
    doc.add_heading('Seuil de rentabilit√© √©conomique', level=1)
    
    # Cr√©er le tableau Seuil de Rentabilit√© √âconomique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Ann√©e 1', 0):,.0f} $"
        row_cells_seuil[2].text = f"{row.get('Ann√©e 2', 0):,.0f} $"
        row_cells_seuil[3].text = f"{row.get('Ann√©e 3', 0):,.0f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de t√©l√©chargement
    st.download_button(
        label="T√©l√©charger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a √©t√© g√©n√©r√© avec succ√®s !")


def calculate_remboursements_emprunts(financements, years=3):
    """
    Votre fonction de calcul existante...
    """
    remboursements = [0.0 for _ in range(years)]  # Initialiser les remboursements pour chaque ann√©e

    for loan_name, loan_info in financements.items():
        # V√©rifier que loan_info est un dictionnaire et commence par "Pr√™t "
        if isinstance(loan_info, dict) and loan_name.startswith("Pr√™t "):
            required_keys = {"montant", "duree", "taux"}
            if not required_keys.issubset(loan_info.keys()):
                st.warning(f"Le pr√™t '{loan_name}' est incomplet et sera ignor√©.")
                continue  # Ignorer les financements incomplets

            montant = loan_info.get("montant", 0.0)
            duree_mois = loan_info.get("duree", 60)  # Par d√©faut 60 mois
            taux_annuel = loan_info.get("taux", 5.0)  # Par d√©faut 5%
            principal_mensuel =  montant / duree_mois if duree_mois > 0 else 0.0

            # Calcul des remboursements par ann√©e bas√©s sur principal_mensuel
            # Principal Year 1
            if duree_mois > 12:
                principal_year1 = principal_mensuel * 12
            else:
                principal_year1 = principal_mensuel * duree_mois

            # Principal Year 2
            if duree_mois - 12 < 0:
                principal_year2 = 0.0
            elif duree_mois > 24:
                principal_year2 = principal_mensuel * 12
            else:
                principal_year2 = principal_mensuel * (duree_mois - 12)

            # Principal Year 3
            if duree_mois - 24 < 0:
                principal_year3 = 0.0
            elif duree_mois > 36:
                principal_year3 = principal_mensuel * 12
            else:
                principal_year3 = principal_mensuel * (duree_mois - 24)

            # Ajouter les remboursements principaux au total par ann√©e
            remboursements[0] += round(principal_year1, 2)
            if years >= 2:
                remboursements[1] += round(principal_year2, 2)
            if years >= 3:
                remboursements[2] += round(principal_year3, 2)
        else:
            # Ignorer les financements qui ne sont pas des pr√™ts (e.g., Apports, Subventions)
            continue

    return remboursements

def page_capacite_autofinancement():
    """
    Affiche le tableau de Capacit√© d'Autofinancement en utilisant les donn√©es de la session.
    """
    st.title("Capacit√© d'autofinancement")
    
    # V√©rifier si les donn√©es sont pr√©sentes dans la session
    if "data" not in st.session_state:
        st.error("Les donn√©es ne sont pas initialis√©es. Veuillez initialiser la session.")
        return
    
    # R√©cup√©rer les donn√©es de la session
    data = st.session_state["data"]
    
    # R√©cup√©rer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # R√©cup√©rer les soldes interm√©diaires de gestion
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    
    # Fonction pour convertir les valeurs en float, remplacer les erreurs par 0.0
    def safe_float_conversion(values):
        return [float(x) if isinstance(x, (int, float)) else 0.0 for x in values]
    
    # R√©cup√©rer et convertir les donn√©es n√©cessaires
    resultat_exercice = safe_float_conversion(soldes_intermediaires.get("resultat_exercice", [0.0, 0.0, 0.0]))
    dotations_aux_amortissements = safe_float_conversion(soldes_intermediaires.get("dotations_aux_amortissements", [0.0, 0.0, 0.0]))
    capacite_autofinancement = safe_float_conversion(soldes_intermediaires.get("capacite_autofinancement", [0.0, 0.0, 0.0]))
    
    # R√©cup√©rer les financements
    financements = data.get("financements", {})
    
    # Filtrer uniquement les pr√™ts (dictionnaires) nomm√©s avec "Pr√™t " pour √©viter les subventions
    pret_financements = {
        k: v for k, v in financements.items()
        if isinstance(v, dict) and k.startswith("Pr√™t ")
    }
    
    # Calculer les remboursements des emprunts
    remboursements_emprunts = calculate_remboursements_emprunts(pret_financements, years=3)
    
    # Autofinancement net = Capacit√© d'autofinancement - Remboursements des emprunts
    autofinancement_net = [
        capacite_autofinancement[i] - remboursements_emprunts[i]
        for i in range(3)
    ]
    
    # Pr√©parer les valeurs mon√©taires
    values = {
        "Ann√©e 1": [
            resultat_exercice[0],
            dotations_aux_amortissements[0],
            capacite_autofinancement[0],
            remboursements_emprunts[0],
            autofinancement_net[0]
        ],
        "Ann√©e 2": [
            resultat_exercice[1],
            dotations_aux_amortissements[1],
            capacite_autofinancement[1],
            remboursements_emprunts[1],
            autofinancement_net[1]
        ],
        "Ann√©e 3": [
            resultat_exercice[2],
            dotations_aux_amortissements[2],
            capacite_autofinancement[2],
            remboursements_emprunts[2],
            autofinancement_net[2]
        ]
    }
    
    # Pr√©parer le tableau final avec les labels
    capacite_fonc = [
        "R√©sultat de l'exercice",
        "+ Dotation aux amortissements",
        "Capacit√© d'autofinancement",
        "- Remboursements des emprunts",
        "Autofinancement net"
    ]
    
    data_table = {
        "Capacit√© d'autofinancement": capacite_fonc,
        "Ann√©e 1": values["Ann√©e 1"],
        "Ann√©e 2": values["Ann√©e 2"],
        "Ann√©e 3": values["Ann√©e 3"]
    }
    
    # Cr√©er le DataFrame avec les donn√©es
    df = pd.DataFrame(data_table)
    
    # D√©finir l'ordre des colonnes
    columns_order = ["Capacit√© d'autofinancement",
                     "Ann√©e 1",
                     "Ann√©e 2",
                     "Ann√©e 3"]
    df = df[columns_order]
    
    # D√©finir la fonction de formatage
    def format_value(x):
        if x == 0.0:
            return "-"
        else:
            return f"{x:,.2f} $"
    
    # Afficher le tableau avec une mise en forme am√©lior√©e
    st.dataframe(
        df.style.format({
            "Ann√©e 1": format_value,
            "Ann√©e 2": format_value,
            "Ann√©e 3": format_value,
        }).set_properties(**{
            'text-align': 'right'
        }).set_table_styles([{
            'selector': 'th',
            'props': [('text-align', 'center')]
        }])
    )
    
    # Stocker les r√©sultats dans les donn√©es
    data["capacite_autofinancement"] = {
        "resultat_exercice": resultat_exercice,
        "dotations_aux_amortissements": dotations_aux_amortissements,
        "capacite_autofinancement": capacite_autofinancement,
        "remboursements_emprunts": remboursements_emprunts,
        "autofinancement_net": autofinancement_net
    }
    
    # Enregistrer les donn√©es mises √† jour dans la session
    st.session_state["data"] = data
    
    # Pr√©parer les donn√©es d'exportation pour Capacit√© d'Autofinancement
    export_table_capacite = []
    for idx, label in enumerate(capacite_fonc):
        export_table_capacite.append({
            "Description": label,
            "Ann√©e 1": values["Ann√©e 1"][idx],
            "Ann√©e 2": values["Ann√©e 2"][idx],
            "Ann√©e 3": values["Ann√©e 3"][idx]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_capacite_autofinancement'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_capacite
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    if st.button("T√©l√©charger le Document Word Complet", key="download_word_complet_investissements_et_financements"):
        telecharger_document_complet()

    
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_seuil_rentabilite_economique():
    st.title("Seuil de rentabilit√© √©conomique")
    
    # V√©rifier si les donn√©es sont pr√©sentes dans la session
    if "data" not in st.session_state:
        st.error("Les donn√©es ne sont pas initialis√©es. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]

    # R√©cup√©rer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # R√©cup√©rer les donn√©es n√©cessaires avec les cl√©s exactes
    compte_resultat = data.get("compte_de_resultat", {})
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    
    # R√©cup√©rer 'ventes_production_reelle' et 'achats_consommes' du 'soldes_intermediaires'
    ventes_production_reelle = soldes_intermediaires.get("ventes_production_reelle", [0.0, 0.0, 0.0])
    achats_consommes = soldes_intermediaires.get("achats_consommes", [0.0, 0.0, 0.0])
    
    # R√©cup√©rer les charges n√©cessaires du 'soldes_intermediaires'
    charges_externes = soldes_intermediaires.get("charges_externes", [0.0, 0.0, 0.0])
    impots_et_taxes = soldes_intermediaires.get("impots_et_taxes", [0.0, 0.0, 0.0])
    charges_personnel = soldes_intermediaires.get("charges_personnel", [0.0, 0.0, 0.0])
    dotations_aux_amortissements = soldes_intermediaires.get("dotations_aux_amortissements", [0.0, 0.0, 0.0])
    charges_financieres = soldes_intermediaires.get("charges_financieres", [0.0, 0.0, 0.0])
    
    # R√©cup√©rer 'resultat_avant_impots' du 'compte_resultat'
    resultat_avant_impots = compte_resultat.get("resultat_avant_impots", [0.0, 0.0, 0.0])
    
    # V√©rifier si les listes ont 3 √©l√©ments
    if not (len(ventes_production_reelle) == len(achats_consommes) == len(charges_externes) == len(impots_et_taxes) == len(charges_personnel) == len(dotations_aux_amortissements) == len(charges_financieres) == len(resultat_avant_impots) == 3):
        st.error("Les listes de donn√©es ne contiennent pas exactement 3 √©l√©ments. Veuillez v√©rifier les donn√©es.")
        return
    
    # Calcul des Co√ªts fixes pour chaque ann√©e
    couts_fixes = []
    for i in range(3):
        cout_fix = (
            charges_externes[i] +
            impots_et_taxes[i] +
            charges_personnel[i] +
            dotations_aux_amortissements[i] +
            charges_financieres[i]
        )
        couts_fixes.append(cout_fix)
   
    # Total des co√ªts variables = Achats consomm√©s
    total_couts_variables = achats_consommes.copy()
    
    # Marge sur co√ªts variables
    marge_sur_couts_variables = []
    for i in range(3):
        marge = ventes_production_reelle[i] - total_couts_variables[i]
        marge_sur_couts_variables.append(marge)
    
    # Taux de marge sur co√ªts variables
    taux_marge_sur_couts_variables = []
    for i in range(3):
        if ventes_production_reelle[i] != 0:
            taux_marge = marge_sur_couts_variables[i] / ventes_production_reelle[i]
        else:
            taux_marge = 0.0
        taux_marge_sur_couts_variables.append(taux_marge)
    
    # Total des charges
    total_charges = []
    for i in range(3):
        total_charge = couts_fixes[i] + total_couts_variables[i]
        total_charges.append(total_charge)
   
    # Seuil de rentabilit√© (CA)
    seuil_rentabilite_ca = []
    for i in range(3):
        if taux_marge_sur_couts_variables[i] != 0:
            seuil_ca = couts_fixes[i] / taux_marge_sur_couts_variables[i]
        else:
            seuil_ca = 0.0
        seuil_rentabilite_ca.append(seuil_ca)
    
    # Exc√©dent / insuffisance
    excedent_insuffisance = []
    for i in range(3):
        excedent = ventes_production_reelle[i] - seuil_rentabilite_ca[i]
        excedent_insuffisance.append(excedent)
    
    # Point mort
    point_mort_ca_par_jour_ouvre = []
    for i in range(3):
        point_mort = seuil_rentabilite_ca[i] / 250
        point_mort_ca_par_jour_ouvre.append(point_mort)
    
    # Pr√©paration des donn√©es pour le tableau
    data_table = {
        "Seuil de rentabilite_economique": [
            "Ventes + Production r√©elle",
            "Achats consomm√©s",
            "Total des co√ªts variables",
            "Marge sur co√ªts variables",
            "Taux de marge sur co√ªts variables",
            "Co√ªts fixes",
            "Total des charges",
            "R√©sultat courant avant imp√¥ts",
            "Seuil de rentabilite (chiffre d'affaires)",
            "Exc√©dent / insuffisance",
            "Point mort en chiffre d'affaires par jour ouvr√©"
        ],
        "Ann√©e 1": [
            ventes_production_reelle[0],
            achats_consommes[0],
            total_couts_variables[0],
            marge_sur_couts_variables[0],
            taux_marge_sur_couts_variables[0],
            couts_fixes[0],
            total_charges[0],
            resultat_avant_impots[0],
            seuil_rentabilite_ca[0],
            excedent_insuffisance[0],
            point_mort_ca_par_jour_ouvre[0]
        ],
        "Ann√©e 2": [
            ventes_production_reelle[1],
            achats_consommes[1],
            total_couts_variables[1],
            marge_sur_couts_variables[1],
            taux_marge_sur_couts_variables[1],
            couts_fixes[1],
            total_charges[1],
            resultat_avant_impots[1],
            seuil_rentabilite_ca[1],
            excedent_insuffisance[1],
            point_mort_ca_par_jour_ouvre[1]
        ],
        "Ann√©e 3": [
            ventes_production_reelle[2],
            achats_consommes[2],
            total_couts_variables[2],
            marge_sur_couts_variables[2],
            taux_marge_sur_couts_variables[2],
            couts_fixes[2],
            total_charges[2],
            resultat_avant_impots[2],
            seuil_rentabilite_ca[2],
            excedent_insuffisance[2],
            point_mort_ca_par_jour_ouvre[2]
        ]
    }
    
    # Cr√©er le DataFrame
    df = pd.DataFrame(data_table)
    
    # D√©finir "Seuil de rentabilite_economique" comme index
    df.set_index("Seuil de rentabilite_economique", inplace=True)
    
    # √âtape 4: D√©finir une fonction de formatage
    # √âtape 4: D√©finir une fonction de formatage
    def format_row(row):
        if row.name == "Taux de marge sur co√ªts variables":
            # Formater en pourcentage avec deux d√©cimales
            return row.apply(lambda x: "{:.2f} %".format(x))
        else:
            # Formater en dollars avec s√©parateurs de milliers et sans d√©cimales
            return row.apply(lambda x: "{:,.0f} $".format(x) if isinstance(x, (int, float)) else x)

    # √âtape 5: Appliquer le formatage
    df_formatted = df.apply(format_row, axis=1)

    # √âtape 6: Afficher le tableau format√© avec Streamlit
    st.table(df_formatted)
    # Stocker les r√©sultats dans les donn√©es
    data["seuil_rentabilite_economique"] = {
        "ventes_production_reelle": ventes_production_reelle,
        "achats_consommes": achats_consommes,
        "total_couts_variables": total_couts_variables,
        "marge_sur_couts_variables": marge_sur_couts_variables,
        "taux_marge_sur_couts_variables": taux_marge_sur_couts_variables,
        "couts_fixes": couts_fixes,
        "total_charges": total_charges,
        "resultat_courant_avant_impots": resultat_avant_impots,
        "seuil_rentabilite_ca": seuil_rentabilite_ca,
        "excedent_insuffisance": excedent_insuffisance,
        "point_mort_ca_par_jour_ouvre": point_mort_ca_par_jour_ouvre
    }
    
    # Enregistrer les donn√©es dans la session
    st.session_state["data"] = data
    
    # Pr√©parer les donn√©es d'exportation pour Seuil de rentabilit√© √©conomique
    export_table_seuil = []
    for idx, label in enumerate(data_table["Seuil de rentabilite_economique"]):
        export_table_seuil.append({
            "Description": label,
            "Ann√©e 1": data_table["Ann√©e 1"][idx],
            "Ann√©e 2": data_table["Ann√©e 2"][idx],
            "Ann√©e 3": data_table["Ann√©e 3"][idx]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_seuil_rentabilite_economique'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_seuil
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le fichier Word complet contenant tous les tableaux avec une cl√© unique
    if st.button("T√©l√©charger le Document Word Complet", key="download_word_complet_seuil_rentabilite_economique"):
        telecharger_document_complet()


def telecharger_document_complet():
    # R√©cup√©rer les donn√©es export√©es de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    
    # V√©rifiez que toutes les donn√©es sont pr√©sentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data")
    ]):
        st.error("Toutes les sections doivent √™tre remplies avant de t√©l√©charger le document complet.")
        return
    
    # Cr√©er un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Cr√©er le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Dur√©e (mois)", "")) if row.get("Dur√©e (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes sp√©cifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Cr√©er le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Ann√©e 1", "")
        row_cells[2].text = row.get("Ann√©e 2", "")
        row_cells[3].text = row.get("Ann√©e 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section D√©tail des Amortissements ###
    doc.add_heading('D√©tail des Amortissements', level=1)
    
    # Cr√©er le tableau D√©tail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les donn√©es Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Ann√©e 1", "")
        row_cells_amort[2].text = row.get("Ann√©e 2", "")
        row_cells_amort[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calcul√©s en fonction de la dur√©e d'amortissement sp√©cifi√©e.")
    
    ### 4. Ajouter la section Compte de R√©sultats Pr√©visionnel ###
    doc.add_heading('Compte de R√©sultats Pr√©visionnel', level=1)
    
    # Cr√©er le tableau Compte de R√©sultats Pr√©visionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Ann√©e 1", "")
        row_cells_compte[2].text = row.get("Ann√©e 2", "")
        row_cells_compte[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 5. Ajouter la section Soldes Interm√©diaires de Gestion ###
    doc.add_heading('Soldes interm√©diaires de gestion', level=1)
    
    # Cr√©er le tableau Soldes interm√©diaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 6. Ajouter la section Capacit√© d'Autofinancement ###
    doc.add_heading('Capacit√© d\'autofinancement', level=1)
    
    # Cr√©er le tableau Capacit√© d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilit√© √âconomique ###
    doc.add_heading('Seuil de rentabilit√© √©conomique', level=1)
    
    # Cr√©er le tableau Seuil de Rentabilit√© √âconomique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Ann√©e 1', 0):,.1f} "
        row_cells_seuil[2].text = f"{row.get('Ann√©e 2', 0):,.1f} "
        row_cells_seuil[3].text = f"{row.get('Ann√©e 3', 0):,.1f} "
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 8. Ajouter la section Besoin en Fonds de Roulement ###
    doc.add_heading('Besoin en fonds de roulement', level=1)
    
    # Cr√©er le tableau Besoin en Fonds de Roulement dans Word
    table_bfr = doc.add_table(rows=1, cols=5)
    table_bfr.style = 'Light List Accent 1'
    table_bfr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_bfr = table_bfr.rows[0].cells
    headers_bfr = ["Analyse clients / fournisseurs", "D√©lai jours", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_bfr):
        hdr_cells_bfr[i].text = header
        for paragraph in hdr_cells_bfr[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_bfr[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_bfr['table_data']:
        row_cells_bfr = table_bfr.add_row().cells
        row_cells_bfr[0].text = row.get("Analyse clients / fournisseurs", "")
        row_cells_bfr[1].text = row.get("D√©lai jours", "")
        row_cells_bfr[2].text = row.get("Ann√©e 1", "")
        row_cells_bfr[3].text = row.get("Ann√©e 2", "")
        row_cells_bfr[4].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de t√©l√©chargement
    st.download_button(
        label="T√©l√©charger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a √©t√© g√©n√©r√© avec succ√®s !")

def page_besoin_fonds_roulement():
    st.title("Besoin en fonds de roulement")
    
    # V√©rifier si les donn√©es sont pr√©sentes dans la session
    if "data" not in st.session_state:
        st.error("Les donn√©es ne sont pas initialis√©es. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]
    
    # R√©cup√©rer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")

    # R√©cup√©rer les d√©lais clients et fournisseurs depuis "besoin_fonds_roulement"
    besoin_fonds = data.get("fonds_roulement", {})
    delai_clients = besoin_fonds.get("duree_credits_clients", 0)  # Dur√©e moyenne des cr√©dits accord√©s aux clients en jours
    delai_fournisseurs = besoin_fonds.get("duree_dettes_fournisseurs", 0)  # Dur√©e moyenne des cr√©dits accord√©s aux fournisseurs en jours

    st.write("---")
    
    # R√©cup√©rer "Ventes + Production r√©elle" et "Achats consomm√©s" depuis "soldes_intermediaires_de_gestion"
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    ventes_production_reelle = soldes_intermediaires.get("ventes_production_reelle", [0.0, 0.0, 0.0])
    achats_consommes = soldes_intermediaires.get("achats_consommes", [0.0, 0.0, 0.0])
    
    # V√©rifier si les d√©lais sont non nuls
    if delai_clients == 0 or delai_fournisseurs == 0:
        st.error("Les d√©lais de paiement des clients et des fournisseurs ne sont pas renseign√©s. Veuillez les saisir dans la section 'Votre besoin en fonds de roulement'.")
        return
    
    # V√©rifier si les listes ont 3 √©l√©ments
    if not (len(ventes_production_reelle) == len(achats_consommes) == 3):
        st.error("Les listes de 'Ventes + Production r√©elle' ou 'Achats consomm√©s' ne contiennent pas exactement 3 √©l√©ments. Veuillez v√©rifier les donn√©es.")
        return
    
    # Calculer le Volume cr√©dit client HT = Ventes + Production r√©elle / (delai_jours * 365)
    volume_credit_client_ht = []
    for i in range(3):
        vcc_ht = (ventes_production_reelle[i] * delai_clients) / 365
        volume_credit_client_ht.append(vcc_ht)
    
    # Calculer le Volume dettes fournisseurs HT = Achats consomm√©s / (delai_jours * 365)
    volume_dettes_fournisseurs_ht = []
    for i in range(3):
        vdf_ht = (achats_consommes[i] * delai_fournisseurs) / 365
        volume_dettes_fournisseurs_ht.append(vdf_ht)
    
    # Calculer le Besoin en fonds de roulement (BFR) = Volume cr√©dit client HT - Volume dettes fournisseurs HT
    bfr = [volume_credit_client_ht[i] - volume_dettes_fournisseurs_ht[i] for i in range(3)]
    
    # Afficher les r√©sultats interm√©diaires
    st.write("### R√©sultats des Calculs")
    st.write(f"**Volume cr√©dit client HT Ann√©e 1** : {volume_credit_client_ht[0]:.2f} $")
    st.write(f"**Volume dettes fournisseurs HT Ann√©e 1** : {volume_dettes_fournisseurs_ht[0]:.2f} $")
    st.write(f"**Besoin en fonds de roulement Ann√©e 1** : {bfr[0]:.2f} $")
    st.write(f"**Volume cr√©dit client HT Ann√©e 2** : {volume_credit_client_ht[1]:.2f} $")
    st.write(f"**Volume dettes fournisseurs HT Ann√©e 2** : {volume_dettes_fournisseurs_ht[1]:.2f} $")
    st.write(f"**Besoin en fonds de roulement Ann√©e 2** : {bfr[1]:.2f} $")
    st.write(f"**Volume cr√©dit client HT Ann√©e 3** : {volume_credit_client_ht[2]:.2f} $")
    st.write(f"**Volume dettes fournisseurs HT Ann√©e 3** : {volume_dettes_fournisseurs_ht[2]:.2f} $")
    st.write(f"**Besoin en fonds de roulement Ann√©e 3** : {bfr[2]:.2f} $")
    
    # Pr√©parer les donn√©es pour le tableau
    data_table = {
        "Analyse clients / fournisseurs": [
            "Besoins",
            "Volume cr√©dit client HT",
            "Ressources",
            "Volume dettes fournisseurs HT",
            "Besoin en fonds de roulement"
        ],
        "D√©lai jours": [
            "",
            f"{delai_clients}",
            "",
            f"{delai_fournisseurs}",
            ""
        ],
        "Ann√©e 1": [
            "",
            f"{volume_credit_client_ht[0]:.2f} $",
            "",
            f"{volume_dettes_fournisseurs_ht[0]:.2f} $",
            f"{bfr[0]:.2f} $"
        ],
        "Ann√©e 2": [
            "",
            f"{volume_credit_client_ht[1]:.2f} $",
            "",
            f"{volume_dettes_fournisseurs_ht[1]:.2f} $",
            f"{bfr[1]:.2f} $"
        ],
        "Ann√©e 3": [
            "",
            f"{volume_credit_client_ht[2]:.2f} $",
            "",
            f"{volume_dettes_fournisseurs_ht[2]:.2f} $",
            f"{bfr[2]:.2f} $"
        ]
    }
    
    df = pd.DataFrame(data_table)
    
    # Afficher le tableau
    st.write("### Tableau du Besoin en fonds de roulement")
    st.table(df)
    
    # Stocker les r√©sultats dans les donn√©es
    data["besoin_fonds_roulement"] = {
        "delai_clients": delai_clients,
        "delai_fournisseurs": delai_fournisseurs,
        "volume_credit_client_ht": volume_credit_client_ht,
        "volume_dettes_fournisseurs_ht": volume_dettes_fournisseurs_ht,
        "bfr": bfr
    }
    
    # Enregistrer les donn√©es dans la session
    st.session_state["data"] = data
    
    # Pr√©parer les donn√©es d'exportation pour Besoin en fonds de roulement
    export_table_bfr = []
    for idx, label in enumerate(data_table["Analyse clients / fournisseurs"]):
        export_table_bfr.append({
            "Analyse clients / fournisseurs": label,
            "D√©lai jours": data_table["D√©lai jours"][idx],
            "Ann√©e 1": data_table["Ann√©e 1"][idx],
            "Ann√©e 2": data_table["Ann√©e 2"][idx],
            "Ann√©e 3": data_table["Ann√©e 3"][idx]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_besoin_fonds_roulement'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_bfr
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le fichier Word complet contenant tous les tableaux avec une cl√© unique
    if st.button("T√©l√©charger le Document Word Complet", key="download_word_complet_besoin_fonds_roulement"):
        telecharger_document_complet()


def telecharger_document_complet():
    # R√©cup√©rer les donn√©es export√©es de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
    
    # V√©rifiez que toutes les donn√©es sont pr√©sentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data"),
        export_data_plan_financement.get("table_data")
    ]):
        st.error("Toutes les sections doivent √™tre remplies avant de t√©l√©charger le document complet.")
        return
    
    # Cr√©er un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Cr√©er le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Dur√©e (mois)", "")) if row.get("Dur√©e (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes sp√©cifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Cr√©er le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Ann√©e 1", "")
        row_cells[2].text = row.get("Ann√©e 2", "")
        row_cells[3].text = row.get("Ann√©e 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section D√©tail des Amortissements ###
    doc.add_heading('D√©tail des Amortissements', level=1)
    
    # Cr√©er le tableau D√©tail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les donn√©es Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Ann√©e 1", "")
        row_cells_amort[2].text = row.get("Ann√©e 2", "")
        row_cells_amort[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calcul√©s en fonction de la dur√©e d'amortissement sp√©cifi√©e.")
    
    ### 4. Ajouter la section Compte de R√©sultats Pr√©visionnel ###
    doc.add_heading('Compte de R√©sultats Pr√©visionnel', level=1)
    
    # Cr√©er le tableau Compte de R√©sultats Pr√©visionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Ann√©e 1", "")
        row_cells_compte[2].text = row.get("Ann√©e 2", "")
        row_cells_compte[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 5. Ajouter la section Soldes Interm√©diaires de Gestion ###
    doc.add_heading('Soldes interm√©diaires de gestion', level=1)
    
    # Cr√©er le tableau Soldes interm√©diaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 6. Ajouter la section Capacit√© d'Autofinancement ###
    doc.add_heading('Capacit√© d\'autofinancement', level=1)
    
    # Cr√©er le tableau Capacit√© d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilit√© √âconomique ###
    doc.add_heading('Seuil de rentabilit√© √©conomique', level=1)
    
    # Cr√©er le tableau Seuil de Rentabilit√© √âconomique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Ann√©e 1', 0):,.2f} "
        row_cells_seuil[2].text = f"{row.get('Ann√©e 2', 0):,.2f} "
        row_cells_seuil[3].text = f"{row.get('Ann√©e 3', 0):,.2f} "
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 8. Ajouter la section Plan de Financement √† Trois Ans ###
    doc.add_heading('Plan de financement √† trois ans', level=1)
    
    # Cr√©er le tableau Plan de Financement √† Trois Ans dans Word
    table_plan = doc.add_table(rows=1, cols=5)
    table_plan.style = 'Light List Accent 1'
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_plan = table_plan.rows[0].cells
    headers_plan = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_plan):
        hdr_cells_plan[i].text = header
        for paragraph in hdr_cells_plan[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_plan[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_plan_financement['table_data']:
        row_cells_plan = table_plan.add_row().cells
        row_cells_plan[0].text = row.get("Plan de financement √† trois ans", "")
        row_cells_plan[1].text = row.get("Ann√©e 1", "")
        row_cells_plan[2].text = row.get("Ann√©e 2", "")
        row_cells_plan[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de t√©l√©chargement
    st.download_button(
        label="T√©l√©charger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a √©t√© g√©n√©r√© avec succ√®s !")




def page_plan_financement_trois_ans(): 
    st.title("Plan de financement √† trois ans")
    
    # V√©rifier si les donn√©es sont pr√©sentes dans la session
    if "data" not in st.session_state:
        st.error("Les donn√©es ne sont pas initialis√©es. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]
    
    # R√©cup√©rer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # R√©cup√©rer les besoins d√©marrage
    besoins_demarrage = data.get("besoins_demarrage", {})
    
    # Calcul des Immobilisations incorporelles et corporelles
    # D√©finissez quels √©l√©ments de "besoins_demarrage" correspondent √† chaque cat√©gorie
    immobilisations_inc = sum([
        besoins_demarrage.get("Frais d‚Äô√©tablissement", 0),
        besoins_demarrage.get("Frais d‚Äôouverture de compteurs", 0),
        besoins_demarrage.get("Logiciels, formations", 0),
        besoins_demarrage.get("D√©p√¥t de marque", 0),
        besoins_demarrage.get("Droits d‚Äôentr√©e", 0),
        besoins_demarrage.get("Achat fonds de commerce ou parts", 0),
        besoins_demarrage.get("Droit au bail", 0),
        besoins_demarrage.get("Caution ou d√©p√¥t de garantie", 0),
        besoins_demarrage.get("Frais de dossier", 0),
        besoins_demarrage.get("Frais de notaire", 0),
    ])
    
    immobilisations_corp = sum([
        besoins_demarrage.get("Enseigne et √©l√©ments de communication", 0),
        besoins_demarrage.get("V√©hicule", 0),
        besoins_demarrage.get("Mat√©riel professionnel", 0),
        besoins_demarrage.get("Mat√©riel autre", 0),
        besoins_demarrage.get("Mat√©riel de bureau", 0)
    ])
    
    immobilisations = [
        immobilisations_inc + immobilisations_corp,  # Ann√©e 1
        0.0,  # Ann√©e 2
        0.0   # Ann√©e 3
    ]
    
    # Acquisition des stocks
    acquisition_stocks = [
        besoins_demarrage.get("Stock de mati√®res et produits", 0),
        0.0,  # Ann√©e 2
        0.0   # Ann√©e 3
    ]
    
    # Variation du Besoin en fonds de roulement (BFR)
    besoin_fonds = data.get("besoin_fonds_roulement", {})
    bfr = besoin_fonds.get("bfr", [0.0, 0.0, 0.0])
    
    # Variation BFR = BFR ann√©e n - BFR ann√©e n-1
    variation_bfr = [
        bfr[0],                    # Variation en ann√©e 1 (BFR ann√©e 1 - BFR ann√©e 0)
        bfr[1] - bfr[0],           # Variation en ann√©e 2
        bfr[2] - bfr[1]            # Variation en ann√©e 3
    ]
    
    # Remboursement d'emprunts
    capacite_autofinancement = data.get("capacite_autofinancement", {})
    remboursements_emprunts = capacite_autofinancement.get("remboursements_emprunts", [0.0, 0.0, 0.0])
    
    # Total des besoins
    total_besoins = [
        immobilisations[0] + acquisition_stocks[0] + variation_bfr[0] + remboursements_emprunts[0],
        immobilisations[1] + acquisition_stocks[1] + variation_bfr[1] + remboursements_emprunts[1],
        immobilisations[2] + acquisition_stocks[2] + variation_bfr[2] + remboursements_emprunts[2]
    ]
    
    # Apport personnel
    financements = data.get("financements", {})
    apport_personnel = financements.get("Apport personnel ou familial", 0.0)
    apports_nature = financements.get("Apports en nature (en valeur)", 0.0)
    apport_total = apport_personnel + apports_nature
    apport_personnel_list = [apport_total, 0.0, 0.0] 
    
    # Emprunts
    pret_1 = financements.get("Pr√™t 1", {}).get("montant", 0.0)
    pret_2 = financements.get("Pr√™t 2", {}).get("montant", 0.0)
    pret_3 = financements.get("Pr√™t 3", {}).get("montant", 0.0)
    total_emprunts = pret_1 + pret_2 + pret_3
    emprunts = [total_emprunts, 0.0, 0.0]  # Supposons que les emprunts sont en ann√©e 1
    
    # Subventions
    subvention_1 = financements.get("Subvention 1", {}).get("montant", 0.0)
    subvention_2 = financements.get("Subvention 2", {}).get("montant", 0.0)
    subventions = subvention_1 + subvention_2
    subventions_list = [subventions, 0.0, 0.0]  # Supposons que les subventions sont en ann√©e 1
    
    # Autres financements
    autres_financements = financements.get("Autre financement", 0.0)
    autres_financements_list = [autres_financements, 0.0, 0.0]  # Supposons que c'est en ann√©e 1
    
    # Capacit√© d'auto-financement
    capacite_autofinancement_values = capacite_autofinancement.get("capacite_autofinancement", [0.0, 0.0, 0.0])
    
    # Total des ressources
    total_ressources = [
        apport_personnel_list[0] + emprunts[0] + subventions_list[0] + autres_financements_list[0] + capacite_autofinancement_values[0],
        apport_personnel_list[1] + emprunts[1] + subventions_list[1] + autres_financements_list[1] + capacite_autofinancement_values[1],
        apport_personnel_list[2] + emprunts[2] + subventions_list[2] + autres_financements_list[2] + capacite_autofinancement_values[2]
    ]
    
    # Variation de tr√©sorerie
    variation_tresorerie = [
        total_ressources[0] - total_besoins[0],
        total_ressources[1] - total_besoins[1],
        total_ressources[2] - total_besoins[2]
    ]
    
    # Exc√©dent de tr√©sorerie (cumulatif)
    excedent_tresorerie = []
    cumul_excedent = 0.0
    for i in range(3):
        cumul_excedent += variation_tresorerie[i]
        excedent_tresorerie.append(cumul_excedent)
    
    # Pr√©paration des donn√©es pour le tableau
    data_table = {
        "Plan de financement √† trois ans": [
            "Immobilisations",
            "Acquisition des stocks",
            "Variation du Besoin en fonds de roulement",
            "Remboursement d'emprunts",
            "Total des besoins",
            "Apport personnel",
            "Emprunts",
            "Subventions",
            "Autres financements",
            "Capacit√© d'auto-financement",
            "Total des ressources",
            "Variation de tr√©sorerie",
            "Exc√©dent de tr√©sorerie"
        ],
        "Ann√©e 1": [
            f"{immobilisations[0]:,.2f} $",
            f"{acquisition_stocks[0]:,.2f} $",
            f"{variation_bfr[0]:,.2f} $",
            f"{remboursements_emprunts[0]:,.2f} $",
            f"{total_besoins[0]:,.2f} $",
            f"{apport_personnel_list[0]:,.2f} $",
            f"{emprunts[0]:,.2f} $",
            f"{subventions_list[0]:,.2f} $",
            f"{autres_financements_list[0]:,.2f} $",
            f"{capacite_autofinancement_values[0]:,.2f} $",
            f"{total_ressources[0]:,.2f} $",
            f"{variation_tresorerie[0]:,.2f} $",
            f"{excedent_tresorerie[0]:,.2f} $"
        ],
        "Ann√©e 2": [
            f"{immobilisations[1]:,.2f} $",
            f"{acquisition_stocks[1]:,.2f} $",
            f"{variation_bfr[1]:,.2f} $",
            f"{remboursements_emprunts[1]:,.2f} $",
            f"{total_besoins[1]:,.2f} $",
            f"{apport_personnel_list[1]:,.2f} $",
            f"{emprunts[1]:,.2f} $",
            f"{subventions_list[1]:,.2f} $",
            f"{autres_financements_list[1]:,.2f} $",
            f"{capacite_autofinancement_values[1]:,.2f} $",
            f"{total_ressources[1]:,.2f} $",
            f"{variation_tresorerie[1]:,.2f} $",
            f"{excedent_tresorerie[1]:,.2f} $"
        ],
        "Ann√©e 3": [
            f"{immobilisations[2]:,.2f} $",
            f"{acquisition_stocks[2]:,.2f} $",
            f"{variation_bfr[2]:,.2f} $",
            f"{remboursements_emprunts[2]:,.2f} $",
            f"{total_besoins[2]:,.2f} $",
            f"{apport_personnel_list[2]:,.2f} $",
            f"{emprunts[2]:,.2f} $",
            f"{subventions_list[2]:,.2f} $",
            f"{autres_financements_list[2]:,.2f} $",
            f"{capacite_autofinancement_values[2]:,.2f} $",
            f"{total_ressources[2]:,.2f} $",
            f"{variation_tresorerie[2]:,.2f} $",
            f"{excedent_tresorerie[2]:,.2f} $"
        ]
    }
    
    df = pd.DataFrame(data_table)
    st.write("### Tableau du Plan de financement √† trois ans")
    st.table(df)
    
    # Stocker les r√©sultats dans les donn√©es
    data["plan_financement"] = {
        "immobilisations": immobilisations,
        "acquisition_stocks": acquisition_stocks,
        "variation_bfr": variation_bfr,
        "remboursements_emprunts": remboursements_emprunts,
        "total_besoins": total_besoins,
        "apport_personnel": apport_personnel_list,
        "emprunts": emprunts,
        "subventions": subventions_list,
        "autres_financements": autres_financements_list,
        "capacite_autofinancement": capacite_autofinancement_values,
        "total_ressources": total_ressources,
        "variation_tresorerie": variation_tresorerie,
        "excedent_tresorerie": excedent_tresorerie
    }
    
    # Enregistrer les donn√©es dans la session
    st.session_state["data"] = data   
    
    # Pr√©parer les donn√©es d'exportation pour Plan de Financement √† Trois Ans
    export_table_plan_financement = []
    for idx, label in enumerate(data_table["Plan de financement √† trois ans"]):
        export_table_plan_financement.append({
            "Plan de financement √† trois ans": label,
            "Ann√©e 1": data_table["Ann√©e 1"][idx],
            "Ann√©e 2": data_table["Ann√©e 2"][idx],
            "Ann√©e 3": data_table["Ann√©e 3"][idx]
        })
    
    # Stocker les donn√©es d'exportation dans la session
    st.session_state['export_data_plan_financement_trois_ans'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_plan_financement
    }
    
    # Section Export
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le fichier Word complet contenant tous les tableaux avec une cl√© unique
    if st.button("T√©l√©charger le Document Word Complet", key="download_word_complet_plan_financement_trois_ans"):
        telecharger_document_complet()
        


def telecharger_document_complet():
    # R√©cup√©rer les donn√©es export√©es de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
    export_data_budget_tresorerie_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
    export_data_budget_tresorerie_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})
    print(export_data_plan_financement)
    st.write(export_data_plan_financement)
    # V√©rifiez que toutes les donn√©es n√©cessaires sont pr√©sentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data"),
        export_data_plan_financement.get("table_data"),
        export_data_budget_tresorerie_part1.get("table_data"),
        export_data_budget_tresorerie_part2.get("table_data")
    ]):
        st.error("Toutes les sections doivent √™tre remplies avant de t√©l√©charger le document complet.")
        return
    
    # Cr√©er un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    
    # Cr√©er le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Dur√©e (mois)", "")) if row.get("Dur√©e (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes sp√©cifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Cr√©er le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Ann√©e 1", "")
        row_cells[2].text = row.get("Ann√©e 2", "")
        row_cells[3].text = row.get("Ann√©e 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section D√©tail des Amortissements ###
    doc.add_heading('D√©tail des Amortissements', level=1)
    
    # Cr√©er le tableau D√©tail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les donn√©es Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Ann√©e 1", "")
        row_cells_amort[2].text = row.get("Ann√©e 2", "")
        row_cells_amort[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calcul√©s en fonction de la dur√©e d'amortissement sp√©cifi√©e.")
    
    ### 4. Ajouter la section Compte de R√©sultats Pr√©visionnel ###
    doc.add_heading('Compte de R√©sultats Pr√©visionnel', level=1)
    
    # Cr√©er le tableau Compte de R√©sultats Pr√©visionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Ann√©e 1", "")
        row_cells_compte[2].text = row.get("Ann√©e 2", "")
        row_cells_compte[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 5. Ajouter la section Soldes Interm√©diaires de Gestion ###
    doc.add_heading('Soldes interm√©diaires de gestion', level=1)
    
    # Cr√©er le tableau Soldes interm√©diaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 6. Ajouter la section Capacit√© d'Autofinancement ###
    doc.add_heading('Capacit√© d\'autofinancement', level=1)
    
    # Cr√©er le tableau Capacit√© d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Ann√©e 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Ann√©e 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Ann√©e 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilit√© √âconomique ###
    doc.add_heading('Seuil de rentabilit√© √©conomique', level=1)
    
    # Cr√©er le tableau Seuil de Rentabilit√© √âconomique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Ann√©e 1', 0):,.2f} "
        row_cells_seuil[2].text = f"{row.get('Ann√©e 2', 0):,.2f} "
        row_cells_seuil[3].text = f"{row.get('Ann√©e 3', 0):,.2f} "
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 8. Ajouter la section Plan de Financement √† Trois Ans ###
    doc.add_heading('Plan de financement √† trois ans', level=1)
    
    # Cr√©er le tableau Plan de Financement √† Trois Ans dans Word
    table_plan = doc.add_table(rows=1, cols=4)
    table_plan.style = 'Light List Accent 1'
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_plan = table_plan.rows[0].cells
    headers_plan = ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"]
    for i, header in enumerate(headers_plan):
        hdr_cells_plan[i].text = header
        for paragraph in hdr_cells_plan[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_plan[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_plan_financement['table_data']:
        row_cells_plan = table_plan.add_row().cells
        row_cells_plan[0].text = row.get("Plan de financement √† trois ans", "")
        row_cells_plan[1].text = row.get("Ann√©e 1", "")
        row_cells_plan[2].text = row.get("Ann√©e 2", "")
        row_cells_plan[3].text = row.get("Ann√©e 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")
    
    ### 9. Ajouter la section Budget Pr√©visionnel de Tr√©sorerie Partie 1 ###
    doc.add_heading('Budget pr√©visionnel de tr√©sorerie - Partie 1', level=1)
    
    # Cr√©er le premier tableau Budget pr√©visionnel de tr√©sorerie
    table_budget_part1 = doc.add_table(rows=1, cols=len(export_data_budget_tresorerie_part1['table_data'][0]))
    table_budget_part1.style = 'Light List Accent 1'
    table_budget_part1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ajouter les en-t√™tes
    headers_budget_part1 = export_data_budget_tresorerie_part1['table_data'][0].keys()
    hdr_cells_budget_part1 = table_budget_part1.rows[0].cells
    for i, header in enumerate(headers_budget_part1):
        hdr_cells_budget_part1[i].text = header
        for paragraph in hdr_cells_budget_part1[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_budget_part1[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les donn√©es du premier tableau
    for row in export_data_budget_tresorerie_part1['table_data'][1:]:
        row_cells = table_budget_part1.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = value
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les donn√©es du budget pr√©visionnel de tr√©sorerie - Partie 1 sont bas√©es sur les estimations fournies.")
    
    ### 10. Ajouter la section Budget Pr√©visionnel de Tr√©sorerie Partie 2 ###
    doc.add_heading('Budget pr√©visionnel de tr√©sorerie - Partie 2', level=1)
    
    # Cr√©er le deuxi√®me tableau Budget pr√©visionnel de tr√©sorerie
    table_budget_part2 = doc.add_table(rows=1, cols=len(export_data_budget_tresorerie_part2['table_data'][0]))
    table_budget_part2.style = 'Light List Accent 1'
    table_budget_part2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ajouter les en-t√™tes
    headers_budget_part2 = export_data_budget_tresorerie_part2['table_data'][0].keys()
    hdr_cells_budget_part2 = table_budget_part2.rows[0].cells
    for i, header in enumerate(headers_budget_part2):
        hdr_cells_budget_part2[i].text = header
        for paragraph in hdr_cells_budget_part2[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_budget_part2[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les donn√©es du deuxi√®me tableau
    for row in export_data_budget_tresorerie_part2['table_data'][1:]:
        row_cells = table_budget_part2.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = value
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les donn√©es du budget pr√©visionnel de tr√©sorerie - Partie 2 sont bas√©es sur les estimations fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de t√©l√©chargement
    st.download_button(
        label="T√©l√©charger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a √©t√© g√©n√©r√© avec succ√®s !")


import streamlit as st
import pandas as pd

def page_budget_previsionnel_tresorerie():
    st.title("Budget pr√©visionnel de tr√©sorerie")
    
    data = st.session_state.get("data", {})
    
    if not data:
        st.error("Les donn√©es ne sont pas initialis√©es. Veuillez initialiser la session.")
        return
    
    # R√©cup√©rer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "N/A")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "N/A")
    
    st.write(f"**Projet :** {projet} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp; **(Hors TVA)**")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Cr√©ation du budget pr√©visionnel pour la premi√®re ann√©e, mois par mois
    months = [f"Mois {i+1}" for i in range(12)] + ["TOTAL"]
    
    # Initialisation des structures de donn√©es
    encaissements = {}
    decaissements = {}
    solde_precedent = [0.0] * 12
    solde_mois = [0.0] * 12
    solde_tresorerie_cumul = [0.0] * 12
    
    # ----------------------------
    # Encaissements
    # ----------------------------
    
    # R√©cup√©rer les encaissements depuis "Plan de financement √† trois ans"
    plan_financement = data.get("plan_financement", {})
    apport_personnel = plan_financement.get("apport_personnel", [0.0, 0.0, 0.0])[0]
    emprunts = plan_financement.get("emprunts", [0.0, 0.0, 0.0])[0]
    subventions = plan_financement.get("subventions", [0.0, 0.0, 0.0])[0]
    autres_financements = plan_financement.get("autres_financements", [0.0, 0.0, 0.0])[0]
    
    encaissements["Apport personnel"] = [apport_personnel] + [0.0]*11
    encaissements["Emprunts"] = [emprunts] + [0.0]*11
    encaissements["Subventions"] = [subventions] + [0.0]*11
    encaissements["Autres financements"] = [autres_financements] + [0.0]*11
    
    # R√©cup√©rer les ventes depuis "Chiffre d'Affaires Pr√©visionnel"
    chiffre_affaires = data.get("chiffre_affaires", {})
    
    # Initialiser les listes de ventes mensuelles
    vente_marchandises_mensuel = []
    vente_services_mensuel = []
    
    # Remplir les ventes mensuelles de Marchandises
    for i in range(1, 13):
        key_ca = f"Marchandises_Mois {i}_ca"
        ca = chiffre_affaires.get(key_ca, 0.0)
        vente_marchandises_mensuel.append(ca)
    
    # Remplir les ventes mensuelles de Services
    for i in range(1, 13):
        key_ca = f"Services_Mois {i}_ca"
        ca = chiffre_affaires.get(key_ca, 0.0)
        vente_services_mensuel.append(ca)
    
    encaissements["Vente de marchandises"] = vente_marchandises_mensuel
    encaissements["Vente de services"] = vente_services_mensuel
    encaissements["Chiffre d'affaires (total)"] = [vente_marchandises_mensuel[i] + vente_services_mensuel[i] for i in range(12)]
    
    # Total des encaissements
    total_encaissements = []
    for i in range(12):
        total = (
            encaissements["Apport personnel"][i] +
            encaissements["Emprunts"][i] +
            encaissements["Subventions"][i] +
            encaissements["Autres financements"][i] +
            encaissements["Vente de marchandises"][i] +
            encaissements["Vente de services"][i]
        )
        total_encaissements.append(total)
    total_total_encaissements = sum(total_encaissements)
    total_encaissements.append(total_total_encaissements)
    
    # ----------------------------
    # D√©caissements
    # ----------------------------
    
    # R√©cup√©rer les donn√©es n√©cessaires pour les d√©caissements
    besoins_demarrage = data.get("besoins_demarrage", {})
    charges_variables = data.get("charges_variables", {})
    compte_resultat = data.get("compte_de_resultat", {})
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    capacite_autofinancement = data.get("capacite_autofinancement", {})
    
    # Immobilisations incorporelles et corporelles depuis "besoins_demarrage"
    immobilisations_incorporelles = sum([
        besoins_demarrage.get("Frais d‚Äô√©tablissement", 0.0),
        besoins_demarrage.get("Frais d‚Äôouverture de compteurs", 0.0),
        besoins_demarrage.get("Logiciels, formations", 0.0),
        besoins_demarrage.get("D√©p√¥t de marque", 0.0),
        besoins_demarrage.get("Droits d‚Äôentr√©e", 0.0),
        besoins_demarrage.get("Achat fonds de commerce ou parts", 0.0),
        besoins_demarrage.get("Droit au bail", 0.0),
        besoins_demarrage.get("Caution ou d√©p√¥t de garantie", 0.0),
        besoins_demarrage.get("Frais de dossier", 0.0),
        besoins_demarrage.get("Frais de notaire", 0.0),
    ])
    
    immobilisations_corporelles = sum([
        besoins_demarrage.get("Enseigne et √©l√©ments de communication", 0.0),
        besoins_demarrage.get("V√©hicule", 0.0),
        besoins_demarrage.get("Mat√©riel professionnel", 0.0),
        besoins_demarrage.get("Mat√©riel autre", 0.0),
        besoins_demarrage.get("Mat√©riel de bureau", 0.0)
    ])
    
    immobilisations_total = immobilisations_incorporelles + immobilisations_corporelles
    
    decaissements["Immobilisations incorporelles"] = [immobilisations_incorporelles] + [0.0]*11
    decaissements["Immobilisations corporelles"] = [immobilisations_corporelles] + [0.0]*11
    decaissements["Immobilisations (total)"] = [immobilisations_total] + [0.0]*11
    
    # Acquisition des stocks depuis "Stock de mati√®res et produits"
    acquisition_stocks = besoins_demarrage.get("Stock de mati√®res et produits", 0.0)
    decaissements["Acquisition stocks"] = [acquisition_stocks] + [0.0]*11
    
    # √âch√©ances emprunt : "Principal ann√©e 1" divis√© par 12
    remboursements_emprunts = capacite_autofinancement.get("remboursements_emprunts", [0.0, 0.0, 0.0])
    principal_annee1 = remboursements_emprunts[0]
    echeances_emprunt_mensuel = principal_annee1 / 12.0 if principal_annee1 > 0 else 0.0
    decaissements["√âch√©ances emprunt"] = [echeances_emprunt_mensuel] * 12
    
    # Achats de marchandises : "Vente de marchandises" * "le co√ªt d'achat de vos marchandises" de "Charges Variables"
    cout_achat_marchandises_pct = charges_variables.get("cout_achat_marchandises_pct", 100.0)
    if cout_achat_marchandises_pct == 0.0:
        cout_achat_marchandises_pct = 100.0  # Supposer 100% si non renseign√©
    
    achats_marchandises_mensuel = [vente_marchandises_mensuel[i] * cout_achat_marchandises_pct / 100.0 for i in range(12)]
    decaissements["Achats de marchandises"] = achats_marchandises_mensuel
    
    # Charges externes : R√©cup√©rer depuis "soldes_intermediaires_de_gestion"
    charges_externes_annee1 = soldes_intermediaires.get("charges_externes", [0.0, 0.0, 0.0])[0]
    charges_externes_mensuel = charges_externes_annee1 / 12.0 if charges_externes_annee1 > 0 else 0.0
    decaissements["Charges externes"] = [charges_externes_mensuel] * 12
    
    # Imp√¥ts et taxes
    impots_et_taxes_annee1 = compte_resultat.get("impots_et_taxes", [0.0, 0.0, 0.0])[0]
    impots_et_taxes_mensuel = impots_et_taxes_annee1 / 12.0 if impots_et_taxes_annee1 > 0 else 0.0
    decaissements["Imp√¥ts et taxes"] = [impots_et_taxes_mensuel] * 12
    
    # Salaires employ√©s, Charges sociales employ√©s, Pr√©l√®vement dirigeant(s), Charges sociales dirigeant(s), Frais bancaires, charges financi√®res
    salaires_employes_annee1 = compte_resultat.get("salaires_employes", [0.0, 0.0, 0.0])[0]
    charges_sociales_employes_annee1 = compte_resultat.get("charges_sociales_employes", [0.0, 0.0, 0.0])[0]
    prelevement_dirigeants_annee1 = compte_resultat.get("salaires_dirigeants", [0.0, 0.0, 0.0])[0]
    charges_sociales_dirigeants_annee1 = compte_resultat.get("charges_sociales_dirigeants", [0.0, 0.0, 0.0])[0]
    frais_bancaires_annuels = compte_resultat.get("total_frais_financiers", [0.0, 0.0, 0.0])[0]
    
    salaires_employes_mensuel = [salaires_employes_annee1 / 12.0] * 12
    charges_sociales_employes_mensuel = [charges_sociales_employes_annee1 / 12.0] * 12
    prelevement_dirigeants_mensuel = [prelevement_dirigeants_annee1 / 12.0] * 12
    charges_sociales_dirigeants_mensuel = [charges_sociales_dirigeants_annee1 / 12.0] * 12
    frais_bancaires_mensuel = [frais_bancaires_annuels / 12.0] * 12 if frais_bancaires_annuels > 0 else [0.0] * 12
    
    decaissements["Salaires employ√©s"] = salaires_employes_mensuel
    decaissements["Charges sociales employ√©s"] = charges_sociales_employes_mensuel
    decaissements["Pr√©l√®vement dirigeant(s)"] = prelevement_dirigeants_mensuel
    decaissements["Charges sociales dirigeant(s)"] = charges_sociales_dirigeants_mensuel
    decaissements["Frais bancaires, charges financi√®res"] = frais_bancaires_mensuel
    
    # ----------------------------
    # Total charges de personnel
    # ----------------------------
    # Calculer le total des charges de personnel pour chaque mois
    total_charges_personnel_mensuel = [
        salaires_employes_mensuel[i] + charges_sociales_employes_mensuel[i] +
        prelevement_dirigeants_mensuel[i] + charges_sociales_dirigeants_mensuel[i]
        for i in range(12)
    ]
    decaissements["Total charges de personnel"] = total_charges_personnel_mensuel
    
    # ----------------------------
    # Total des d√©caissements
    # ----------------------------
    
    # D√©finir les cl√©s √† inclure dans le total des d√©caissements
    decaissements_keys = [
        "Immobilisations (total)",
        "Acquisition stocks",
        "√âch√©ances emprunt",
        "Achats de marchandises",
        "Charges externes",
        "Imp√¥ts et taxes",
        "Total charges de personnel",
        "Frais bancaires, charges financi√®res"
    ]
    
    total_decaissements = []
    for i in range(12):
        total = sum([decaissements[key][i] for key in decaissements_keys])
        total_decaissements.append(total)
    total_total_decaissements = sum(total_decaissements)
    total_decaissements.append(total_total_decaissements)
    
    # ----------------------------
    # Calcul des Soldes
    # ----------------------------
    
    for i in range(12):
        solde_mois[i] = total_encaissements[i] - total_decaissements[i]
        solde_tresorerie_cumul[i] = solde_tresorerie_cumul[i - 1] + solde_mois[i] if i > 0 else solde_mois[i]
        solde_precedent[i] = solde_tresorerie_cumul[i - 1] if i > 0 else 0.0
    
    # Append totals to solde_mois, solde_precedent, solde_tresorerie_cumul
    total_solde_mois = sum(solde_mois)
    solde_mois.append(total_solde_mois)
    
    # Pour solde_precedent, le total n'est pas significatif, on peut ajouter une cha√Æne vide
    solde_precedent.append("")
    
    # Pour solde_tresorerie_cumul, on peut ajouter la derni√®re valeur cumulative
    solde_tresorerie_cumul.append(solde_tresorerie_cumul[-1])
    
    # ----------------------------
    # Pr√©paration des donn√©es pour le tableau
    # ----------------------------
    
    table_data = {"Description": months}
    
    # Encaissements
    for key in encaissements:
        amounts = encaissements[key]
        total = sum(amounts)
        amounts_with_total = amounts + [total]
        table_data[key] = [f"{value:,.2f} $" if value != 0 else "-" for value in amounts_with_total]
    
    # D√©caissements
    for key in decaissements:
        # Inclure toutes les lignes de d√©caissements
        amounts = decaissements[key]
        total = sum(amounts)
        # Remplacer 0 par '-' si n√©cessaire
        amounts_with_total = [f"{value:,.2f} $" if value != 0 else "-" for value in amounts] + [f"{total:,.2f} $" if total != 0 else "-"]
        table_data[key] = amounts_with_total
    
    # Totaux et soldes
    table_data["Total des encaissements"] = [f"{value:,.2f} $" if value != 0 else "-" for value in total_encaissements]
    table_data["Total des d√©caissements"] = [f"{value:,.2f} $" if value != 0 else "-" for value in total_decaissements]
    solde_precedent_formatted = [f"{value:,.2f} $" if isinstance(value, (int, float)) and value != 0 else "-" for value in solde_precedent]
    table_data["Solde pr√©c√©dent"] = solde_precedent_formatted
    table_data["Solde du mois"] = [f"{value:,.2f} $" if value != 0 else "-" for value in solde_mois]
    table_data["Solde de tr√©sorerie (cumul)"] = [f"{value:,.2f} $" if value != 0 else "-" for value in solde_tresorerie_cumul]
    
    # Assurer que toutes les listes ont la m√™me longueur
    max_length = max(len(lst) for lst in table_data.values())
    for key in table_data:
        if len(table_data[key]) < max_length:
            difference = max_length - len(table_data[key])
            table_data[key] += [""] * difference  # Remplir avec des cha√Ænes vides si n√©cessaire
        elif len(table_data[key]) > max_length:
            table_data[key] = table_data[key][:max_length]  # Tronquer si trop long
    
    # Cr√©ation du DataFrame complet
    df_full = pd.DataFrame(table_data)
    df_full.set_index("Description", inplace=True)
    df_full = df_full.T  # Transposer pour avoir les mois comme colonnes
    
    # S√©paration en deux tableaux
    # Tableau 1 : Mois 1 √† Mois 5
    columns_part1 = ["Mois 1", "Mois 2", "Mois 3", "Mois 4", "Mois 5"]
    df_part1 = df_full[columns_part1]
    
    # Tableau 2 : Mois 6 √† Mois 12 + TOTAL
    columns_part2 = ["Mois 6", "Mois 7", "Mois 8", "Mois 9", "Mois 10", "Mois 11", "Mois 12", "TOTAL"]
    df_part2 = df_full[columns_part2]
    
    ### 3. Ajouter la section Budget Pr√©visionnel de Tr√©sorerie ###
    # (Les deux tableaux seront ajout√©s dans telecharger_document_complet())
    
    ### 4. Affichage des tableaux s√©par√©s ###
    st.subheader("Budget pr√©visionnel de tr√©sorerie")
    st.table(df_part1)
    
    st.subheader("Budget pr√©visionnel de tr√©sorerie (suite)")
    st.table(df_part2)
    
    # ----------------------------
    # Stockage des r√©sultats dans les donn√©es
    # ----------------------------
    
    data["budget_previsionnel_tresorerie"] = {
        "encaissements": encaissements,
        "decaissements": decaissements,
        "total_encaissements": total_encaissements,
        "total_decaissements": total_decaissements,
        "solde_precedent": solde_precedent,
        "solde_mois": solde_mois,
        "solde_tresorerie_cumul": solde_tresorerie_cumul
    }
    
    # Enregistrer les donn√©es dans la session
    st.session_state["data"] = data   
    
    # ----------------------------
    # Pr√©paration des donn√©es d'exportation pour Budget Pr√©visionnel de Tr√©sorerie Partie 1
    # ----------------------------
    
    export_table_budget_part1 = []
    headers_part1 = df_part1.columns.tolist()
    export_table_budget_part1.append(dict(zip(["Description"] + headers_part1, [""] + headers_part1)))
    for index, row in df_part1.iterrows():
        export_table_budget_part1.append(dict(zip(["Description"] + headers_part1, [index] + row.tolist())))
    
    # Stocker les donn√©es d'exportation pour Partie 1
    st.session_state['export_data_budget_previsionnel_tresorerie_part1'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_budget_part1
    }
    
    # ----------------------------
    # Pr√©paration des donn√©es d'exportation pour Budget Pr√©visionnel de Tr√©sorerie Partie 2
    # ----------------------------
    
    export_table_budget_part2 = []
    headers_part2 = df_part2.columns.tolist()
    export_table_budget_part2.append(dict(zip(["Description"] + headers_part2, [""] + headers_part2)))
    for index, row in df_part2.iterrows():
        export_table_budget_part2.append(dict(zip(["Description"] + headers_part2, [index] + row.tolist())))
    
    # Stocker les donn√©es d'exportation pour Partie 2
    st.session_state['export_data_budget_previsionnel_tresorerie_part2'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_budget_part2
    }
    
    # ----------------------------
    # Section Export
    # ----------------------------
    
    st.header("Exporter les donn√©es")
    
    # Bouton pour t√©l√©charger le fichier Word complet contenant tous les tableaux avec une cl√© unique
    if st.button("T√©l√©charger le Document Word Complet", key="download_word_complet_budget_previsionnel_tresorerie"):
        telecharger_document_complet()

# Section 15 : Tableaux d'Analyse Financi√®re
def page_douze_tableaux():
    st.title("Tableaux d'Analyse Financi√®re")
    
    data = st.session_state["data"]
    
    st.markdown("""
    Cette section pr√©sente les principaux indicateurs financiers bas√©s sur les donn√©es que vous avez saisies.
    """)
    
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    total_salaires_annee1 = data.get("total_salaires_annee1", 0.0)
    charges_sociales_dirigeant_annee1 = data.get("charges_sociales", {}).get("dirigeants", {}).get("annee1", 0.0)
    charges_sociales_employes_annee1 = data.get("charges_sociales", {}).get("employes", {}).get("annee1", 0.0)
    amortissements_annee1 = data.get("amortissements", {}).get("total", {}).get("annee1", 0.0)
    
    # Calcul du r√©sultat net
    resultat_net = total_ca_annee1 - total_charges_fixes_annee1 - total_charges_variables - total_salaires_annee1 - charges_sociales_dirigeant_annee1 - charges_sociales_employes_annee1 - amortissements_annee1
    
    # Capacit√© d'autofinancement (simplifi√©e)
    capacite_autofinancement = resultat_net + amortissements_annee1  # Les amortissements sont r√©int√©gr√©s
    
    st.write(f"**R√©sultat Net Ann√©e 1 :** {resultat_net:.2f} $")
    st.write(f"**Capacit√© d'Autofinancement Ann√©e 1 :** {capacite_autofinancement:.2f} $")
    
    # Vous pouvez r√©p√©ter les calculs pour les ann√©es 2 et 3 si n√©cessaire
    
    st.write("---")
    
    st.session_state["data"] = data
    
def load_and_split_documents(file_path):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    raw_documents = PyPDFLoader(file_path).load()
    return text_splitter.split_documents(raw_documents)

def create_faiss_db(documents):
    if not documents:
        raise ValueError("Aucun document trouv√© pour cr√©er la base de donn√©es FAISS.")
    embeddings = OpenAIEmbeddings(api_key=api_key)
    return FAISS.from_documents(documents, embeddings)

def generate_section(system_message, query, documents, combined_content, tableau_financier, business_model):
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    llm = ChatOpenAI(api_key=api_key)
    if documents:
        db = create_faiss_db(documents)
        qa_chain = ConversationalRetrievalChain.from_llm(llm, retriever=db.as_retriever(), memory=memory, verbose=True)
        combined_info = qa_chain.run({'question': query})
        full_content = combined_content + " " + combined_info + " " + query+ " "+tableau_financier
    else:
        full_content = combined_content + " " + query+ "Dans ce donn√©es o√π vous allez recuperer les informations generales de l'entreprises "+ tableau_financier+ "utiliser les donn√©es financier pour enrichir les arguments aussi sachez que le nom du projet  correspond nom de l'entreprise. Voici les autres informations √† considerer c'est les informations du business model et ca doit etre tenue compte lors de la generation:"+ business_model
    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": full_content}
        ],
        temperature=0.9
    )
    return completion.choices[0].message.content

def extract_company_name(text):
    match = re.search(r"(nom de l'entreprise est|Nom de l'entreprise|La vision de) ([\w\s]+)", text, re.IGNORECASE)
    if match:
        return match.group(2).strip()
    return "Nom de l'entreprise non trouv√©"

def generate_markdown(results):
    markdown_content = "# Business Plan\n\n"
    for sec_name, content in results.items():
        markdown_content += f"## {sec_name}\n\n"
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            if paragraph.startswith('- '):  # Points de liste
                markdown_content += f"- {paragraph[2:]}\n"
            elif re.match(r'^\d+\.\s', paragraph):  # Points num√©rot√©s
                markdown_content += f"{paragraph}\n"
            else:
                markdown_content += f"{paragraph}\n"
        markdown_content += "\n"

    return markdown_content

def convert_table_to_markdown(table_name, table_data):
    """
    Convertit les donn√©es d'une table en format Markdown.
    
    Args:
        table_name (str): Nom de la table.
        table_data (list of dict): Donn√©es de la table.
    
    Returns:
        str: Table au format Markdown.
    """
    if not table_data:
        return "Aucune donn√©e disponible."
    
    # Extraire les en-t√™tes de colonnes
    headers = list(table_data[0].keys())
    markdown_table = "| " + " | ".join(headers) + " |\n"
    markdown_table += "| " + " | ".join(['---'] * len(headers)) + " |\n"
    
    # Ajouter les lignes
    for row in table_data:
        markdown_table += "| " + " | ".join([str(row.get(header, "")) for header in headers]) + " |\n"
    
    return markdown_table

def convert_all_tables_to_markdown(tables):
    """
    Convertit toutes les tables en une seule cha√Æne de caract√®res au format Markdown.
    
    Args:
        tables (dict): Dictionnaire contenant les tables financi√®res.
    
    Returns:
        str: Toutes les tables concat√©n√©es en Markdown.
    """
    markdown = ""
    for table_name, table_data in tables.items():
        markdown += f"### {table_name}\n\n"
        markdown += convert_table_to_markdown(table_name, table_data) + "\n\n"
    return markdown


def markdown_to_word_via_text(markdown_content):
    # Cr√©er un nouveau document Word
    doc = Document()
    doc.add_heading('Business Plan', 0)

    # Diviser le contenu en lignes
    lines = markdown_content.split('\n')
    table_data = []
    inside_table = False
    plain_text_output = []  # Pour collecter le texte brut

    for line in lines:
        line = line.strip()
        if not line:
            # Si ligne vide et donn√©es de table en cours, ajouter le tableau au document
            if table_data:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell.strip()
                table_data = []
                inside_table = False
            continue

        if line.startswith('## '):
            # Sous-titre
            doc.add_heading(line[3:], level=2)
            plain_text_output.append(line[3:])
        elif line.startswith('- '):
            # Liste √† puces
            doc.add_paragraph(line[2:], style='List Bullet')
            plain_text_output.append(f"‚Ä¢ {line[2:]}")
        elif re.match(r'^\d+\.\s', line):
            # Liste num√©rot√©e
            doc.add_paragraph(line, style='List Number')
            plain_text_output.append(line)
        elif line.startswith('|'):
            # D√©tection des lignes de tableau (√©vite les lignes de s√©paration)
            if re.match(r'\|?\s*[-:]+\s*\|', line):
                inside_table = True
                continue  # Ignorer les lignes de s√©paration
            else:
                inside_table = True
                table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
        elif line.startswith('**') and line.endswith('**'):
            # Texte en gras
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:-2])
            run.bold = True
            plain_text_output.append(line[2:-2])
        elif not inside_table:
            # Paragraphe normal
            doc.add_paragraph(line)
            plain_text_output.append(line)

    # Traiter les donn√©es de table restantes
    if table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell.strip()

    # Sauvegarder le document dans un buffer m√©moire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return "\n".join(plain_text_output), buffer

# Fonction pour convertir un dictionnaire en texte format√©
def format_table_data(data, title):
    if not data:
        return f"### {title}\nAucune donn√©e disponible pour cette section.\n\n"
    
    text = f"### {title}\n\n"
    
    # Formatage des donn√©es en tableau markdown si possible
    if isinstance(data, dict) and any(isinstance(v, dict) for v in data.values()):
        # Essayer de cr√©er un tableau markdown
        first_item = next((v for v in data.values() if isinstance(v, dict)), None)
        if first_item:
            headers = list(first_item.keys())
            text += "| " + " | ".join(headers) + " |\n"
            text += "|" + "---|" * len(headers) + "\n"
            
            for key, value in data.items():
                if isinstance(value, dict):
                    row_values = [str(value.get(h, '')) for h in headers]
                    text += "| " + " | ".join(row_values) + " |\n"
    else:
        # Format simple
        for key, value in data.items():
            if isinstance(value, dict):
                text += f"**{key}** :\n"
                for sub_key, sub_value in value.items():
                    text += f"  - {sub_key} : {sub_value}\n"
            elif isinstance(value, list):
                text += f"**{key}** : {', '.join(map(str, value))}\n"
            else:
                text += f"**{key}** : {value}\n"
    
    # Ajouter une analyse contextuelle sp√©cifique √† la RDC
    analysis = generate_rdc_analysis(title, data)
    text += f"\n**Analyse pour la RDC :** {analysis}\n\n"
    
    return text

def generate_rdc_analysis(title, data):
    """
    G√©n√®re une analyse sp√©cifique au contexte √©conomique de la RDC.
    """
    analyses_rdc = {
        "Investissements et financements": "Dans le contexte congolais, il est crucial de diversifier les sources de financement (banques locales, institutions de microfinance, partenaires internationaux). Consid√©rer l'impact du taux de change USD/FC sur les investissements.",
        
        "Salaires et Charges Sociales": "Les charges sociales en RDC incluent l'INSS, l'IPR et autres taxes. Il faut tenir compte du salaire minimum interprofessionnel garanti (SMIG) et des sp√©cificit√©s du march√© du travail local.",
        
        "Compte de r√©sultats pr√©visionnel": "Dans l'√©conomie congolaise, int√©grer les fluctuations mon√©taires, la saisonnalit√© des activit√©s et l'impact des infrastructures sur les co√ªts op√©rationnels.",
        
        "Seuil de rentabilit√© √©conomique": "En RDC, le seuil doit tenir compte des contraintes logistiques, des co√ªts √©nerg√©tiques √©lev√©s et de la volatilit√© du march√© local.",
        
        "Besoin en fonds de roulement": "Consid√©rer les d√©lais de paiement clients souvent allong√©s en RDC et pr√©voir une tr√©sorerie suffisante pour les variations saisonni√®res.",
        
        "Plan de financement √† trois ans": "Prendre en compte l'inflation, l'√©volution du taux de change et les opportunit√©s de financement via les institutions congolaises et internationales.",
        
        "Budget pr√©visionnel de tr√©sorerie": "Int√©grer les sp√©cificit√©s du syst√®me bancaire congolais, les co√ªts de transfert et la n√©cessit√© de maintenir des r√©serves pour les impr√©vus."
    }
    
    return analyses_rdc.get(title, "Adapter cette analyse aux r√©alit√©s √©conomiques de la RDC : infrastructure, r√©glementation locale, march√©s et comportements des consommateurs congolais.")

def page_generation_business_plan():
    st.title("G√©n√©rateur de Business Plan")

    uploaded_file = st.file_uploader("T√©l√©chargez votre fichier PDF", type="pdf")
    user_text_input = st.text_area("Entrez des informations suppl√©mentaires ou un texte alternatif:", height=200)

    if uploaded_file or user_text_input:
        documents = []
        combined_content = user_text_input  

        if uploaded_file:
            file_path = "uploaded_document.pdf"
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())
            documents = load_and_split_documents(file_path)    

        # Cr√©er un dictionnaire pour stocker les r√©sultats
        results = {}
        
        # Messages syst√®me et requ√™tes pour chaque section
        system_messages = {
            "Couverture": """
                G√©n√©rer cette section du business plan:
                Voici les textes √† afficher sous forme :
                
                # Canevas de Plans d‚ÄôAffaires

                Nom du projet ou entreprise
                
                 

            """,
            "Sommaire": """
                G√©n√©rer cette section du business plan:
                Voici les textes √† afficher sous forme de liste:
                ## Sommaire
                I. R√©sum√© Ex√©cutif ¬´ Executive Summary ¬ª / Pitch
                II. Pr√©sentation de votre entreprise/projet
                III. Pr√©sentation de l‚Äôoffre de produit(s) et/ou service(s)  
                IV. √âtude de march√©
                V. Strat√©gie marketing, communication et politique commerciale
                VI. Moyens de production et organisation 
                VII. √âtude des risques/hypoth√®ses  
                VIII. Plan financier 
                
            """,
            "R√©sum√© Ex√©cutif": """
                G√©n√©rer cette section du business plan:
                
                ## I. R√©sum√© Ex√©cutif ¬´ Executive Summary ¬ª / Pitch
                G√©n√©rer deux grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                Attirer l'attention du lecteur en 5 minutes et lui donner envie d'en savoir plus.
                D√©crire le projet en quelques phrases simples et impactantes.
                Ne pas essayer de tout couvrir, soyez concis et pr√©cis.

                Les elements cl√©s √† generer et qui doivent etre contenue dans les paragraphes:
                - **Pr√©sentation de la PME** : Nom de l‚Äôentreprise et br√®ve description du service/produit fourni.
                - **Pr√©sentation des porteurs de projet** : Nom, pr√©nom, coordonn√©es, situation de famille, formation et dipl√¥mes, exp√©rience professionnelle, activit√©s extra ou para-professionnelles (Joindre CV en annexe).
                - **Potentiel en termes de taille et de profit** : D√©montrez par des calculs simples comment votre PME fera du profit.
                - **Votre besoin financier**.

            """,
            "Pr√©sentation de votre entreprise": """
                G√©n√©rer cette section du business plan:

                ## II. Pr√©sentation de votre entreprise/projet

                G√©n√©rer 6 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Parler de votre entreprise/projet de mani√®re plus d√©taill√©e.
                - Pr√©senter l‚Äô√©quipe manag√©riale cl√©.

                Les elements cl√©s √† generer et qui doivent etre contenue dans les paragraphes:
                - **Informations g√©n√©rales sur la PME** :
                - Forme juridique : Ets, Sarlu, Sarl, SAS, SA.
                - Si√®ge social : Adresse juridique de l‚Äôentreprise.
                - Coordonn√©es bancaires : Num√©ro de compte (avec 23 chiffres) de l‚Äôentreprise ainsi que la banque o√π est log√© le compte (joindre le Swift Copy).
                - Couverture g√©ographique de l‚Äôentreprise et ses activit√©s : lieu d‚Äôimplantation de l‚Äôentreprise et diff√©rentes zones couvertes.
                - **Description d√©taill√©e de la PME et objectifs de son projet** : Pr√©sentez l‚Äôentreprise, son origine, introduisez ses atouts/opportunit√©s et enfin d√©crivez le projet de l‚Äôentreprise.
                - **Stade d‚Äôavancement de l‚Äôentreprise ou du projet** :
                - D√©crivez ce qui a √©t√© fait et les projets √† mener dans le futur.
                - Parlez du niveau de maturit√© de la PME ou du projet.
                - Lister √©ventuellement les financements d√©j√† acquis.
                - **Pr√©sentation de l‚Äô√©quipe manag√©riale** : D√©crivez l‚Äôorganigramme et l‚Äôorganisation des ressources humaines, pr√©sentez les associ√©s de la PME ainsi que leurs parts sociales.
                - **Analyse SWOT** : Forces, faiblesses, opportunit√©s, contraintes/menaces. de preference ca doit etre presenter sous forme de tableau.
                - **Business Mod√®le Canevas** : Ins√©rer votre business mod√®le canevas avec les 9 rubriques bien remplies.

            """,
            "Pr√©sentation de l‚Äôoffre de produit": """
                G√©n√©rer cette section du business plan :

                ## III. Pr√©sentation de l‚Äôoffre de produit(s) et/ou service(s)
                G√©n√©rer 6 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Parler de l‚Äôoffre de produits/services de mani√®re d√©taill√©e.
                - Pr√©senter la proposition de valeur diff√©renciante de la PME ou de son offre.

                Les elements cl√©s √† generer et qui doivent etre contenue dans les paragraphes:
                - **Noms du/des produit(s) ou service(s)**.
                - **Besoins identifi√©s** sur le march√© auxquels r√©pond votre offre.
                - **Description du/des produit(s) ou service(s)** r√©pondant √† ces besoins.
                - **Proposition de valeur unique**.
                - **Prise en compte de l‚Äôaspect genre** dans le fonctionnement de la PME ou du projet de l‚Äôentreprise.
                - **Prise en compte de l‚Äôenvironnement** :
                - Identification des impacts environnementaux et sociaux des activit√©s de la PME.
                - Mise en place de mesures d‚Äôatt√©nuation.
                - Existence d‚Äôun Plan de Gestion Environnemental et Social.

            """,
            "√âtude de march√©": """
                G√©n√©rer cette section du business plan :

                ## IV. √âtude de march√©

                G√©n√©rer 8 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Expliquer la m√©thode utilis√©e pour la conduite de l‚Äô√©tude de march√©.
                - Pr√©senter les r√©sultats de l‚Äô√©tude de march√©.

                Les elements cl√©s √† generer et qui doivent etre contenue dans les paragraphes, les numeros doivent etre respecter:
                1. **Description des hypoth√®ses et m√©thodes de l‚Äô√©tude de march√©** :
                - Citer le produit ou service pr√©-cibl√©.
                - Pr√©ciser le march√© pr√©-cibl√© : secteur d‚Äôactivit√© dans lequel le produit s‚Äôinscrit.
                - Pr√©senter les m√©thodes choisies pour r√©aliser l‚Äô√©tude de march√© : questionnaire, √©tude documentaire, √©tude de concurrence, √©tude m√©tier, etc.

                2. **Approche g√©n√©rale du march√© (pr√©cisez les sources √† chaque √©tape)** :
                - D√©crire le march√©, ses principales caract√©ristiques, historique et perspectives.
                - Pr√©senter les r√©sultats : march√© cible, march√© potentiel, march√© r√©el.
                - Pr√©senter les menaces et opportunit√©s du march√©.

                3. **Caract√©ristiques de la demande** :
                - Pr√©senter le volume de la demande, l‚Äô√©volution de la demande sur le march√© cibl√© et les tendances de consommation.
                - D√©tailler les diff√©rents types de client√®le (segmentation).
                - Lister les prescripteurs (partenaires qui peuvent apporter des clients).

                4. **Caract√©ristiques de l‚Äôoffre** :
                - Pr√©senter la concurrence directe et indirecte : lister les concurrents et d√©crire leur offre de services/produits.
                - Lister les points forts et les points faibles de la concurrence : avantages concurrentiels de la concurrence sur le march√©.
                - Comment vous diff√©renciez-vous de ces concurrents indirects ?

                5. **Caract√©ristiques de l‚Äôenvironnement** :
                - D√©crire l‚Äôenvironnement des affaires relatif au d√©veloppement de la PME/projet : le cadre l√©gal, r√©glementaire, les facteurs externes au march√© lui-m√™me, l‚Äô√©volution des technologies.
                - Lister les menaces et opportunit√©s li√©es √† l‚Äôenvironnement.

                6. **Partenariats** :
                - Pr√©ciser les partenariats strat√©giques nou√©s ou √† mettre en place pour faire cro√Ætre l‚Äôentreprise : il peut s‚Äôagir des acteurs en amont et en aval de votre cha√Æne de production/distribution (fournisseurs, distributeurs, partenaires commerciaux, etc.).

                7. **Cr√©ation d‚Äôemplois** :
                - D√©montrer l‚Äôimpact de la PME/projet en termes d‚Äôemplois directs d√©j√† cr√©√©s ou √† cr√©er.

                8. **Chiffre d‚Äôaffaires** :
                - Pr√©ciser la part de march√© vis√©e et le volume de chiffre d‚Äôaffaires pr√©visible √† horizon 1 an, 2 ans, 3 ans.

            """,
            "Strat√©gie Marketing":  """
                G√©n√©rer cette section du business plan :

                ## V. Strat√©gie Marketing, Communication et Politique Commerciale

                G√©n√©rer cette section, l'objectif pour cette section est de :
                - Pr√©senter la strat√©gie marketing et commerciale √† court et moyen terme.

                Les elements cl√©s √† generer et qui doivent etre contenue dans les paragraphes, les numeros doivent etre respecter:
                1. **Choix de segments de client√®le** :
                - Expliquer quels segments de client√®le vont constituer la cible de la PME/projet et pourquoi ce choix.
                - Expliquer dans les grandes lignes le positionnement strat√©gique.

                2. **Marketing-mix (4P : Produit ‚Äì Prix ‚Äì Place ‚Äì Promotion)** :
                - Pr√©senter la politique marketing g√©n√©rale :
                    - Choix du nom, du logo et des couleurs.
                    - Choix du message, du slogan.
                - Tableau synth√©tique des segments :

                    | Segment de client√®le | Produit propos√© | Positionnement en termes de prix | Lieu de distribution | Style et mode de communication |
                    |-----------------------|-----------------|----------------------------------|-----------------------|---------------------------------|
                    | Segment 1            |                 |                                  |                       |                                 |
                    | Segment 2            |                 |                                  |                       |                                 |
                    | Segment 3            |                 |                                  |                       |                                 |

                3. **Plan Marketing et actions commerciales**  :
                - Pr√©senter le plan marketing : lister les actions commerciales et actions de communication pr√©vues ; inscrire leur co√ªt si possible.

                    | Types d‚Äôactions       | Janvier | F√©vrier | Mars | ... | D√©cembre |
                    |-----------------------|---------|---------|------|-----|----------|
                    | Action 1             |         |         |      |     |          |
                    | Action 2             |         |         |      |     |          |

                4. **Moyens et partenaires sollicit√©s** :
                - Lister les moyens √† mettre en ≈ìuvre et les partenaires sollicit√©s pour les actions commerciales et de communication.

            """,
            "Moyens de production et organisation": """
                G√©n√©rer cette section du business plan:

                ## VI. Moyens de production et organisation

                G√©n√©rer 4 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Sp√©cifier les moyens humains et mat√©riels √† disposition de la PME.

                Les elements cl√©s √† generer et qui doivent etre contenue dans les paragraphes:
                - **Locaux** :
                - Liste des locaux, bail de location, conditions n√©goci√©es, co√ªt, utilit√©.
                - **Mat√©riel** :
                - Liste, mode d‚Äôacquisition ou de location, co√ªt, utilit√©, renouvellement.
                - **Moyens humains** :
                - Personnel, plannings, horaires, co√ªt, charges sociales ; indiquer une r√©partition claire des t√¢ches.
                - **Fournisseurs et sous-traitants** :
                - Liste des fournisseurs et/ou sous-traitants, devis obtenus, tarifs, conditions n√©goci√©es.

            """,
            "√âtude des risques": """
                G√©n√©rer cette section du business plan:

                ## VII. √âtude des risques/hypoth√®ses

                G√©n√©rer cette section, l'objectif pour cette section est de :
                - Pr√©senter la synth√®se des risques et mesures d‚Äôatt√©nuation identifi√©s quant au d√©veloppement de la PME/projet.

                Les elements cl√©s √† generer et qui doivent etre contenue dans les paragraphes:
                - **Tableau des risques** :

                | Nature de risque          | Description              | Strat√©gie de traitement    |
                |---------------------------|--------------------------|----------------------------|
                | Risques li√©s √† l‚Äôenvironnement g√©n√©ral |                          |                            |
                | Risques li√©s au march√©    |                          |                            |
                | Risques li√©s aux outils   |                          |                            |
                | Risques li√©s aux personnes |                          |                            |
                | Risques li√©s aux tiers    |                          |                            |
                | Autres risques (sp√©cifiez) |                          |                            |

                √âtude des risques/hypoth√®ses:

            """,
            "Annexes": """
                G√©n√©rer cette section du business plan:
                
                ## VII. √âtude des risques/hypoth√®ses

                ### Objectif
                - Pr√©senter la synth√®se des risques et mesures d‚Äôatt√©nuation identifi√©s quant au d√©veloppement de la PME/projet.

                ### Contenu attendu
                - **Tableau des risques** :

                | Nature de risque          | Description              | Strat√©gie de traitement    |
                |---------------------------|--------------------------|----------------------------|
                | Risques li√©s √† l‚Äôenvironnement g√©n√©ral |                          |                            |
                | Risques li√©s au march√©    |                          |                            |
                | Risques li√©s aux outils   |                          |                            |
                | Risques li√©s aux personnes |                          |                            |
                | Risques li√©s aux tiers    |                          |                            |
                | Autres risques (sp√©cifiez) |                          |                            |

            """,
            "Annexes": """
                G√©n√©rer cette section du business plan:

                7 ‚Äì ANNEXES
                Renvoyer en annexe les documents trop volumineux ou difficiles √† lire : - - - -
                √©tude de march√© compl√®te,
                contrats,
                conditions

                Annexes du projet:

            """,
            "Plan financier": """
                G√©n√©rer cette section du business plan:

                ## VIII. Plan financier

                G√©n√©rer cette section avec des analyses d√©taill√©es, l'objectif pour cette section est de :
                - Pr√©senter l'ensemble des √©l√©ments financiers du projet avec des analyses approfondies
                - D√©montrer la viabilit√© financi√®re et la rentabilit√© du projet
                - Fournir des conclusions strat√©giques bas√©es sur les donn√©es financi√®res

                ### Contenu √† g√©n√©rer :

                1. **Introduction g√©n√©rale** :
                - Pr√©senter l'importance de l'analyse financi√®re pour la viabilit√© du projet
                - Expliquer la m√©thodologie utilis√©e pour les projections

                2. **Analyse des investissements et financements** :
                - Analyser les besoins en investissements initiaux
                - Commenter la structure de financement propos√©e
                - √âvaluer l'ad√©quation entre besoins et ressources

                3. **Analyse des charges et de la rentabilit√©** :
                - Analyser l'√©volution pr√©visionnelle des charges (salaires, amortissements)
                - Commenter la structure des co√ªts et leur optimisation
                - √âvaluer les marges et la rentabilit√© progressive

                4. **Analyse des soldes interm√©diaires de gestion** :
                - Interpr√©ter les indicateurs cl√©s (CA, valeur ajout√©e, EBE, r√©sultat)
                - Analyser l'√©volution de la performance sur 3 ans
                - Identifier les leviers d'am√©lioration

                5. **Analyse de la capacit√© d'autofinancement et du seuil de rentabilit√©** :
                - Commenter la g√©n√©ration de cash-flow
                - Analyser le point mort et sa progression
                - √âvaluer la soutenabilit√© financi√®re

                6. **Analyse du besoin en fonds de roulement** :
                - Expliquer l'√©volution du BFR
                - Analyser l'impact sur la tr√©sorerie
                - Proposer des actions d'optimisation

                7. **Analyse du plan de financement sur 3 ans** :
                - V√©rifier l'√©quilibre emplois/ressources
                - Analyser la structure financi√®re
                - √âvaluer les besoins de financement compl√©mentaires

                8. **Analyse de la tr√©sorerie pr√©visionnelle** :
                - Commenter l'√©volution mensuelle de la tr√©sorerie
                - Identifier les p√©riodes critiques
                - Proposer des solutions de gestion de tr√©sorerie

                9. **Synth√®se et conclusions financi√®res** :
                - Synth√©tiser les principaux enseignements de l'analyse
                - √âvaluer la viabilit√© globale du projet
                - Formuler des recommandations strat√©giques
                - Identifier les facteurs cl√©s de succ√®s financier
                - Proposer des indicateurs de suivi de performance

                **Instructions sp√©ciales** :
                - Sous chaque analyse de tableau, ajouter une conclusion de 2-3 phrases
                - Cr√©er des liens logiques entre les diff√©rentes analyses
                - Terminer par une conclusion g√©n√©rale qui guide vers les prochaines √©tapes
                - Utiliser les donn√©es financi√®res concr√®tes pour √©tayer chaque analyse

            """
        }

        queries = {
            "Couverture": "Afficher seulement le texte fournies",
            "Sommaire": "Afficher seulement le texte fournises",
            "R√©sum√© Ex√©cutif": "D√©crire bri√®vement le projet, son potentiel de profit et les qualifications de l'√©quipe.",
            "Pr√©sentation de votre entreprise": "Fournir une analyse d√©taill√©e de l'entreprise, incluant son origine, ses objectifs et son organisation.",
            "Pr√©sentation de l‚Äôoffre de produit": "D√©crire les produits ou services, leur proposition de valeur unique, et les besoins du march√© qu'ils adressent.",
            "√âtude de march√©": "Analyser le march√© cible, les tendances de consommation, et la concurrence directe et indirecte.",
            "Strat√©gie Marketing": "D√©crire la strat√©gie marketing, y compris les segments cibles, le positionnement, le mix marketing (Produit, Prix, Place, Promotion) et les actions commerciales pr√©vues.",
            "Moyens de production et organisation": "D√©crire les moyens humains et mat√©riels, ainsi que l'organisation op√©rationnelle de l'entreprise.",
            "√âtude des risques": "Identifier les risques potentiels et proposer des strat√©gies pour les att√©nuer.",
            "Plan financier": "Analyser en d√©tail tous les √©l√©ments financiers du projet, commenter chaque tableau avec des conclusions, √©tablir des liens logiques entre les analyses, et formuler une synth√®se strat√©gique de la viabilit√© financi√®re.",
            "Annexes": "Inclure tous les documents annexes pertinents pour √©tayer le plan d'affaires."
        }

        # Espaces r√©serv√©s pour chaque section
        placeholders = {name: st.empty() for name in system_messages.keys()}
        
        data = st.session_state.get("data", {})
        tables = data.get("tables", {})
        
        
        
        
        section_order = list(system_messages.keys())
        # D√©finir le point de s√©paration
        split_section = "Pr√©sentation de votre entreprise"
        # S√©parer les sections en deux groupes
        first_part = []
        second_part = []
        for section in section_order:
            if section == split_section:
                first_part.append(section)
                second_part = section_order[section_order.index(section)+1:]
                break
            else:
                first_part.append(section)

        # R√©cup√©rer toutes les donn√©es des √©tapes pr√©c√©dentes
        business_model_precedent = st.session_state.get('business_model_precedent', '')
        persona_data = st.session_state.get('persona_data', {})
        marche_data = st.session_state.get('marche_data', {})
        concurrence_data = st.session_state.get('concurrence_data', {})
        facteurs_limitants = st.session_state.get('facteurs_limitants', {})
        
        # Construire un contexte enrichi avec toutes les donn√©es pr√©c√©dentes
        contexte_complet = f"""
        DONN√âES DU BUSINESS MODEL: {business_model_precedent}
        
        INFORMATIONS PERSONA: {persona_data}
        
        ANALYSE DE MARCH√â: {marche_data}
        
        ANALYSE CONCURRENCE: {concurrence_data}
        
        FACTEURS LIMITANTS: {facteurs_limitants}
        """
        
        results_first_part = {}
        results_second_part = {}
        
        # R√©cup√©rer les donn√©es export√©es de toutes les sections
        # R√©cup√©rer les donn√©es export√©es de toutes les sections 
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
        export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
        export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
        export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
        export_data_budget_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
        export_data_budget_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})

        # Concat√©ner toutes les sections
        final_text = ""
        final_text += format_table_data(export_data_investissements, "Investissements et financements")
        final_text += format_table_data(export_data_salaires, "Salaires et Charges Sociales")
        final_text += format_table_data(export_data_amortissements, "D√©tail des Amortissements")
        final_text += format_table_data(export_data_compte, "Compte de r√©sultats pr√©visionnel")
        final_text += format_table_data(export_data_soldes, "Soldes interm√©diaires de gestion")
        final_text += format_table_data(export_data_capacite, "Capacit√© d'autofinancement")
        final_text += format_table_data(export_data_seuil, "Seuil de rentabilit√© √©conomique")
        final_text += format_table_data(export_data_bfr, "Besoin en fonds de roulement")

        # Ajouter les nouvelles sections
        final_text += format_table_data(export_data_plan_financement, "Plan de financement √† trois ans")
        final_text += format_table_data(export_data_budget_part1, "Budget pr√©visionnel de tr√©sorerie")
        final_text += format_table_data(export_data_budget_part2, "Budget pr√©visionnel de tr√©sorerie(suite)")

        

        # G√©n√©rer toutes les sections automatiquement
       # G√©n√©ration du Business Plan et t√©l√©chargement des fichiers      
        for section_name in first_part:
            with st.spinner(f"G√©n√©ration de {section_name}..."):
                system_message = system_messages[section_name]
                query = queries[section_name]
                
                try:
                    # V√©rifier si la section est "Couverture" ou "Sommaire"
                    if section_name in ["Couverture", "Sommaire"]:
                        results_first_part[section_name] = generate_section(system_message, query, documents, combined_content, "", business_model="")
                    else:
                        results_first_part[section_name] = generate_section(system_message, query, documents, combined_content + contexte_complet, final_text, business_model=contexte_complet)
                except ValueError as e:
                    results_first_part[section_name] = f"Erreur: {str(e)}"
                
                combined_content += " " + results_first_part[section_name]
                placeholders[section_name].markdown(f"\n\n### {section_name}\n{results_first_part[section_name]}")


        # Extraction du nom de l'entreprise
        #first_section_content = results.get("R√©sum√© Ex√©cutif", "")
        #company_name = extract_company_name(first_section_content)

        # Cr√©er le contenu Markdown principal
        markdown_content1 = generate_markdown(results_first_part)
        
        #content_result, word_buffer = markdown_to_word_via_text(arkdown_contentm)

        # Ajouter les tableaux au contenu Markdown
        #  markdown_content = append_tables_to_markdown(markdown_content)
        
        
        # G√©n√©ration de la seconde partie
        for section_name in second_part:
            with st.spinner(f"G√©n√©ration de {section_name}..."):
                system_message = system_messages[section_name]
                query = queries[section_name]
                
                try:
                    results_second_part[section_name] = generate_section(system_message, query, documents, combined_content + contexte_complet, final_text, business_model=contexte_complet)
                except ValueError as e:
                    results_second_part[section_name] = f"Erreur: {str(e)}"
                
                combined_content += " " + results_second_part[section_name]
                placeholders[section_name].markdown(f"\n\n### {section_name}\n{results_second_part[section_name]}")

        
        
        markdown_content2 = generate_markdown(results_second_part)

        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(markdown_content2))
        pdf.meta["title"] = "Business Plan" 
        pdf_file_path = "business_plan.pdf"
        pdf.save(pdf_file_path)
        
        
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement


        from docx import Document
        import re

        def markdown_to_word_via_text(markdown_contents, doc):
            doc.add_heading('Business Plan', 0)

            # Diviser le contenu en lignes
            lines = markdown_contents.split('\n')
            table_data = []
            inside_table = False

            for line in lines:
                line = line.strip()
                if not line:
                    # Si ligne vide et donn√©es de table en cours, ajouter le tableau au document
                    if table_data:
                        add_table_with_borders(doc, table_data)
                        table_data = []
                        inside_table = False
                    continue

                if line.startswith('# '):  # Titre niveau 1
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):  # Titre niveau 2
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):  # Titre niveau 3
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):  # Titre niveau 4
                    doc.add_heading(line[5:], level=4)
                elif re.match(r'^\d+\.\s', line):  # Liste num√©rot√©e
                    # V√©rifier s'il y a du texte en gras dans la liste num√©rot√©e
                    match = re.match(r'^(\d+\.\s)(\*\*.+?\*\*)', line)
                    if match:
                        paragraph = doc.add_paragraph(style='List Number')
                        paragraph.add_run(match.group(1))  # Num√©ro
                        bold_run = paragraph.add_run(match.group(2)[2:-2])  # Texte en gras sans `**`
                        bold_run.bold = True
                    else:
                        doc.add_paragraph(line, style='List Number')
                elif line.startswith('- ') or line.startswith('‚Ä¢'):  # Liste √† puces
                    match = re.match(r'^(‚Ä¢|-)\s\*\*(.+?)\*\*(.*)', line)
                    if match:
                        paragraph = doc.add_paragraph(style='List Bullet')
                        bold_run = paragraph.add_run(match.group(2))  # Texte en gras
                        bold_run.bold = True
                        if match.group(3):  # Texte apr√®s le gras
                            paragraph.add_run(match.group(3).strip())
                    else:
                        doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('|'):  # D√©tection des lignes de tableau
                    if re.match(r'\|?\s*[-:]+\s*\|', line):
                        inside_table = True
                        continue  # Ignorer les lignes de s√©paration
                    else:
                        inside_table = True
                        table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
                elif re.match(r'^\*\*.+?\*\*\s*:', line):  # Texte en gras suivi de texte normal
                    match = re.match(r'^\*\*(.+?)\*\*\s*:(.*)', line)
                    if match:
                        paragraph = doc.add_paragraph()
                        bold_run = paragraph.add_run(match.group(1))  # Texte en gras
                        bold_run.bold = True
                        if match.group(2):  # Texte normal apr√®s le `:`
                            paragraph.add_run(f":{match.group(2)}")
                elif re.match(r'^\*\*.+?\*\*$', line):  # Texte enti√®rement en gras
                    paragraph = doc.add_paragraph()
                    bold_run = paragraph.add_run(line[2:-2])  # Texte sans `**`
                    bold_run.bold = True
                elif re.match(r'^\*\*.+?\*\*\s[\d.,]+\s?[$$%]$', line):  # Nombres avec symboles mon√©taires
                    match = re.match(r'^\*\*(.+?)\*\*\s([\d.,]+\s?[$$%])$', line)
                    if match:
                        paragraph = doc.add_paragraph()
                        bold_run = paragraph.add_run(match.group(1))  # Texte en gras
                        bold_run.bold = True
                        paragraph.add_run(f" {match.group(2)}")  # Montant avec symbole
                elif not inside_table:  # Paragraphe normal
                    doc.add_paragraph(line)

            # Traiter les donn√©es de table restantes
            if table_data:
                add_table_with_borders(doc, table_data)

        def add_table_with_borders(doc, table_data):
            """
            Ajoute un tableau au document Word avec bordures et gestion du texte en gras dans les cellules.
            """
            num_cols = len(table_data[0])
            table = doc.add_table(rows=len(table_data), cols=num_cols)
            table.style = 'Table Grid'  # Appliquer un style de tableau avec bordures

            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    cell_content = table.cell(i, j).paragraphs[0]
                    parts = re.split(r'(\*\*.+?\*\*)', cell)  # Diviser par texte en gras
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):  # Texte en gras
                            run = cell_content.add_run(part[2:-2])
                            run.bold = True
                        else:  # Texte normal
                            cell_content.add_run(part.strip())




        # Ajouter la cr√©ation et le t√©l√©chargement du fichier Word
        
        doc = Document()
        markdown_to_word_via_text(markdown_content1, doc)
        generer_docx_business_model(nom_entreprise, datetime.date.today(), st.session_state.business_model_precedent, doc, value=2)
        markdown_to_word_via_text(markdown_content2, doc)
        # doc.add_paragraph(content_result)
        
        # V√©rifier et ajouter le contenu
       

        # Ajouter les sections du Business Plan
        """ for section_name, content in results.items():
            doc.add_heading(section_name, level=1)
            doc.add_paragraph(content)"""

        # R√©cup√©rer les donn√©es des tableaux depuis la session Streamlit
       # R√©cup√©rer les donn√©es des tableaux depuis la session Streamlit
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
        export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
        export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})

        # Ajouter une section pour les tableaux
        doc.add_heading('R√©sum√© des Donn√©es Financi√®res', level=1)

        # Fonction pour ajouter un tableau dans le document Word
        def ajouter_tableau(donnees, headers, titre):
            """
            Ajoute un tableau au document Word avec bordures et gestion du texte en gras dans les cellules.
            """
            doc.add_heading(titre, level=2)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light List Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Ajouter les en-t√™tes
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Ajouter les donn√©es des tableaux
            for row in donnees:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    cell_value = row.get(header, "")
                    cell_text = str(cell_value)  # Convertir en cha√Æne de caract√®res
                    row_cells[i].text = cell_text
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Ajouter une note
            doc.add_paragraph()
            doc.add_paragraph("Les r√©sultats sont calcul√©s selon les donn√©es fournies.")


        # Ajouter les diff√©rents tableaux
        if export_data_investissements.get("table_data"):
            ajouter_tableau(export_data_investissements["table_data"], ["Investissements", "Taux (%)", "Dur√©e (mois)", "Montant ($)"], "Investissements et Financements")
        if export_data_salaires.get("table_data"):
            ajouter_tableau(export_data_salaires["table_data"], ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"], "Salaires et Charges Sociales")
        if export_data_amortissements.get("amortissements"):
            ajouter_tableau(export_data_amortissements["amortissements"], ["Amortissement", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"], "D√©tail des Amortissements")
        if export_data_compte.get("table_data"):
            ajouter_tableau(export_data_compte["table_data"], ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"], "Compte de R√©sultats Pr√©visionnel")
        if export_data_soldes.get("table_data"):
            ajouter_tableau(export_data_soldes["table_data"], ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"], "Soldes Interm√©diaires de Gestion")
        if export_data_capacite.get("table_data"):
            ajouter_tableau(export_data_capacite["table_data"], ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"], "Capacit√© d'Autofinancement")
        if export_data_seuil.get("table_data"):
            ajouter_tableau(export_data_seuil["table_data"], ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"], "Seuil de Rentabilit√© √âconomique")
        if export_data_bfr.get("table_data"):
            ajouter_tableau(export_data_bfr["table_data"], ["Analyse clients / fournisseurs", "D√©lai jours", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"], "Besoin en Fonds de Roulement")

        # **Nouvelles sections ajout√©es¬†:**

        # Ajouter la section Plan de Financement √† Trois Ans
        export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
        if export_data_plan_financement.get("table_data"):
            ajouter_tableau(
                export_data_plan_financement["table_data"],
                ["Description", "Ann√©e 1", "Ann√©e 2", "Ann√©e 3"],
                "Plan de Financement √† Trois Ans"
            )

        # Ajouter la section Budget Pr√©visionnel de Tr√©sorerie Partie 1
        export_data_budget_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
        if export_data_budget_part1.get("table_data"):
            ajouter_tableau(
                export_data_budget_part1["table_data"],
                ["Description", "Mois 1", "Mois 2", "Mois 3", "Mois 4", "Mois 5", "TOTAL"],
                "Budget Pr√©visionnel de Tr√©sorerie - Partie 1"
            )

        # Ajouter la section Budget Pr√©visionnel de Tr√©sorerie Partie 2
        export_data_budget_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})
        if export_data_budget_part2.get("table_data"):
            ajouter_tableau(
                export_data_budget_part2["table_data"],
                ["Description", "Mois 6", "Mois 7", "Mois 8", "Mois 9", "Mois 10", "Mois 11", "Mois 12", "TOTAL"],
                "Budget Pr√©visionnel de Tr√©sorerie - Partie 2"
            )

        # Enregistrer le document dans un buffer
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        # T√©l√©charger les fichiers g√©n√©r√©s
        st.success("Le PDF et le document Word ont √©t√© g√©n√©r√©s avec succ√®s.")
        with open(pdf_file_path, "rb") as f:
            st.download_button("T√©l√©chargez le PDF", f, file_name="business_plan.pdf", mime="application/pdf")

        st.download_button("T√©l√©chargez le document Word", word_buffer, file_name="business_plan.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# S√©lection du type d'entreprise et nom

st.title('Business Plan')
st.sidebar.header("Configuration Initiale pour le business model")
type_entreprise = st.sidebar.selectbox("Type d'entreprise", ["PME", "Startup"], key="type_entreprise")
nom_entreprise = st.sidebar.text_input("Nom de l'entreprise", value="", key="nom_entreprise")

if not nom_entreprise:
    st.sidebar.warning("Veuillez entrer le nom de votre entreprise.")


# Initialiser les variables dans la session si ce n'est pas d√©j√† fait
if 'business_model_precedent' not in st.session_state:
    st.session_state.business_model_precedent = ""

# Noms des nouveaux onglets du Business Model Canvas
business_model_tab_names = [
    "Collecte des Donn√©es",
    "G√©n√©rer Business Model",
]

# Fonctions correspondantes pour les nouveaux onglets
business_model_sections = [
    page_collecte_donnees,
    page_generer_business_model,
]

# Liste des noms d'onglets existants

# Mise √† jour des noms d'onglets
tab_names = [
    "Informations G√©n√©rales", "Besoins de D√©marrage", "Financement",
    "Charges Fixes", "Chiffre d'Affaires", "Charges Variables",
    "Fonds de Roulement", "Salaires", "Rentabilit√©", "Tr√©sorerie","G√©n√©ration du Business Plan"
]

# Mise √† jour de la liste des fonctions correspondantes
sections = [
    page_informations_generales, page_besoins_demarrage, page_financement,
    page_charges_fixes, page_chiffre_affaires, page_charges_variables,
    page_fonds_roulement, page_salaires, page_rentabilite, page_tresorerie,
    page_generation_business_plan
]

# Trouver l'index de "G√©n√©ration du Business Plan"
try:
    index_generation_bp = tab_names.index("G√©n√©ration du Business Plan")
except ValueError:
    st.error("L'onglet 'G√©n√©ration du Business Plan' n'a pas √©t√© trouv√© dans la liste des onglets.")
    index_generation_bp = len(tab_names)  # Ajouter √† la fin si non trouv√©

# Ins√©rer les nouveaux onglets avant "G√©n√©ration du Business Plan"
tab_names =business_model_tab_names + tab_names
sections = business_model_sections + sections


# Cr√©ation des onglets
tabs = st.tabs(tab_names)

# Parcours des onglets
for i, tab in enumerate(tabs):
    with tab:
        sections[i]()


