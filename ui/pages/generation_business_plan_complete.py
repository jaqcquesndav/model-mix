"""
Service de génération de business plan avec intégration des tableaux financiers
Version cyclique adaptée d'Origin.txt avec système de templates professionnels
"""

import streamlit as st
from typing import Dict, Any, List
from services.ai.content_generation import generate_section, tester_connexion_openai
from services.financial.calculations import calculer_tableaux_financiers_5_ans
from services.document.generation import format_table_to_markdown
from templates.business_plan_prompts import (
    get_system_messages_origin_style,
    get_queries_origin_style,
    get_sections_configuration_origin_style,
    get_template_context
)
from templates import get_metaprompt, get_system_messages  # Ancien système pour compatibilité

def generate_section_origin(system_message, query, documents, combined_content, tableau_financier, business_model):
    """
    Fonction de génération EXACTE copiée d'Origin.txt
    """
    return generate_section(
        system_message=system_message,
        user_query=query,
        additional_context=combined_content + "\n\n" + tableau_financier + "\n\n" + str(business_model),
        section_name=""
    )
import pandas as pd
import tempfile
import os

def page_generation_business_plan_integree():
    """Page de génération du business plan avec tableaux financiers intégrés - Version cyclique"""
    st.title("🎯 Générateur de Business Plan Complet")
    
    # Test de connectivité OpenAI au début (comme dans Origin.txt)
    st.markdown("### 🔌 Statut de la connexion IA")
    with st.expander("Vérifier la connexion OpenAI", expanded=False):
        if st.button("🧪 Tester la connexion API"):
            with st.spinner("Test de connexion en cours..."):
                status = tester_connexion_openai()
                
                if status["status"] == "success":
                    st.success(f"✅ {status['message']}")
                    st.info(f"📊 {status['details']}")
                elif status["status"] == "warning":
                    st.warning(f"⚠️ {status['message']}")
                    st.info(f"📋 {status['details']}")
                else:
                    st.error(f"❌ {status['message']}")
                    st.error(f"📋 {status['details']}")
                    st.stop()  # Arrêter l'exécution si l'API ne fonctionne pas
    
    # Récupération du template sélectionné
    template_actuel = st.session_state.get('template_selectionne', 'COPA TRANSFORME')
    
    st.info(f"🎨 **Template actuel :** {template_actuel}")
    st.info("📋 Cette version génère un business plan complet section par section, avec intégration automatique des tableaux financiers.")
    
    # Options de génération
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📄 Option 1: Génération avec document")
        uploaded_file = st.file_uploader("Téléchargez un document de référence (PDF)", type="pdf", key="upload_doc")
        
        if uploaded_file:
            st.success(f"✅ Document chargé: {uploaded_file.name}")
    
    with col2:
        st.markdown("### ✍️ Option 2: Génération par description")
        user_text_input = st.text_area(
            "Décrivez votre projet/entreprise:",
            height=150,
            placeholder="Décrivez votre entreprise, ses produits/services, son marché cible, ses objectifs...",
            key="user_description"
        )
    
    # Options de génération
    st.markdown("### ⚙️ Options de génération")
    col_opt1, col_opt2, col_opt3 = st.columns(3)
    
    with col_opt1:
        use_workflow_data = st.checkbox(
            "📊 Utiliser les données du workflow",
            value=True,
            help="Utilise les données saisies dans l'application (Business Model, finances, etc.)"
        )
    
    with col_opt2:
        show_progress = st.checkbox(
            "📈 Affichage en temps réel",
            value=True,
            help="Affiche les sections au fur et à mesure de leur génération"
        )
    
    with col_opt3:
        split_generation = st.checkbox(
            "🔄 Génération en deux phases",
            value=True,
            help="Sépare la génération en deux parties comme dans Origin.txt"
        )
    
    # Validation et génération
    can_generate = uploaded_file is not None or user_text_input.strip() != "" or use_workflow_data
    
    if not can_generate:
        st.warning("⚠️ Veuillez soit télécharger un document, soit saisir une description, soit cocher 'Utiliser les données du workflow'")
        return
    
    if st.button("🚀 Générer le Business Plan Complet", type="primary", disabled=not can_generate):
        generate_complete_business_plan_origin_exact(
            uploaded_file=uploaded_file,
            user_text_input=user_text_input,
            template_nom=template_actuel,
            use_workflow_data=use_workflow_data,
            show_progress=show_progress,
            split_generation=split_generation
        )

def generate_complete_business_plan_origin_exact(uploaded_file=None, user_text_input="", template_nom="COPA TRANSFORME", 
                                              use_workflow_data=True, show_progress=True, split_generation=True):
    """Génère un business plan avec la logique EXACTE d'Origin.txt adaptée pour templates RDC"""
    
    # 1. Traitement des documents (EXACT Origin.txt)
    documents = []
    combined_content = user_text_input if user_text_input else ""
    
    if uploaded_file:
        documents = process_uploaded_pdf(uploaded_file)
        if documents:
            st.success(f"✅ {len(documents)} documents PDF traités")
    
    # 2. Récupération des données financières (EXACT Origin.txt)
    business_data = collect_all_business_data() if use_workflow_data else {}
    
    # Récupérer les données exportées de toutes les sections (EXACT Origin.txt)
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
    export_data_budget_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
    export_data_budget_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})

    # Concaténer toutes les sections financières (EXACT Origin.txt)
    final_text = ""
    final_text += format_table_data_origin(export_data_investissements, "Investissements et financements")
    final_text += format_table_data_origin(export_data_salaires, "Salaires et Charges Sociales")
    final_text += format_table_data_origin(export_data_amortissements, "Détail des Amortissements")
    final_text += format_table_data_origin(export_data_compte, "Compte de résultats prévisionnel")
    final_text += format_table_data_origin(export_data_soldes, "Soldes intermédiaires de gestion")
    final_text += format_table_data_origin(export_data_capacite, "Capacité d'autofinancement")
    final_text += format_table_data_origin(export_data_seuil, "Seuil de rentabilité économique")
    final_text += format_table_data_origin(export_data_bfr, "Besoin en fonds de roulement")
    final_text += format_table_data_origin(export_data_plan_financement, "Plan de financement à trois ans")
    final_text += format_table_data_origin(export_data_budget_part1, "Budget prévisionnel de trésorerie")
    final_text += format_table_data_origin(export_data_budget_part2, "Budget prévisionnel de trésorerie(suite)")

    # 3. Configuration des sections selon template (Origin.txt + templates)
    system_messages = get_system_messages_origin_style(template_nom)
    queries = get_queries_origin_style()
    
    # 4. Espaces réservés pour affichage (EXACT Origin.txt)
    placeholders = {name: st.empty() for name in system_messages.keys()}
    
    # 5. Séparation en deux parties (EXACT Origin.txt)
    section_order = list(system_messages.keys())
    split_section = "Présentation de votre entreprise"
    
    first_part = []
    second_part = []
    for section in section_order:
        if section == split_section:
            first_part.append(section)
            second_part = section_order[section_order.index(section)+1:]
            break
        else:
            first_part.append(section)
    
    results_first_part = {}
    results_second_part = {}
    
    # 6. Génération première partie (EXACT Origin.txt)
    st.markdown("### 🔄 **Phase 1 : Sections fondamentales**")
    
    for section_name in first_part:
        with st.spinner(f"Génération de {section_name}..."):
            system_message = system_messages[section_name]
            query = queries[section_name]
            
            try:
                # Logique EXACTE d'Origin.txt
                if section_name in ["Couverture", "Sommaire"]:
                    results_first_part[section_name] = generate_section(
                        system_message=system_message, 
                        user_query=query, 
                        additional_context=combined_content,
                        section_name=section_name
                    )
                else:
                    # Récupérer le business model (EXACT Origin.txt)
                    business_model = st.session_state.get('business_model_precedent', '')
                    results_first_part[section_name] = generate_section(
                        system_message=system_message, 
                        user_query=query, 
                        additional_context=combined_content,
                        section_name=section_name,
                        financial_context=final_text,
                        business_model=business_model
                    )
            except ValueError as e:
                results_first_part[section_name] = f"Erreur: {str(e)}"
            
            # Accumulation progressive du contexte (EXACT Origin.txt)
            combined_content += " " + results_first_part[section_name]
            
            # Affichage en temps réel
            if show_progress:
                placeholders[section_name].markdown(f"\n\n### {section_name}\n{results_first_part[section_name]}")

    st.success("✅ Phase 1 terminée")
    
    # 7. Génération seconde partie (EXACT Origin.txt)
    st.markdown("### 🔄 **Phase 2 : Sections avancées**")
    
    for section_name in second_part:
        with st.spinner(f"Génération de {section_name}..."):
            system_message = system_messages[section_name]
            query = queries[section_name]
            
            try:
                business_model = st.session_state.get('business_model_precedent', '')
                results_second_part[section_name] = generate_section(
                    system_message=system_message, 
                    user_query=query, 
                    additional_context=combined_content,
                    section_name=section_name,
                    financial_context=final_text,
                    business_model=business_model
                )
            except ValueError as e:
                results_second_part[section_name] = f"Erreur: {str(e)}"
            
            # Accumulation continue (EXACT Origin.txt)
            combined_content += " " + results_second_part[section_name]
            
            # Affichage en temps réel
            if show_progress:
                placeholders[section_name].markdown(f"\n\n### {section_name}\n{results_second_part[section_name]}")

    st.success("✅ Phase 2 terminée")
    
    # 8. Combiner tous les résultats
    all_results = {**results_first_part, **results_second_part}
    
    # 9. Génération des fichiers de sortie (Origin.txt style)
    create_export_files_origin_style(all_results, business_data, template_nom)

def format_table_data_origin(export_data, section_title):
    """Formate les données de tableau dans le style Origin.txt"""
    if not export_data:
        return ""
    
    formatted_text = f"\n\n=== {section_title} ===\n"
    
    for key, value in export_data.items():
        if isinstance(value, dict):
            formatted_text += f"\n{key}:\n"
            for sub_key, sub_value in value.items():
                formatted_text += f"  {sub_key}: {sub_value}\n"
        else:
            formatted_text += f"{key}: {value}\n"
    
    return formatted_text

def create_export_files_origin_style(results: Dict[str, str], business_data: Dict[str, Any], template_nom: str):
    """Fonction d'export dans le style Origin.txt"""
    return create_export_files_cyclique(results, business_data, template_nom)

def generate_section_cyclique(section_name, section_config, documents, combined_content, 
                            financial_tables_text, business_data, results, placeholders, template_nom):
    """Génère une section individuelle avec la logique cyclique inspirée d'Origin.txt"""
    
    try:
        # Utiliser la nouvelle configuration
        if "system_message" in section_config and "user_query" in section_config:
            system_message = section_config["system_message"]
            query = section_config["user_query"]
        else:
            # Fallback
            system_message = section_config.get("system_message", "")
            query = section_config.get("query", section_config.get("user_query", ""))
        
        # Logique inspirée d'Origin.txt : contexte différent selon la section
        if section_name in ["Couverture", "Sommaire"]:
            # Sections simples comme dans Origin.txt
            content = generate_section(
                system_message=system_message,
                user_query=query,
                additional_context="",  # Pas de contexte pour éviter la confusion
                section_name=section_name
            )
        else:
            # Sections avec contexte business comme dans Origin.txt
            business_model = st.session_state.get('business_model_precedent', '')
            
            # Construire le contexte comme dans Origin.txt
            context_info = f"""Dans ces données où vous allez récupérer les informations générales de l'entreprise {financial_tables_text} utiliser les données financières pour enrichir les arguments aussi sachez que le nom du projet correspond au nom de l'entreprise. Voici les autres informations à considérer c'est les informations du business model et ça doit être tenu compte lors de la génération: {business_model}"""
            
            # Ajouter un contexte limité pour éviter les doublons
            limited_context = combined_content[-500:] if combined_content else ""  # Seulement les 500 derniers caractères
            
            full_content = limited_context + " " + query + " " + context_info
            
            content = generate_section(
                system_message=system_message,
                user_query=full_content,
                additional_context="",
                section_name=section_name
            )
        
        # Nettoyer le contenu généré
        content = clean_generated_content_origin_style(content, section_name)
        
        results[section_name] = content
        
        # Affichage en temps réel
        if placeholders and section_name in placeholders:
            placeholders[section_name].markdown(f"### {section_name}\n{content}")
            
    except Exception as e:
        error_msg = f"❌ Erreur lors de la génération de {section_name}: {str(e)}"
        st.error(f"Erreur détaillée: {str(e)}")
        
        results[section_name] = error_msg
        
        if placeholders and section_name in placeholders:
            placeholders[section_name].error(error_msg)

def clean_generated_content_origin_style(content: str, section_name: str) -> str:
    """Nettoie le contenu dans le style d'Origin.txt - Supprime les explications"""
    
    # Supprimer les phrases d'introduction de l'IA
    intro_phrases = [
        "Voici une",
        "Voici un",
        "Créer une page de couverture",
        "Pour créer",
        "Il est important de",
        "Assurez-vous de",
        "N'oubliez pas de",
        "Ce document",
        "Cette section",
        "Voici comment",
        "Dans le cadre de",
        "Pour analyser",
        "En résumé",
        "Pour conclure"
    ]
    
    lines = content.split('\n')
    cleaned_lines = []
    skip_paragraph = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # Détecter les phrases d'introduction à supprimer
        should_skip = any(phrase in line for phrase in intro_phrases)
        
        if should_skip:
            skip_paragraph = True
            continue
            
        # Arrêter de skipper quand on trouve une nouvelle section ou un contenu utile
        if skip_paragraph and (line_stripped.startswith('#') or line_stripped.startswith('**') or line_stripped.startswith('|')):
            skip_paragraph = False
        
        if not skip_paragraph:
            cleaned_lines.append(line)
    
    content = '\n'.join(cleaned_lines)
    
    # Supprimer les doublons de sections
    content = remove_section_duplicates(content)
    
    return content.strip()

def remove_section_duplicates(content: str) -> str:
    """Supprime les doublons de sections dans le contenu"""
    
    lines = content.split('\n')
    seen_headers = set()
    result_lines = []
    
    for line in lines:
        # Identifier les en-têtes de section
        if line.startswith('#') or line.startswith('**') and line.endswith('**'):
            header_clean = line.strip('#').strip('*').strip()
            
            if header_clean not in seen_headers:
                seen_headers.add(header_clean)
                result_lines.append(line)
            else:
                # Dupliquer trouvé, on ne l'ajoute pas
                continue
        else:
            result_lines.append(line)
    
    return '\n'.join(result_lines)

def collect_all_business_data() -> Dict[str, Any]:
    """Collecte toutes les données business de l'application avec logique cyclique"""
    data = st.session_state.get("data", {})
    
    return {
        "business_model": st.session_state.get('business_model_precedent', '') or st.session_state.get('business_model_initial', {}),
        "persona": st.session_state.get('persona_data', {}),
        "marche": st.session_state.get('analyse_marche', {}),
        "concurrence": st.session_state.get('concurrence', {}),
        "facteurs_limitants": st.session_state.get('facteurs_limitants_data', {}),
        "arbre_probleme": st.session_state.get('arbre_probleme', {}),
        "informations_generales": data.get("informations_generales", {}),
        "ca_previsions": data.get("ca_previsions", {}),
        "charges_variables": data.get("charges_variables", {}),
        "charges_fixes": data.get("charges_fixes", {}),
        "salaires": data.get("salaires", {}),
        "investissements": data.get("besoins_demarrage", {}),
        "financements": data.get("financements", {})
    }

def generate_all_financial_tables() -> Dict[str, Any]:
    """Génère tous les tableaux financiers et les formate pour le business plan cyclique"""
    
    data = st.session_state.get("data", {})
    
    try:
        # Utiliser notre service de calculs 5 ans
        tableaux_5_ans = calculer_tableaux_financiers_5_ans()
        
        # Formater tous les tableaux pour inclusion dans le business plan
        formatted_tables = format_financial_tables_for_business_plan_cyclique(tableaux_5_ans)
        
        return {
            "raw_data": tableaux_5_ans,
            "formatted_text": formatted_tables,
            "tables_list": list(tableaux_5_ans.keys()) if tableaux_5_ans else []
        }
        
    except Exception as e:
        st.warning(f"Erreur lors de la génération des tableaux financiers: {str(e)}")
        
        # Fallback vers les données exportées du session state (comme Origin.txt)
        return get_financial_tables_from_session()

def get_financial_tables_from_session() -> Dict[str, Any]:
    """Récupère les tableaux financiers depuis le session state (logique Origin.txt)"""
    
    # Récupérer les données exportées de toutes les sections (comme dans Origin.txt)
    export_data = {
        'investissements': st.session_state.get('export_data_investissements', {}),
        'salaires': st.session_state.get('export_data_salaires_charges_sociales', {}),
        'amortissements': st.session_state.get('export_data_detail_amortissements', {}),
        'compte': st.session_state.get('export_data_compte_resultats_previsionnel', {}),
        'soldes': st.session_state.get('export_data_soldes_intermediaires_de_gestion', {}),
        'capacite': st.session_state.get('export_data_capacite_autofinancement', {}),
        'seuil': st.session_state.get('export_data_seuil_rentabilite_economique', {}),
        'bfr': st.session_state.get('export_data_besoin_fonds_roulement', {}),
        'plan_financement': st.session_state.get('export_data_plan_financement_trois_ans', {}),
        'budget_part1': st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {}),
        'budget_part2': st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})
    }
    
    # Formater le texte final (comme dans Origin.txt)
    final_text = ""
    section_titles = {
        'investissements': "Investissements et financements",
        'salaires': "Salaires et Charges Sociales", 
        'amortissements': "Détail des Amortissements",
        'compte': "Compte de résultats prévisionnel",
        'soldes': "Soldes intermédiaires de gestion",
        'capacite': "Capacité d'autofinancement",
        'seuil': "Seuil de rentabilité économique",
        'bfr': "Besoin en fonds de roulement",
        'plan_financement': "Plan de financement à trois ans",
        'budget_part1': "Budget prévisionnel de trésorerie",
        'budget_part2': "Budget prévisionnel de trésorerie (suite)"
    }
    
    for key, title in section_titles.items():
        if export_data[key]:
            final_text += format_table_data_cyclique(export_data[key], title)
    
    return {
        "raw_data": export_data,
        "formatted_text": final_text,
        "tables_list": list(section_titles.keys())
    }

def format_financial_tables_for_business_plan_cyclique(tableaux_data: Dict[str, Any]) -> str:
    """Formate tous les tableaux financiers en texte pour inclusion dans le business plan (style cyclique)"""
    
    if not tableaux_data:
        return "⚠️ Aucune donnée financière disponible"
    
    formatted_text = "\n\n### TABLEAUX FINANCIERS DÉTAILLÉS\n\n"
    
    # Tableaux principaux à inclure (priorité cyclique Origin.txt)
    tables_config = {
        "compte_resultats_5ans": {
            "title": "Compte de Résultats Prévisionnel (5 ans)",
            "description": "Évolution prévisionnelle des revenus et charges sur 5 années"
        },
        "plan_financement_5ans": {
            "title": "Plan de Financement (5 ans)", 
            "description": "Équilibre emplois/ressources sur 5 années"
        },
        "soldes_intermediaires_5ans": {
            "title": "Soldes Intermédiaires de Gestion",
            "description": "Indicateurs de performance financière"
        },
        "capacite_autofinancement_5ans": {
            "title": "Capacité d'Autofinancement",
            "description": "Ressources internes générées par l'activité"
        },
        "seuil_rentabilite_5ans": {
            "title": "Analyse du Seuil de Rentabilité",
            "description": "Point mort et analyse de rentabilité"
        },
        "bfr_5ans": {
            "title": "Besoin en Fonds de Roulement",
            "description": "Calcul et évolution du BFR"
        }
    }
    
    for table_key, config in tables_config.items():
        if table_key in tableaux_data:
            table_data = tableaux_data[table_key]
            
            formatted_text += f"\n#### {config['title']}\n\n"
            formatted_text += f"*{config['description']}*\n\n"
            
            # Formater le tableau selon son type
            if isinstance(table_data, dict) and "table_data" in table_data:
                formatted_text += format_table_to_markdown(table_data["table_data"])
            elif isinstance(table_data, dict) and "data" in table_data:
                formatted_text += format_table_to_markdown(table_data["data"])
            else:
                formatted_text += "Données non disponibles pour ce tableau\n\n"
            
            formatted_text += "---\n\n"
    
    return formatted_text

def format_table_data_cyclique(data: Dict[str, Any], title: str) -> str:
    """Formate les données de tableau en texte (style Origin.txt)"""
    if not data:
        return ""
    
    text = f"\n\n### {title}\n\n"
    
    # Si c'est une structure de table avec table_data
    if isinstance(data, dict) and "table_data" in data:
        table_data = data["table_data"]
        if isinstance(table_data, list) and table_data:
            # Convertir en DataFrame pour le formatage
            try:
                df = pd.DataFrame(table_data)
                text += df.to_markdown(index=False)
                text += "\n\n"
            except:
                # Fallback simple
                for row in table_data:
                    if isinstance(row, dict):
                        for key, value in row.items():
                            text += f"  {key}: {value}\n"
                    text += "\n"
    
    # Si c'est un dictionnaire simple
    elif isinstance(data, dict):
        for key, value in data.items():
            if key != "table_data":  # Éviter la redondance
                text += f"  **{key}**: {value}\n"
        text += "\n"
    
    return text


# Fonction manquante pour compatibilité avec le système existant
def prepare_section_context(section_name: str, business_data: Dict, financial_tables: Dict, user_input: str) -> str:
    """Prépare le contexte spécifique pour chaque section (compatibilité)"""
    
    base_context = f"""
DONNÉES BUSINESS MODEL: {business_data.get('business_model', '')}

INFORMATIONS GÉNÉRALES: {business_data.get('informations_generales', {})}

INFORMATIONS UTILISATEUR: {user_input}
"""
    
    # Contexte spécifique selon la section
    if section_name == "Plan financier":
        # Pour la section financière, inclure tous les tableaux
        return base_context + f"\n\nTABLEAUX FINANCIERS:\n{financial_tables.get('formatted_text', '')}"
    
    elif section_name in ["Résumé Exécutif", "Présentation de votre entreprise"]:
        # Inclure les données business principales
        return base_context + f"""
PERSONA: {business_data.get('persona', {})}
ANALYSE MARCHÉ: {business_data.get('marche', {})}
CONCURRENCE: {business_data.get('concurrence', {})}
"""
    
    else:
        return base_context

def prepare_section_context(section_name: str, business_data: Dict, financial_tables: Dict, user_input: str) -> str:
    """Prépare le contexte spécifique pour chaque section"""
    
    base_context = f"""
DONNÉES BUSINESS MODEL: {business_data.get('business_model', '')}

INFORMATIONS GÉNÉRALES: {business_data.get('informations_generales', {})}

INFORMATIONS UTILISATEUR: {user_input}
"""
    
    # Contexte spécifique selon la section
    if section_name == "Plan financier":
        # Pour la section financière, inclure tous les tableaux
        return base_context + f"\n\nTABLEAUX FINANCIERS:\n{financial_tables.get('formatted_text', '')}"
    
    elif section_name in ["Résumé Exécutif", "Présentation de votre entreprise"]:
        # Inclure les données business principales
        return base_context + f"""
PERSONA: {business_data.get('persona', {})}
ANALYSE MARCHÉ: {business_data.get('marche', {})}
CONCURRENCE: {business_data.get('concurrence', {})}
"""
    
    else:
        return base_context

def get_business_plan_sections_by_template(template_nom="COPA TRANSFORME") -> Dict[str, Dict[str, str]]:
    """
    Définit les sections du business plan selon le template sélectionné avec système professionnel
    Utilise le nouveau système business_plan_prompts.py
    """
    try:
        # Utiliser le nouveau système professionnel
        sections_config = get_sections_configuration(template_nom)
        return sections_config
        
    except Exception as e:
        st.warning(f"Erreur lors du chargement de la configuration du template {template_nom}: {str(e)}")
        
        # Fallback vers l'ancien système
        return get_fallback_sections_configuration(template_nom)

def get_fallback_sections_configuration(template_nom: str) -> Dict[str, Dict[str, str]]:
    """
    Configuration de fallback utilisant l'ancien système template_manager
    """
    # Récupérer les messages système de l'ancien template
    old_system_messages = get_system_messages(template_nom)
    
    # Structure de base des sections
    base_sections = {
        "Couverture": {
            "system_message": f"Générer une couverture professionnelle pour le template {template_nom}",
            "user_query": "Créer une page de couverture professionnelle"
        },
        "Sommaire": {
            "system_message": f"Générer un sommaire structuré pour le template {template_nom}",
            "user_query": "Afficher le sommaire du business plan"
        },
        "Résumé Exécutif": {
            "system_message": old_system_messages.get("business_plan", f"Générer un résumé exécutif pour {template_nom}"),
            "user_query": "Décrire brièvement le projet et son potentiel"
        },
        "Présentation de votre entreprise": {
            "system_message": old_system_messages.get("business_plan", f"Présenter l'entreprise selon {template_nom}"),
            "user_query": "Présenter l'entreprise de façon complète"
        }
    }
    
    return base_sections


def create_export_files_cyclique(results: Dict[str, str], business_data: Dict[str, Any], template_nom: str):
    """Crée les fichiers d'export avec style cyclique Origin.txt - VERSION CORRIGÉE"""
    
    if not results:
        st.warning("⚠️ Aucun contenu à exporter. Générez d'abord le business plan.")
        return
    
    st.subheader("📥 Téléchargements")
    
    # Créer le contenu complet sans doublons
    complete_content = f"""# Business Plan Complet - Template {template_nom}

**Date de génération :** {pd.Timestamp.now().strftime('%d/%m/%Y %H:%M')}  
**Entreprise :** {business_data.get('informations_generales', {}).get('nom_entreprise', 'Non spécifiée')}  
**Template utilisé :** {template_nom}  

---

"""
    
    # Ajouter toutes les sections générées SANS duplication
    sections_order = ["Couverture", "Sommaire", "Résumé Exécutif", "Présentation de votre entreprise", 
                     "Présentation de l'offre de produit", "Étude de marché", "Stratégie Marketing",
                     "Moyens de production et organisation", "Étude des risques", "Plan financier", "Annexes"]
    
    for section_name in sections_order:
        if section_name in results and results[section_name]:
            # Nettoyer le contenu avant ajout
            clean_content = clean_generated_content_origin_style(results[section_name], section_name)
            complete_content += f"\n\n{clean_content}\n\n---\n"
    
    # Options d'export en colonnes
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### 📄 Format Word")
        if st.button("📄 Générer Word", key="btn_word"):
            try:
                word_buffer = generate_word_document_cyclique(results, business_data, template_nom)
                st.download_button(
                    label="⬇️ Télécharger Business Plan.docx",
                    data=word_buffer,
                    file_name=f"business_plan_{template_nom.lower().replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word"
                )
                st.success("✅ Document Word généré")
            except Exception as e:
                st.error(f"❌ Erreur génération Word : {str(e)}")
    
    with col2:
        st.markdown("### 📑 Format Markdown")
        if st.button("📑 Générer MD", key="btn_md"):
            st.download_button(
                label="⬇️ Télécharger Business Plan.md",
                data=complete_content,
                file_name=f"business_plan_{template_nom.lower().replace(' ', '_')}.md",
                mime="text/markdown",
                key="download_md"
            )
            st.success("✅ Document Markdown prêt")
    
    with col3:
        st.markdown("### 📊 Données JSON")
        if st.button("📊 Générer JSON", key="btn_json"):
            try:
                export_data = {
                    "template": template_nom,
                    "timestamp": pd.Timestamp.now().isoformat(),
                    "business_data": business_data,
                    "sections": {k: clean_generated_content_origin_style(v, k) for k, v in results.items()},
                    "metadata": {
                        "total_sections": len(results),
                        "enterprise_name": business_data.get('informations_generales', {}).get('nom_entreprise', 'Non spécifiée')
                    }
                }
                
                import json
                json_data = json.dumps(export_data, indent=2, ensure_ascii=False, default=str)
                
                st.download_button(
                    label="⬇️ Télécharger Données.json",
                    data=json_data,
                    file_name=f"business_data_{template_nom.lower().replace(' ', '_')}.json",
                    mime="application/json",
                    key="download_json"
                )
                st.success("✅ Données JSON prêtes")
            except Exception as e:
                st.error(f"❌ Erreur export JSON : {str(e)}")
    
    # Aperçu du contenu
    st.markdown("### 📋 Aperçu du contenu généré")
    with st.expander("Voir le contenu complet", expanded=False):
        st.markdown(complete_content)

def generate_word_document_cyclique(results: Dict[str, str], business_data: Dict[str, Any], template_nom: str):
    """Génère un document Word avec style cyclique et template"""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from io import BytesIO
    import re
    
    doc = Document()
    
    # En-tête avec informations du template
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"Business Plan - Template {template_nom}"
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Titre principal
    title = doc.add_heading(f'Business Plan Complet', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Sous-titre avec template
    subtitle = doc.add_paragraph(f"Template : {template_nom}")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].bold = True
    
    # Informations métadonnées
    meta_para = doc.add_paragraph()
    meta_para.add_run("Date de génération : ").bold = True
    meta_para.add_run(pd.Timestamp.now().strftime('%d/%m/%Y %H:%M'))
    meta_para.add_run("\nEntreprise : ").bold = True
    meta_para.add_run(business_data.get('informations_generales', {}).get('nom_entreprise', 'Non spécifiée'))
    
    doc.add_page_break()
    
    # Ajouter chaque section avec traitement markdown basique
    for section_name, content in results.items():
        
        # Titre de section
        doc.add_heading(section_name, 1)
        
        # Traitement du contenu markdown (adapté d'Origin.txt)
        lines = content.split('\n')
        table_data = []
        inside_table = False
        
        for line in lines:
            line = line.strip()
            if not line:
                if table_data:
                    add_table_with_borders_cyclique(doc, table_data)
                    table_data = []
                    inside_table = False
                continue
            
            if line.startswith('###'):
                doc.add_heading(line[3:].strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line[2:].strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line[1:].strip(), level=1)
            elif line.startswith('|') and '|' in line:
                # Tableau détecté
                if not inside_table:
                    inside_table = True
                if not line.replace('|', '').replace('-', '').replace(' ', '').strip():
                    continue  # Ligne de séparation
                table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
            elif re.match(r'^\d+\.', line):
                doc.add_paragraph(line, style='List Number')
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                if inside_table and table_data:
                    add_table_with_borders_cyclique(doc, table_data)
                    table_data = []
                    inside_table = False
                
                # Traitement du texte en gras **texte**
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)
        
        # Traiter la dernière table si elle existe
        if table_data:
            add_table_with_borders_cyclique(doc, table_data)
    
    # Sauvegarder en buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def add_table_with_borders_cyclique(doc, table_data):
    """Ajoute un tableau avec bordures au document Word (style Origin.txt)"""
    if not table_data:
        return
    
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    num_cols = len(table_data[0])
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(table_data):
        for j, cell_content in enumerate(row):
            if j < len(table.columns):
                cell = table.cell(i, j)
                cell.text = str(cell_content)
                
                # Style de la première ligne (en-tête)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def process_uploaded_pdf(uploaded_file):
    """Traite un fichier PDF uploadé (comme dans Origin.txt)"""
    try:
        # Sauvegarder temporairement le fichier
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_file_path = tmp_file.name
        
        # Charger et diviser le document (utilise les fonctions LangChain existantes)
        from services.ai.content_generation import load_and_split_documents
        documents = load_and_split_documents(tmp_file_path)
        
        # Nettoyer le fichier temporaire
        os.unlink(tmp_file_path)
        
        return documents
        
    except Exception as e:
        st.error(f"Erreur lors du traitement du PDF : {str(e)}")
        return []

def process_uploaded_pdf(uploaded_file):
    """Traite un fichier PDF uploadé"""
    # À implémenter si nécessaire
    return []