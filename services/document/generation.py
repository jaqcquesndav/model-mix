"""
Service de génération et export de documents
"""

import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import pandas as pd
from datetime import datetime, date
from typing import Dict, List, Any
import re

def generer_docx_business_model(nom_entreprise: str, date_creation: date, business_model: str, doc: Document, value: int = 1) -> Document:
    """
    Génère un document Word contenant le business model
    
    Args:
        nom_entreprise (str): Nom de l'entreprise
        date_creation (date): Date de création
        business_model (str): Contenu du business model
        doc (Document): Document Word existant
        value (int): Mode de génération
    
    Returns:
        Document: Document Word modifié
    """
    if value == 1:  # Nouvelle page
        doc.add_page_break()
    
    # Titre principal
    title = doc.add_heading('Business Model Canvas', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Informations entreprise
    doc.add_heading('Informations Générales', level=1)
    info_para = doc.add_paragraph()
    info_para.add_run(f"Entreprise : ").bold = True
    info_para.add_run(nom_entreprise)
    info_para.add_run(f"\nDate de création : ").bold = True
    info_para.add_run(date_creation.strftime("%d/%m/%Y"))
    
    # Contenu du business model
    doc.add_heading('Business Model Canvas Détaillé', level=1)
    
    # Diviser le contenu en sections
    sections = business_model.split('\n\n')
    for section in sections:
        if section.strip():
            # Détecter les titres (lignes commençant par #, ** ou des numéros)
            if (section.startswith('#') or 
                section.startswith('**') and section.endswith('**') or
                re.match(r'^\d+\.', section.strip())):
                
                # Titre de section
                titre = section.strip().lstrip('#').strip('*').strip()
                doc.add_heading(titre, level=2)
            else:
                # Contenu normal
                paragraphe = doc.add_paragraph(section.strip())
    
    return doc

def markdown_to_word_via_text(markdown_content: str, doc: Document) -> Document:
    """
    Convertit le contenu Markdown en document Word
    
    Args:
        markdown_content (str): Contenu au format Markdown
        doc (Document): Document Word existant
    
    Returns:
        Document: Document Word modifié
    """
    lines = markdown_content.split('\n')
    table_data = []
    inside_table = False

    for line in lines:
        line = line.strip()
        if not line:
            # Si ligne vide et données de table en cours, ajouter le tableau au document
            if table_data:
                add_table_with_borders(doc, table_data)
                table_data = []
                inside_table = False
            continue

        if line.startswith('# '):  # Titre niveau 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):  # Titre niveau 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):  # Titre niveau 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):  # Titre niveau 4
            doc.add_heading(line[5:], level=4)
        elif re.match(r'^\d+\.\s', line):  # Liste numérotée
            # Vérifier s'il y a du texte en gras dans la liste numérotée
            match = re.match(r'^(\d+\.\s)(\*\*.+?\*\*)', line)
            if match:
                paragraph = doc.add_paragraph(style='List Number')
                paragraph.add_run(match.group(1))  # Numéro
                bold_run = paragraph.add_run(match.group(2)[2:-2])  # Texte en gras sans `**`
                bold_run.bold = True
            else:
                doc.add_paragraph(line, style='List Number')
        elif line.startswith('- ') or line.startswith('•'):  # Liste à puces
            match = re.match(r'^(•|-)\s\*\*(.+?)\*\*(.*)', line)
            if match:
                paragraph = doc.add_paragraph(style='List Bullet')
                bold_run = paragraph.add_run(match.group(2))  # Texte en gras
                bold_run.bold = True
                if match.group(3):  # Texte après le gras
                    paragraph.add_run(match.group(3).strip())
            else:
                doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):  # Détection des lignes de tableau
            if re.match(r'\|?\s*[-:]+\s*\|', line):
                inside_table = True
                continue  # Ignorer les lignes de séparation
            else:
                inside_table = True
                table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
        elif re.match(r'^\*\*.+?\*\*\s*:', line):  # Texte en gras suivi de texte normal
            match = re.match(r'^\*\*(.+?)\*\*\s*:(.*)', line)
            if match:
                paragraph = doc.add_paragraph()
                bold_run = paragraph.add_run(match.group(1))  # Texte en gras
                bold_run.bold = True
                if match.group(2):  # Texte normal après le `:`
                    paragraph.add_run(f":{match.group(2)}")
        elif re.match(r'^\*\*.+?\*\*$', line):  # Texte entièrement en gras
            paragraph = doc.add_paragraph()
            bold_run = paragraph.add_run(line[2:-2])  # Texte sans `**`
            bold_run.bold = True
        elif not inside_table:  # Paragraphe normal
            doc.add_paragraph(line)

    # Traiter les données de table restantes
    if table_data:
        add_table_with_borders(doc, table_data)
        
    return doc

def add_table_with_borders(doc: Document, table_data: List[List[str]]):
    """
    Ajoute un tableau au document Word avec bordures
    
    Args:
        doc (Document): Document Word
        table_data (list): Données du tableau
    """
    if not table_data:
        return
        
    num_cols = len(table_data[0])
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'  # Appliquer un style de tableau avec bordures

    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            if j < len(table.columns):  # Vérifier que la colonne existe
                cell_content = table.cell(i, j).paragraphs[0]
                parts = re.split(r'(\*\*.+?\*\*)', cell)  # Diviser par texte en gras
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):  # Texte en gras
                        run = cell_content.add_run(part[2:-2])
                        run.bold = True
                    else:  # Texte normal
                        cell_content.add_run(part.strip())

def ajouter_tableau_financier(doc: Document, donnees: Dict[str, Any], headers: List[str], titre: str):
    """
    Ajoute un tableau financier au document Word
    
    Args:
        doc (Document): Document Word
        donnees (dict): Données du tableau
        headers (list): En-têtes du tableau
        titre (str): Titre du tableau
    """
    doc.add_heading(titre, level=2)
    
    if not donnees.get("table_data"):
        doc.add_paragraph("Aucune donnée disponible pour cette section.")
        return
    
    table_data = donnees["table_data"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ajouter les en-têtes
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Ajouter les données des tableaux
    for row in table_data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            cell_value = row.get(header, "")
            cell_text = str(cell_value)  # Convertir en chaîne de caractères
            row_cells[i].text = cell_text
            row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")

def format_table_to_markdown(table_data: List[Dict[str, Any]], title: str = "") -> str:
    """
    Formate des données de tableau en Markdown
    
    Args:
        table_data (list): Liste de dictionnaires représentant les lignes du tableau
        title (str): Titre optionnel du tableau
    
    Returns:
        str: Tableau formaté en Markdown
    """
    if not table_data:
        return f"### {title}\n\nAucune donnée disponible.\n\n" if title else "Aucune donnée disponible.\n\n"
    
    markdown = ""
    if title:
        markdown += f"### {title}\n\n"
    
    # Récupérer les en-têtes (clés du premier dictionnaire)
    headers = list(table_data[0].keys())
    
    # Ligne d'en-têtes
    markdown += "| " + " | ".join(headers) + " |\n"
    
    # Ligne de séparation
    markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    
    # Lignes de données
    for row in table_data:
        values = []
        for header in headers:
            value = row.get(header, "")
            # Formater les valeurs numériques
            if isinstance(value, (int, float)):
                if abs(value) >= 1000:
                    value = f"{value:,.0f}".replace(",", " ")
                else:
                    value = f"{value:.2f}"
            values.append(str(value))
        markdown += "| " + " | ".join(values) + " |\n"
    
    markdown += "\n"
    return markdown

def generer_markdown(resultats: Dict[str, str]) -> str:
    """
    Génère le contenu Markdown à partir des résultats
    
    Args:
        resultats (dict): Résultats des sections générées
    
    Returns:
        str: Contenu Markdown complet
    """
    markdown_content = ""
    
    for section_name, content in resultats.items():
        markdown_content += f"\n\n# {section_name}\n\n"
        markdown_content += content
        markdown_content += "\n\n---\n"
    
    return markdown_content

def exporter_donnees_json(donnees: Dict[str, Any]) -> str:
    """
    Exporte les données au format JSON
    
    Args:
        donnees (dict): Données à exporter
    
    Returns:
        str: JSON formaté
    """
    import json
    from datetime import datetime, date
    
    def json_serializer(obj):
        if isinstance(obj, (datetime, date)):
            return obj.isoformat()
        return str(obj)
    
    return json.dumps(donnees, default=json_serializer, indent=2, ensure_ascii=False)

def generer_rapport_excel(donnees_financieres: Dict[str, Any]) -> BytesIO:
    """
    Génère un rapport Excel avec tous les tableaux financiers
    
    Args:
        donnees_financieres (dict): Données financières
    
    Returns:
        BytesIO: Fichier Excel en mémoire
    """
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        
        # Feuille de synthèse
        synthese_data = {
            'Indicateur': ['Date de génération', 'Entreprise', 'Total Investissements'],
            'Valeur': [
                datetime.now().strftime("%d/%m/%Y %H:%M"),
                st.session_state.get('nom_entreprise', 'N/A'),
                donnees_financieres.get('investissements', {}).get('total_investissement', 0)
            ]
        }
        
        df_synthese = pd.DataFrame(synthese_data)
        df_synthese.to_excel(writer, sheet_name='Synthèse', index=False)
        
        # Ajouter chaque tableau financier
        tableaux = {
            'Investissements': donnees_financieres.get('investissements', {}),
            'Compte Résultats': donnees_financieres.get('compte_resultats', {}),
            'Soldes Intermédiaires': donnees_financieres.get('soldes_intermediaires', {}),
            'CAF': donnees_financieres.get('capacite_autofinancement', {}),
            'Seuil Rentabilité': donnees_financieres.get('seuil_rentabilite', {}),
            'BFR': donnees_financieres.get('bfr', {}),
            'Plan Financement': donnees_financieres.get('plan_financement', {}),
        }
        
        for nom_feuille, donnees_tableau in tableaux.items():
            if donnees_tableau.get('table_data'):
                df = pd.DataFrame(donnees_tableau['table_data'])
                df.to_excel(writer, sheet_name=nom_feuille, index=False)
    
    output.seek(0)
    return output

def consolider_donnees_financieres() -> Dict[str, Any]:
    """
    Consolide toutes les données financières du session state
    
    Returns:
        dict: Données consolidées avec synthèse
    """
    donnees = {
        'investissements': st.session_state.get('export_data_investissements', {}),
        'salaires': st.session_state.get('export_data_salaires', {}),
        'amortissements': st.session_state.get('export_data_amortissements', {}),
        'compte_resultats': st.session_state.get('export_data_compte', {}),
        'soldes': st.session_state.get('export_data_soldes', {}),
        'capacite': st.session_state.get('export_data_capacite', {}),
        'seuil': st.session_state.get('export_data_seuil', {}),
        'bfr': st.session_state.get('export_data_bfr', {}),
        'plan_financement': st.session_state.get('export_data_plan_financement', {}),
        'budget_part1': st.session_state.get('export_data_budget_part1', {}),
        'budget_part2': st.session_state.get('export_data_budget_part2', {})
    }
    
    # Créer une synthèse financière
    synthese = f"""
    SYNTHÈSE FINANCIÈRE CONSOLIDÉE:
    
    Total Investissements: {sum([float(str(v).replace('$', '').replace(',', '').strip()) for v in donnees['investissements'].values() if isinstance(v, (int, float, str)) and str(v).replace('$', '').replace(',', '').replace('.', '').isdigit()], start=0)} USD
    
    Analyse de Rentabilité: {donnees.get('seuil', {}).get('point_mort', 'Non calculé')}
    
    Besoins de Financement: {donnees.get('bfr', {}).get('total_bfr', 'Non calculé')}
    
    Capacité d'Autofinancement: {donnees.get('capacite', {}).get('caf_net', 'Non calculé')}
    
    Contexte RDC: Tous les montants sont exprimés en USD pour faciliter l'analyse économique et les comparaisons. 
    L'environnement économique congolais nécessite une attention particulière aux fluctuations monétaires et aux défis logistiques.
    """
    
    return {
        'donnees': donnees,
        'synthese': synthese
    }